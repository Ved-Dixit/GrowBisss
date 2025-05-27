import streamlit as st
import psycopg2
from psycopg2 import sql
import os
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from transformers.pipelines import pipeline
import pandas as pd
import numpy as np
from datetime import datetime, timedelta, timezone, date
import hashlib
import jwt
import time
import docx
from docx.shared import Inches
from io import BytesIO
import pytz
from gtts import gTTS
import speech_recognition as sr
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
from PIL import Image
import json
from fpdf import FPDF
from streamlit_webrtc import webrtc_streamer, AudioProcessorBase, WebRtcMode
import av 
#
# Load environment variables
# Note: Ensure JWT_SECRET is set in your environment variables
# os.environ['JWT_SECRET'] = 'your_super_secret_key_here' # Example, replace with actual secret

# Database Connection
def get_db_connection():
    """Establishes and returns a database connection."""
    # Replace with your actual Railway PostgreSQL connection details
    conn = psycopg2.connect(
        dbname = "railway",
        user = "postgres",
        password = "AtyqCUijTHXVXFHkNTnVXzwEwccYCfhr",
        host = "switchyard.proxy.rlwy.net",
        port = "36704"

    )
    return conn

# Initialize database tables
def init_db():
    """Creates necessary database tables if they don't exist."""
    conn = get_db_connection()
    cur = conn.cursor()
    
    # Create tables if they don't exist
    tables = [
        """
        CREATE TABLE IF NOT EXISTS businesses (
            id SERIAL PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            email VARCHAR(100) UNIQUE NOT NULL,
            password_hash VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            subscription_type VARCHAR(50) DEFAULT 'freemium',
            subscription_expiry DATE DEFAULT (CURRENT_DATE + INTERVAL '3 months')
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS messages (
            id SERIAL PRIMARY KEY,
            sender_type VARCHAR(30) NOT NULL, -- 'business', 'investor', 'service_provider'
            sender_id INTEGER NOT NULL,
            receiver_type VARCHAR(30) NOT NULL,
            receiver_id INTEGER NOT NULL,
            content TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            read_at TIMESTAMP 
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS products (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            name VARCHAR(100) NOT NULL,
            description TEXT,
            price DECIMAL(10,2) NOT NULL,
            quantity INTEGER NOT NULL,
            category VARCHAR(50),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS employees (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            name VARCHAR(100) NOT NULL,
            email VARCHAR(100) NOT NULL,
            position VARCHAR(100),
            department VARCHAR(100),
            salary DECIMAL(12,2),
            join_date DATE,
            last_appraisal_date DATE,
            performance_score INTEGER,
            skills TEXT[],
            UNIQUE(business_id, email)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS projects (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            name VARCHAR(100) NOT NULL,
            description TEXT,
            client VARCHAR(100),
            start_date DATE,
            end_date DATE,
            budget DECIMAL(12,2),
            status VARCHAR(50),
            manager_id INTEGER REFERENCES employees(id),
            progress INTEGER DEFAULT 0
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS documents (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            title VARCHAR(200) NOT NULL,
            content TEXT,
            doc_type VARCHAR(50),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            created_by INTEGER REFERENCES employees(id)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS market_data (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            industry VARCHAR(100),
            metric VARCHAR(100),
            value DECIMAL(12,2),
            date DATE,
            source VARCHAR(100)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS project_assignments (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            project_id INTEGER REFERENCES projects(id) ON DELETE CASCADE,
            employee_id INTEGER REFERENCES employees(id) ON DELETE CASCADE,
            assigned_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(business_id, project_id, employee_id)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS schemes (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            name VARCHAR(200) NOT NULL,
            description TEXT,
            eligibility TEXT,
            benefits TEXT,
            deadline DATE,
            sector VARCHAR(100),
            is_govt BOOLEAN
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS opportunities (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            title VARCHAR(200) NOT NULL,
            description TEXT,
            category VARCHAR(100),
            deadline DATE,
            reward TEXT,
            link VARCHAR(200)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS tax_records (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            financial_year VARCHAR(20),
            total_income DECIMAL(12,2),
            tax_paid DECIMAL(12,2),
            filing_date DATE,
            status VARCHAR(50),
            notes TEXT
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS ipo_data (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            company_name VARCHAR(100),
            issue_size DECIMAL(12,2),
            price_range VARCHAR(50),
            open_date DATE,
            close_date DATE,
            status VARCHAR(50),
            allotment_date DATE,
            listing_date DATE
        )
        """,
        # This is the more complete definition for service_providers
        """
        CREATE TABLE IF NOT EXISTS service_providers (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE SET NULL, 
            name VARCHAR(100) NOT NULL,
            service_type VARCHAR(100) NOT NULL, 
            email VARCHAR(100) UNIQUE NOT NULL, 
            password_hash VARCHAR(255),    
            contact_email VARCHAR(100),   
            rating DECIMAL(3,1),
            experience_years INTEGER,
            pricing TEXT,
            availability BOOLEAN,
            profile_description TEXT,       
            specializations TEXT[],       
            office_address TEXT,          
            website_url VARCHAR(255),     
            linkedin_profile VARCHAR(255),   
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP 
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS attendance (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            employee_id INTEGER REFERENCES employees(id),
            date DATE,
            status VARCHAR(20),
            check_in TIME,
            check_out TIME
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS job_openings (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            title VARCHAR(100),
            department VARCHAR(100),
            description TEXT,
            requirements TEXT[],
            experience_needed VARCHAR(50),
            posted_date DATE,
            status VARCHAR(20)
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS invoices (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            invoice_number VARCHAR(50) UNIQUE,
            customer_name VARCHAR(100),
            customer_email VARCHAR(100),
            issue_date DATE,
            due_date DATE,
            total_amount DECIMAL(12,2),
            tax_amount DECIMAL(12,2),
            status VARCHAR(20) DEFAULT 'pending',
            items JSONB
        )
        """,
        # This is the more complete definition for investors
         """
        CREATE TABLE IF NOT EXISTS investors (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE SET NULL, 
            name VARCHAR(100) NOT NULL,
            firm VARCHAR(100),
            email VARCHAR(100) UNIQUE NOT NULL, 
            password_hash VARCHAR(255),         
            investment_focus VARCHAR(200),
            portfolio_companies TEXT[],
            last_contact DATE,
            profile_description TEXT,          
            website_url VARCHAR(255),         
            linkedin_profile VARCHAR(255),     
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP 
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS user_sessions (
            id SERIAL PRIMARY KEY,
            business_id INTEGER REFERENCES businesses(id) ON DELETE CASCADE,
            session_token VARCHAR(255) UNIQUE,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            expires_at TIMESTAMP
        )
        """
    ]
    
    for table_sql in tables:
        try:
            cur.execute(table_sql)
        except Exception as e:
            st.error(f"Error creating table: {e}\nSQL: {table_sql[:200]}...") # Log part of SQL
    
    # Add indexes for messages table separately
    message_indexes = [
        "CREATE INDEX IF NOT EXISTS idx_messages_conversation_pair1 ON messages (sender_type, sender_id, receiver_type, receiver_id, created_at);",
        "CREATE INDEX IF NOT EXISTS idx_messages_conversation_pair2 ON messages (receiver_type, receiver_id, sender_type, sender_id, created_at);"
    ]
    for index_sql in message_indexes:
        try:
            cur.execute(index_sql)
        except Exception as e:
            st.error(f"Error creating message index: {e}")

    conn.commit()
    cur.close()
    conn.close()

# Security functions
def hash_password(password):
    """Hashes a password using SHA-256."""
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(plain_password, hashed_password):
    """Verifies a plain password against a hashed password."""
    return hash_password(plain_password) == hashed_password

def generate_jwt(business_id):
    """Generates a JWT for a business user."""
    payload = {
        'business_id': business_id,
        'exp': datetime.utcnow() + timedelta(hours=24) # Token expires in 24 hours
    }
    # Ensure JWT_SECRET environment variable is set
    return jwt.encode(payload, os.getenv('JWT_SECRET', 'secret_key'), algorithm='HS256')

def verify_jwt(token):
    """Verifies a JWT and returns the business_id if valid."""
    try:
        payload = jwt.decode(token, os.getenv('JWT_SECRET', 'secret_key'), algorithms=['HS256'])
        return payload.get('business_id')
    except jwt.ExpiredSignatureError:
        # st.warning("Session expired. Please log in again.") # Avoid showing this on every check
        return None
    except jwt.InvalidTokenError:
        # st.warning("Invalid token. Please log in again.") # Avoid showing this on every check
        return None
    except Exception as e:
        # st.error(f"JWT verification error: {e}") # Log unexpected errors
        return None

# Initialize AI models
class AIModels:
    """Manages lazy loading of AI models."""
    def __init__(self):
        st.write("AIModels class instance created. Models will be loaded on first use.")
        self._chatbot_tokenizer_instance = None
        self._chatbot_model_instance = None
        self._text_generator_instance = None
        self._sentiment_analyzer_instance = None
        self._translator_en_hi_instance = None
        self._translator_hi_en_instance = None

    @property
    def chatbot_tokenizer(self):
        if self._chatbot_tokenizer_instance is None:
            st.write("Loading chatbot_tokenizer (facebook/blenderbot-400M-distill)...")
            self._chatbot_tokenizer_instance = AutoTokenizer.from_pretrained("facebook/blenderbot-400M-distill")
            st.write("Chatbot_tokenizer loaded.")
        return self._chatbot_tokenizer_instance

    @property
    def chatbot_model(self):
        if self._chatbot_model_instance is None:
            st.write("Loading chatbot_model (facebook/blenderbot-400M-distill)...")
            self._chatbot_model_instance = AutoModelForSeq2SeqLM.from_pretrained("facebook/blenderbot-400M-distill")
            st.write("Chatbot_model loaded.")
        return self._chatbot_model_instance

    @property
    def text_generator(self):
        if self._text_generator_instance is None:
            st.write("Loading text_generator (gpt2)...")
            self._text_generator_instance = pipeline("text-generation", model="gpt2")
            st.write("Text_generator (gpt2) loaded.")
        return self._text_generator_instance

    @property
    def sentiment_analyzer(self):
        if self._sentiment_analyzer_instance is None:
            st.write("Loading sentiment_analyzer (distilbert-base-uncased-finetuned-sst-2-english)...")
            self._sentiment_analyzer_instance = pipeline("sentiment-analysis", model="distilbert-base-uncased-finetuned-sst-2-english")
            st.write("Sentiment_analyzer loaded.")
        return self._sentiment_analyzer_instance

    @property
    def translator_en_hi(self):
        if self._translator_en_hi_instance is None:
            st.write("Loading translator_en_hi (Helsinki-NLP/opus-mt-en-hi)...")
            self._translator_en_hi_instance = pipeline("translation", model="Helsinki-NLP/opus-mt-en-hi")
            st.write("Translator_en_hi loaded.")
        return self._translator_en_hi_instance

    @property
    def translator_hi_en(self):
        if self._translator_hi_en_instance is None:
            st.write("Loading translator_hi_en (Helsinki-NLP/opus-mt-hi-en)...")
            self._translator_hi_en_instance = pipeline("translation", model="Helsinki-NLP/opus-mt-hi-en")
            st.write("Translator_hi_en loaded.")
        return self._translator_hi_en_instance

    def generate_response(self, prompt):
        """Generates a response using the chatbot model."""
        inputs = self.chatbot_tokenizer([prompt], return_tensors="pt")
        reply_ids = self.chatbot_model.generate(**inputs)
        return self.chatbot_tokenizer.batch_decode(reply_ids, skip_special_tokens=True)[0]

    def generate_text(self, prompt, max_length=150):
        """Generates text using the text generation model."""
        return self.text_generator(prompt, max_length=max_length, num_return_sequences=1)[0]['generated_text']

    def analyze_sentiment(self, text):
        """Analyzes sentiment of text."""
        return self.sentiment_analyzer(text)

    def translate(self, text, target_lang):
        """Translates text between English and Hindi."""
        if target_lang == "Hindi":
            return self.translator_en_hi(text)[0]['translation_text']
        elif target_lang == "English":
            return self.translator_hi_en(text)[0]['translation_text']
        return text

@st.cache_resource
def load_ai_models():
    """Caches and loads AI models."""
    models = AIModels()
    return models

# Registration Pages for Investor and Service Provider
def investor_registration_page():
    """Streamlit page for Investor registration."""
    st.subheader("Investor Registration")
    with st.form("investor_register_form"):
        name = st.text_input("Full Name / Firm Name")
        email = st.text_input("Login Email")
        password = st.text_input("Password", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")
        firm = st.text_input("Firm (if applicable)")
        investment_focus = st.text_area("Investment Focus (e.g., SaaS, Fintech, Early Stage)")
        profile_description = st.text_area("Brief Profile / About Us")
        website_url = st.text_input("Website URL (Optional)")
        linkedin_profile = st.text_input("LinkedIn Profile URL (Optional)")

        submitted = st.form_submit_button("Register as Investor")
        if submitted:
            if password != confirm_password:
                st.error("Passwords do not match.")
            elif not name or not email or not password:
                st.error("Name, Email, and Password are required.")
            else:
                conn = get_db_connection()
                cur = conn.cursor()
                try:
                    hashed_pw = hash_password(password)
                    cur.execute(
                        """INSERT INTO investors (name, email, password_hash, firm, investment_focus, profile_description, website_url, linkedin_profile, business_id)
                           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NULL) RETURNING id""", # business_id is NULL for independent registration
                        (name, email, hashed_pw, firm, investment_focus, profile_description, website_url, linkedin_profile)
                    )
                    investor_id = cur.fetchone()[0]
                    conn.commit()
                    st.success("Investor registration successful! You can now login.")
                    st.session_state.user_type_for_login = 'investor' 
                    st.session_state.registered_email = email 
                    st.session_state.show_login_main = True # Signal to show main login
                    st.rerun()
                except psycopg2.IntegrityError:
                    st.error("This email is already registered as an investor.")
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                finally:
                    cur.close()
                    conn.close()

def service_provider_registration_page():
    """Streamlit page for Service Provider registration."""
    st.subheader("Service Provider Registration (CA, Legal, etc.)")
    with st.form("service_provider_register_form"):
        name = st.text_input("Full Name / Firm Name")
        email = st.text_input("Login Email")
        password = st.text_input("Password", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")
        
        service_type_options = ["legal", "ca", "insurance", "consulting", "other"] 
        service_type = st.selectbox("Primary Service Type", service_type_options)
        
        contact_email_public = st.text_input("Public Contact Email (Optional, if different from login)")
        profile_description = st.text_area("Profile / Service Description")
        specializations_str = st.text_input("Specializations (comma-separated, e.g., Corporate Law, Tax Audit)")
        experience_years = st.number_input("Years of Experience", min_value=0, step=1)
        office_address = st.text_area("Office Address (Optional)")
        website_url = st.text_input("Website URL (Optional)")
        linkedin_profile = st.text_input("LinkedIn Profile URL (Optional)")

        submitted = st.form_submit_button("Register as Service Provider")
        if submitted:
            if password != confirm_password:
                st.error("Passwords do not match.")
            elif not name or not email or not password or not service_type:
                st.error("Name, Email, Password, and Service Type are required.")
            else:
                conn = get_db_connection()
                cur = conn.cursor()
                try:
                    hashed_pw = hash_password(password)
                    specializations_list = [s.strip() for s in specializations_str.split(',') if s.strip()]
                    cur.execute(
                        """INSERT INTO service_providers 
                           (name, email, password_hash, service_type, contact_email, profile_description, specializations, experience_years, office_address, website_url, linkedin_profile, business_id)
                           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NULL) RETURNING id""", # business_id is NULL
                        (name, email, hashed_pw, service_type, contact_email_public or email, profile_description, specializations_list, experience_years, office_address, website_url, linkedin_profile)
                    )
                    provider_id = cur.fetchone()[0]
                    conn.commit()
                    st.success("Service Provider registration successful! You can now login.")
                    st.session_state.user_type_for_login = 'service_provider'
                    st.session_state.registered_email = email
                    st.session_state.show_login_main = True # Signal to show main login
                    st.rerun()
                except psycopg2.IntegrityError:
                    st.error("This email is already registered as a service provider.")
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                finally:
                    cur.close()
                    conn.close()

# Portal Pages for Investor and Service Provider
def investor_portal(investor_id, ai_models):
    """Streamlit portal for logged-in Investors."""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT name, email, firm, investment_focus, profile_description, website_url, linkedin_profile FROM investors WHERE id = %s", (investor_id,))
    investor_data = cur.fetchone()
    cur.close()
    conn.close()

    if not investor_data:
        st.error("Investor profile not found.")
        logout()
        st.rerun()
        return

    st.title(f"🚀 Investor Portal: {investor_data[0]}")
    
    tab1, tab2 = st.tabs(["My Profile", "Messages"])

    with tab1:
        st.subheader("My Profile")
        st.write(f"**Name/Firm:** {investor_data[0]}")
        st.write(f"**Login Email:** {investor_data[1]}")
        st.write(f"**Registered Firm:** {investor_data[2] or 'N/A'}")
        st.write(f"**Investment Focus:** {investor_data[3] or 'N/A'}")
        st.write(f"**Profile Description:** {investor_data[4] or 'N/A'}")
        st.write(f"**Website:** {investor_data[5] or 'N/A'}")
        st.write(f"**LinkedIn:** {investor_data[6] or 'N/A'}")
        # Add profile edit form here later

    with tab2:
        st.subheader("Messages from Businesses")
        chat_module_for_entity(entity_id=investor_id, entity_type='investor', entity_name=investor_data[0], ai_models=ai_models)


def service_provider_portal(provider_id, ai_models):
    """Streamlit portal for logged-in Service Providers."""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT name, email, service_type, contact_email, profile_description, specializations, experience_years, office_address, website_url, linkedin_profile FROM service_providers WHERE id = %s", (provider_id,))
    provider_data = cur.fetchone()
    cur.close()
    conn.close()

    if not provider_data:
        st.error("Service Provider profile not found.")
        logout()
        st.rerun()
        return

    st.title(f"🚀 {provider_data[2].capitalize()} Portal: {provider_data[0]}")

    tab1, tab2 = st.tabs(["My Profile", "Messages"])
    with tab1:
        st.subheader("My Profile")
        st.write(f"**Name/Firm:** {provider_data[0]}")
        st.write(f"**Login Email:** {provider_data[1]}")
        st.write(f"**Service Type:** {provider_data[2].capitalize()}")
        st.write(f"**Public Contact Email:** {provider_data[3] or 'N/A'}")
        st.write(f"**Profile Description:** {provider_data[4] or 'N/A'}")
        st.write(f"**Specializations:** {', '.join(provider_data[5]) if provider_data[5] else 'N/A'}")
        st.write(f"**Experience:** {provider_data[6] or 'N/A'} years")
        st.write(f"**Office Address:** {provider_data[7] or 'N/A'}")
        st.write(f"**Website:** {provider_data[8] or 'N/A'}")
        st.write(f"**LinkedIn:** {provider_data[9] or 'N/A'}")
        
    with tab2:
        st.subheader("Messages from Businesses")
        chat_module_for_entity(entity_id=provider_id, entity_type='service_provider', entity_name=provider_data[0], ai_models=ai_models)

# Generic Chat Module for Non-Business Entities (Investor/Service Provider)
def chat_module_for_entity(entity_id, entity_type, entity_name, ai_models):
    """Handles messaging interface for Investor and Service Provider users."""
    st.header(f"💬 Messaging for {entity_name}")

    conn = get_db_connection()
    cur = conn.cursor()

    # State variables specific to this entity's chat module
    session_prefix = f"chat_entity_{entity_type}_{entity_id}"
    if f'{session_prefix}_partner_type' not in st.session_state:
        st.session_state[f'{session_prefix}_partner_type'] = None 
    if f'{session_prefix}_partner_id' not in st.session_state:
        st.session_state[f'{session_prefix}_partner_id'] = None
    if f'{session_prefix}_partner_name' not in st.session_state:
        st.session_state[f'{session_prefix}_partner_name'] = None

    st.sidebar.subheader("Conversations")
    try:
        # Businesses that have sent a message to this entity OR this entity has sent a message to
        query_partners = """
        SELECT DISTINCT sender_id, sender_type FROM messages
        WHERE receiver_id = %(current_id)s AND receiver_type = %(current_type)s AND sender_type = 'business'
        UNION
        SELECT DISTINCT receiver_id, receiver_type FROM messages
        WHERE sender_id = %(current_id)s AND sender_type = %(current_type)s AND receiver_type = 'business';
        """
        cur.execute(query_partners, {'current_id': entity_id, 'current_type': entity_type})
        partners_db = cur.fetchall()
        
        partner_options = {} 
        for p_id, p_type_db in partners_db: 
            if p_type_db == 'business':
                cur.execute("SELECT name FROM businesses WHERE id = %s", (p_id,))
                name_result = cur.fetchone()
                if name_result:
                    partner_display_name = f"{name_result[0]} (Business)"
                    partner_options[partner_display_name] = ('business', p_id)
        
        selected_partner_display = st.sidebar.selectbox(
            "Select Business to Chat With:",
            options=["-- Select --"] + list(partner_options.keys()),
            key=f"{session_prefix}_select_chat_partner"
        )

        if selected_partner_display != "-- Select --":
            partner_db_type, partner_id_val = partner_options[selected_partner_display]
            if (st.session_state.get(f'{session_prefix}_partner_id') != partner_id_val or
                st.session_state.get(f'{session_prefix}_partner_type') != partner_db_type):
                st.session_state[f'{session_prefix}_partner_type'] = partner_db_type
                st.session_state[f'{session_prefix}_partner_id'] = partner_id_val
                st.session_state[f'{session_prefix}_partner_name'] = selected_partner_display.split(" (")[0]
                st.rerun()

    except Exception as e_partner_list:
        st.sidebar.error(f"Error loading chat partners: {e_partner_list}")

    partner_selected_id = st.session_state.get(f'{session_prefix}_partner_id')
    partner_selected_name = st.session_state.get(f'{session_prefix}_partner_name')

    if partner_selected_id and partner_selected_name:
        st.subheader(f"Chat with {partner_selected_name} (Business)")
        
        # Mark messages as read (messages sent BY the business partner TO this entity)
        try:
            update_read_query = """
                UPDATE messages SET read_at = CURRENT_TIMESTAMP
                WHERE receiver_type = %s AND receiver_id = %s
                  AND sender_type = 'business' AND sender_id = %s
                  AND read_at IS NULL;
            """
            cur.execute(update_read_query, (entity_type, entity_id, partner_selected_id))
            conn.commit()
        except Exception as e_read_update:
            st.warning(f"Could not update read status: {e_read_update}")

        # Fetch messages
        try:
            query_messages = """
                SELECT sender_type, sender_id, content, created_at, read_at
                FROM messages
                WHERE
                    (sender_type = %(current_entity_type)s AND sender_id = %(current_entity_id)s AND receiver_type = 'business' AND receiver_id = %(partner_business_id)s)
                OR
                    (sender_type = 'business' AND sender_id = %(partner_business_id)s AND receiver_type = %(current_entity_type)s AND receiver_id = %(current_entity_id)s)
                ORDER BY created_at ASC;
            """
            cur.execute(query_messages, {
                'current_entity_type': entity_type, 'current_entity_id': entity_id,
                'partner_business_id': partner_selected_id
            })
            messages = cur.fetchall()
            
            chat_container = st.container()
            with chat_container:
                if not messages: st.info("No messages yet. Start the conversation!")
                for msg_sender_type, msg_sender_id, msg_content, msg_created_at, msg_read_at in messages:
                    timestamp_str = msg_created_at.strftime('%Y-%m-%d %H:%M')
                    if msg_sender_type == entity_type and msg_sender_id == entity_id: 
                        read_indicator = " (Read)" if msg_read_at else " (Delivered)"
                        st.markdown(f"<div style='text-align: right; margin-left: 20%; margin-bottom: 5px; padding: 10px; background-color: #DCF8C6; border-radius: 10px;'><b>You ({entity_name})</b> ({timestamp_str}){read_indicator}:<br>{msg_content}</div>", unsafe_allow_html=True)
                    else: 
                        st.markdown(f"<div style='text-align: left; margin-right: 20%; margin-bottom: 5px; padding: 10px; background-color: #FFFFFF; border-radius: 10px; border: 1px solid #E0E0E0;'><b>{partner_selected_name}</b> ({timestamp_str}):<br>{msg_content}</div>", unsafe_allow_html=True)
            if messages:
                 st.markdown(f"<script>window.scrollTo(0,document.body.scrollHeight);</script>", unsafe_allow_html=True)

        except Exception as e_fetch:
            st.error(f"Error fetching messages: {e_fetch}")

        with st.form(f"{session_prefix}_new_message_form", clear_on_submit=True):
            new_message = st.text_area("Your reply:", key=f"{session_prefix}_new_message_content_{partner_selected_id}")
            send_reply_btn = st.form_submit_button("Send Reply")
            if send_reply_btn and new_message.strip():
                try:
                    cur.execute(
                        "INSERT INTO messages (sender_type, sender_id, receiver_type, receiver_id, content) VALUES (%s, %s, %s, %s, %s)",
                        (entity_type, entity_id, 'business', partner_selected_id, new_message.strip())
                    )
                    conn.commit()
                    st.rerun()
                except Exception as e_send:
                    st.error(f"Error sending reply: {e_send}")
            elif send_reply_btn and not new_message.strip():
                st.warning("Message cannot be empty.")
    else:
        st.info("Select a business from the sidebar to view or continue a conversation.")
    
    cur.close()
    conn.close()

# Login Page
def login_page():
    """Streamlit login page for all user types."""
    st.title("GrowBis Login")

    login_user_type_options = ('Business', 'Investor', 'Service Provider')
    default_login_type_index = 0
    if 'user_type_for_login' in st.session_state and st.session_state.user_type_for_login:
        try:
            default_login_type_index = login_user_type_options.index(st.session_state.user_type_for_login.capitalize())
        except ValueError:
            default_login_type_index = 0 # Default if type from session is not in options

    login_user_type = st.radio(
        "Login as:",
        login_user_type_options,
        key='login_user_type_selection',
        horizontal=True,
        index=default_login_type_index
    )

    default_email = st.session_state.get('registered_email', "")

    if login_user_type == 'Business':
        login_tab, register_tab = st.tabs(["Login to Business Account", "Register New Business"])
        with login_tab:
            with st.form("business_login_form"):
                email = st.text_input("Email", value=default_email if st.session_state.get('user_type_for_login') == 'business' else "")
                password = st.text_input("Password", type="password")
                
                if st.form_submit_button("Login as Business"):
                    conn = get_db_connection()
                    cur = conn.cursor()
                    cur.execute("SELECT id, password_hash, name FROM businesses WHERE email = %s", (email,))
                    result = cur.fetchone()
                    
                    if result and verify_password(password, result[1]):
                        business_id = result[0]
                        token = generate_jwt(business_id)
                        expires_at = datetime.utcnow() + timedelta(hours=24) # Match JWT expiry
                        # Store or update session in database
                        cur.execute(
                            "INSERT INTO user_sessions (business_id, session_token, expires_at) VALUES (%s, %s, %s) ON CONFLICT (session_token) DO UPDATE SET expires_at = EXCLUDED.expires_at, business_id = EXCLUDED.business_id",
                            (business_id, token, expires_at)
                        )
                        conn.commit()
                        st.session_state.token = token
                        st.session_state.business_id = business_id # Legacy, ensure it's set
                        st.session_state.user_type = 'business'
                        st.session_state.user_name = result[2]
                        st.session_state.logged_in_entity_id = business_id
                        if 'registered_email' in st.session_state: del st.session_state.registered_email
                        if 'user_type_for_login' in st.session_state: del st.session_state.user_type_for_login
                        st.rerun()
                    else:
                        st.error("Invalid email or password for Business.")
                    cur.close()
                    conn.close()
        with register_tab:
            st.subheader("Register New Business")
            with st.form("business_register_form"):
                name_reg = st.text_input("Business Name")
                email_reg = st.text_input("Business Email")
                password_reg = st.text_input("Password", type="password")
                confirm_password_reg = st.text_input("Confirm Password", type="password")
                
                if st.form_submit_button("Register Business"):
                    if password_reg != confirm_password_reg:
                        st.error("Passwords don't match")
                    elif not name_reg or not email_reg or not password_reg:
                        st.error("Business Name, Email, and Password are required.")
                    else:
                        conn = get_db_connection()
                        cur = conn.cursor()
                        try:
                            password_hash_reg = hash_password(password_reg)
                            cur.execute(
                                "INSERT INTO businesses (name, email, password_hash) VALUES (%s, %s, %s) RETURNING id",
                                (name_reg, email_reg, password_hash_reg)
                            )
                            new_business_id = cur.fetchone()[0]
                            conn.commit()
                            st.success("Business registration successful! Please login.")
                            st.session_state.user_type_for_login = 'business'
                            st.session_state.registered_email = email_reg
                            st.rerun() # Rerun to prefill login form
                        except psycopg2.IntegrityError:
                            st.error("This email is already registered for a business.")
                        except Exception as e:
                            st.error(f"Registration failed: {e}")
                        finally:
                            cur.close()
                            conn.close()
    
    elif login_user_type == 'Investor':
        st.subheader("Investor Login")
        with st.form("investor_login_form"):
            email = st.text_input("Email", value=default_email if st.session_state.get('user_type_for_login') == 'investor' else "")
            password = st.text_input("Password", type="password")

            if st.form_submit_button("Login as Investor"):
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("SELECT id, password_hash, name FROM investors WHERE email = %s", (email,))
                result = cur.fetchone()
                cur.close()
                conn.close()

                if result and result[1] and verify_password(password, result[1]): # Check if password_hash exists
                    st.session_state.user_type = 'investor'
                    st.session_state.logged_in_entity_id = result[0] 
                    st.session_state.user_name = result[2]
                    if 'registered_email' in st.session_state: del st.session_state.registered_email
                    if 'user_type_for_login' in st.session_state: del st.session_state.user_type_for_login
                    st.rerun()
                else:
                    st.error("Invalid email or password for Investor, or investor not fully registered.")

    elif login_user_type == 'Service Provider':
        st.subheader("Service Provider Login")
        with st.form("service_provider_login_form"):
            email = st.text_input("Email", value=default_email if st.session_state.get('user_type_for_login') == 'service_provider' else "")
            password = st.text_input("Password", type="password")

            if st.form_submit_button("Login as Service Provider"):
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("SELECT id, password_hash, name, service_type FROM service_providers WHERE email = %s", (email,))
                result = cur.fetchone()
                cur.close()
                conn.close()

                if result and result[1] and verify_password(password, result[1]): # Check if password_hash exists
                    st.session_state.user_type = 'service_provider'
                    st.session_state.logged_in_entity_id = result[0] 
                    st.session_state.user_name = result[2]
                    st.session_state.service_type = result[3] 
                    if 'registered_email' in st.session_state: del st.session_state.registered_email
                    if 'user_type_for_login' in st.session_state: del st.session_state.user_type_for_login
                    st.rerun()
                else:
                    st.error("Invalid email or password for Service Provider, or provider not fully registered.")
    
    st.markdown("---")
    st.write("Don't have an account of the selected type above? Register here:")
    if login_user_type == 'Investor':
        if st.button("Register as Investor"):
            st.session_state.show_registration_form = 'investor'
            st.rerun()
    elif login_user_type == 'Service Provider':
        if st.button("Register as Service Provider"):
            st.session_state.show_registration_form = 'service_provider'
            st.rerun()
    elif login_user_type == 'Business':
        st.caption("Business registration is available in the 'Register New Business' tab above.")


def check_auth():
    """Checks if a user is authenticated based on session state and token."""
    if 'user_type' in st.session_state:
        if st.session_state.user_type == 'business':
            if 'token' not in st.session_state: return False
            business_id_from_token = verify_jwt(st.session_state.token)
            if not business_id_from_token: return False
            
            conn = get_db_connection()
            cur = conn.cursor()
            try:
                # Validate session token from DB
                cur.execute("SELECT 1 FROM user_sessions WHERE session_token = %s AND expires_at > NOW() AND business_id = %s",
                            (st.session_state.token, business_id_from_token))
                valid_session = cur.fetchone() is not None
                if not valid_session:
                    logout() 
                    return False
                # Ensure session state consistency
                st.session_state.business_id = business_id_from_token 
                st.session_state.logged_in_entity_id = business_id_from_token
                return True
            except Exception as e:
                st.error(f"Session validation error: {e}") 
                return False
            finally:
                cur.close()
                conn.close()

        elif st.session_state.user_type in ['investor', 'service_provider']:
            # For these types, we are relying on session state directly for now
            return ('logged_in_entity_id' in st.session_state and 
                    'user_name' in st.session_state and
                    st.session_state.logged_in_entity_id is not None) # Ensure ID is not None
    return False

def logout():
    """Logs out the current user by clearing session state and invalidating DB session."""
    if 'token' in st.session_state and st.session_state.get('user_type') == 'business':
        conn = get_db_connection()
        cur = conn.cursor()
        try:
            cur.execute(
                "DELETE FROM user_sessions WHERE session_token = %s",
                (st.session_state.token,)
            )
            conn.commit()
        except Exception as e:
            st.warning(f"Error during DB session cleanup: {e}") # Log but don't block logout
        finally:
            cur.close()
            conn.close()
    
    # Clear all session state keys to ensure a clean logout
    keys_to_delete = list(st.session_state.keys())
    for key in keys_to_delete:
        del st.session_state[key]

# Helper Functions for Dashboard and Reports
def time_ago(dt_object):
    """Converts a datetime object or date object to a 'time ago' string."""
    if dt_object is None: return "some time ago"
    now = datetime.now(timezone.utc)
    if isinstance(dt_object, date) and not isinstance(dt_object, datetime):
        dt_object = datetime.combine(dt_object, datetime.min.time(), tzinfo=timezone.utc)
    elif isinstance(dt_object, datetime) and dt_object.tzinfo is None:
        dt_object = dt_object.replace(tzinfo=timezone.utc)
    if not isinstance(dt_object, datetime) or dt_object.tzinfo is None: return "invalid date"
    diff = now - dt_object
    seconds = diff.total_seconds()
    if seconds < 0: return "in the future"
    minutes, hours, days = seconds / 60, seconds / 3600, seconds / 86400
    if seconds < 60: return f"{int(seconds)} seconds ago"
    if minutes < 60: return f"{int(minutes)} minutes ago"
    if hours < 24: return f"{int(hours)} hours ago"
    if days < 7: return f"{int(days)} days ago"
    if days < 30: return f"{int(days // 7)} weeks ago"
    if days < 365: return f"{int(days // 30)} months ago"
    return f"{int(days // 365)} years ago"

def get_quarter_dates(date_obj):
    """Returns (start_date, end_date) for the quarter of date_obj."""
    if isinstance(date_obj, datetime): date_obj = date_obj.date()
    ts = pd.Timestamp(date_obj)
    return ts.to_period('Q').start_time.date(), ts.to_period('Q').end_time.date()

def get_previous_quarter_dates(date_obj):
    """Returns (start_date, end_date) for the quarter before date_obj's quarter."""
    if isinstance(date_obj, datetime): date_obj = date_obj.date()
    ts = pd.Timestamp(date_obj)
    current_quarter_start = ts.to_period('Q').start_time.date()
    return get_quarter_dates(current_quarter_start - timedelta(days=1))

def get_dashboard_financials(business_id, period_start, period_end):
    """Fetches total revenue from invoices within a period for the dashboard."""
    conn = get_db_connection()
    cur = conn.cursor()
    revenue = 0.0
    try:
        cur.execute(
            "SELECT SUM(total_amount) FROM invoices WHERE business_id = %s AND issue_date BETWEEN %s AND %s",
            (business_id, period_start, period_end)
        )
        result = cur.fetchone()
        if result and result[0] is not None: revenue = float(result[0])
    except Exception as e: st.error(f"Error fetching revenue: {e}")
    finally: cur.close(); conn.close()
    return revenue

def get_total_monthly_salary_expense(business_id):
    """Fetches current total monthly salary expense from employee salaries."""
    conn = get_db_connection()
    cur = conn.cursor()
    total_salary = 0.0
    try:
        cur.execute("SELECT SUM(salary) FROM employees WHERE business_id = %s", (business_id,))
        result = cur.fetchone()
        if result and result[0] is not None: total_salary = float(result[0])
    except Exception as e: st.error(f"Error fetching total salaries: {e}")
    finally: cur.close(); conn.close()
    return total_salary

def get_recent_activities_for_dashboard(business_id, limit=4):
    """Fetches recent activities (sales, projects, HR, inventory) for the dashboard."""
    activities_data = []
    conn = get_db_connection()
    cur = conn.cursor()
    queries = [
        ("New Sale", """SELECT customer_name, total_amount, issue_date FROM invoices WHERE business_id = %s ORDER BY issue_date DESC, id DESC LIMIT 1""", lambda r: f"To {r[0]} for ${r[1]:,.2f}"),
        ("Project Update", """SELECT title, content, created_at FROM documents WHERE business_id = %s AND doc_type = 'project_update' ORDER BY created_at DESC LIMIT 1""", lambda r: f"{r[0]}: {r[1][:70]}..." if r[1] else r[0]),
        ("HR - New Hire", """SELECT name, join_date FROM employees WHERE business_id = %s ORDER BY join_date DESC, id DESC LIMIT 1""", lambda r: f"Welcome aboard, {r[0]}!"),
        ("Inventory Alert", """SELECT name, quantity, created_at FROM products WHERE business_id = %s AND quantity < 10 ORDER BY quantity ASC, created_at DESC LIMIT 1""", lambda r: f"Low stock for {r[0]} (Qty: {r[1]})")
    ]
    for act_type, query, detail_formatter in queries:
        try:
            cur.execute(query, (business_id,))
            record = cur.fetchone()
            if record:
                activities_data.append({"type": act_type, "detail": detail_formatter(record), "time_obj": record[-1]})
        except Exception as e: st.warning(f"Error fetching {act_type.lower()} activity: {e}")
    cur.close(); conn.close()
    def get_sort_key(activity):
        time_val = activity.get("time_obj")
        if isinstance(time_val, datetime): 
            if time_val.tzinfo is None: # It's a naive datetime
                return time_val.replace(tzinfo=timezone.utc) # Make it UTC aware
            return time_val
        if isinstance(time_val, date): return datetime.combine(time_val, datetime.min.time(), tzinfo=timezone.utc)
        return datetime.min.replace(tzinfo=timezone.utc)
    activities_data.sort(key=get_sort_key, reverse=True)
    for act in activities_data: act["time_string"] = time_ago(act.get("time_obj"))
    return activities_data[:limit]

# Inventory & Billing Module
def inventory_module(business_id, ai_models):
    """Streamlit module for Inventory and Billing management."""
    st.header("📦 Inventory & Billing Management")
    if "invoice_download_details" not in st.session_state:
        st.session_state.invoice_download_details = None
    conn = get_db_connection()
    cur = conn.cursor()
    tab1, tab2, tab3, tab4 = st.tabs(["View Inventory", "Add Product", "Generate Bill", "Reports"])
    with tab1:
        cur.execute("SELECT id, name, description, price, quantity, category, created_at FROM products WHERE business_id = %s ORDER BY name", (business_id,))
        products = cur.fetchall()
        if products:
            df = pd.DataFrame(products, columns=["ID", "Name", "Description", "Price", "Quantity", "Category", "Created At"])
            st.dataframe(df, hide_index=True)
            df["Quantity"] = pd.to_numeric(df["Quantity"], errors="coerce").fillna(0)
            low_stock = df[df["Quantity"] < 10]
            if not low_stock.empty:
                st.warning("Low Stock Alert!"); st.dataframe(low_stock[["Name", "Quantity"]])
        else: st.info("No products in inventory yet.")
    with tab2:
        with st.form("add_product"):
            name = st.text_input("Product Name", key="product_name")
            description = st.text_area("Description", key="product_desc")
            price = st.number_input("Price", min_value=0.0, step=0.01, key="product_price")
            quantity = st.number_input("Quantity", min_value=0, step=1, key="product_qty")
            category = st.text_input("Category", key="product_cat")
            if st.form_submit_button("Add Product"):
                cur.execute(
                    "INSERT INTO products (business_id, name, description, price, quantity, category) VALUES (%s, %s, %s, %s, %s, %s)",
                    (business_id, name, description, price, quantity, category))
                conn.commit(); st.success("Product added successfully!"); st.rerun()
    with tab3:
        with st.form("create_invoice"):
            customer_name = st.text_input("Customer Name")
            customer_email = st.text_input("Customer Email")
            due_date = st.date_input("Due Date", datetime.now().date() + timedelta(days=14))
            cur.execute("SELECT id, name, price FROM products WHERE business_id = %s", (business_id,))
            products_for_invoice = cur.fetchall()
            items = []
            if products_for_invoice:
                st.write("### Invoice Items")
                # Header row
                cols_header = st.columns([3, 2, 2, 1])
                cols_header[0].write("**Product**"); cols_header[1].write("**Price**"); cols_header[2].write("**Quantity**"); cols_header[3].write("**Total**")

                for i, product_item in enumerate(products_for_invoice):
                    cols_item = st.columns([3, 2, 2, 1])
                    cols_item[0].write(product_item[1]) # Name
                    cols_item[1].write(f"${product_item[2]:.2f}") # Price
                    qty = cols_item[2].number_input(f"Qty", min_value=0, value=0, key=f"qty_{product_item[0]}", label_visibility="collapsed")
                    item_total = product_item[2] * qty
                    cols_item[3].write(f"${item_total:.2f}")
                    if qty > 0: items.append({"product_id": product_item[0], "name": product_item[1], "price": float(product_item[2]), "quantity": qty, "total": float(item_total)})
            
            if st.form_submit_button("Create Invoice"):
                if not items: st.error("Please add at least one item.")
                else:
                    subtotal = sum(item['total'] for item in items)
                    tax_amount = subtotal * 0.18; total_amount = subtotal + tax_amount
                    invoice_number = f"INV-{datetime.now().strftime('%Y%m%d')}-{np.random.randint(1000,9999)}"
                    cur.execute(
                        """INSERT INTO invoices (business_id, invoice_number, customer_name, customer_email, issue_date, due_date, total_amount, tax_amount, items) 
                           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                        (business_id, invoice_number, customer_name, customer_email, datetime.now().date(), due_date, total_amount, tax_amount, json.dumps(items)))
                    for item in items: cur.execute("UPDATE products SET quantity = quantity - %s WHERE id = %s AND business_id = %s", (item['quantity'], item['product_id'], business_id))
                    conn.commit()
                    st.success("Invoice created successfully!")
                    invoice_content_bytes = generate_invoice_pdf(business_id, invoice_number, customer_name, customer_email, datetime.now().date(), due_date, items, subtotal, tax_amount, total_amount)
                    st.session_state.invoice_download_details = {"data": invoice_content_bytes, "file_name": f"invoice_{invoice_number}.pdf", "invoice_number": invoice_number}
                    st.rerun() # Rerun to show download button outside form

        if st.session_state.invoice_download_details:
            details = st.session_state.invoice_download_details
            st.subheader(f"Invoice #{details['invoice_number']} Ready")
            st.download_button(label="Download Invoice (PDF)", data=details["data"], file_name=details["file_name"], mime="application/pdf", key="download_invoice_final_btn")
            if st.button("Create Another Invoice", key="create_another_inv_btn"):
                st.session_state.invoice_download_details = None; st.rerun()
    with tab4:
        st.subheader("Inventory Reports")
        cur.execute("SELECT DATE_TRUNC('month', issue_date) AS month, SUM(total_amount) AS sales FROM invoices WHERE business_id = %s GROUP BY month ORDER BY month", (business_id,))
        sales_data = cur.fetchall()
        if sales_data: st.plotly_chart(px.line(pd.DataFrame(sales_data, columns=["Month", "Sales"]), x="Month", y="Sales", title="Monthly Sales"))
        else: st.info("No sales data.")
        cur.execute("SELECT category, SUM(price * quantity) AS value FROM products WHERE business_id = %s GROUP BY category", (business_id,))
        inv_data = cur.fetchall()
        if inv_data: st.plotly_chart(px.pie(pd.DataFrame(inv_data, columns=["Category", "Value"]), values="Value", names="Category", title="Inventory Value by Category"))
        else: st.info("No inventory data.")
    cur.close(); conn.close()

def generate_invoice_pdf(business_id, invoice_number, customer_name, customer_email, issue_date, due_date, items, subtotal, tax_amount, total_amount):
    """Generates a simple PDF invoice."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    
    # Get business name (optional, can be hardcoded or fetched)
    business_name_display = "Your Business Name" # Placeholder
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT name FROM businesses WHERE id = %s", (business_id,))
        res = cur.fetchone()
        if res: business_name_display = res[0]
        cur.close()
        conn.close()
    except: pass # Ignore if DB call fails for this non-critical info

    pdf.cell(0, 10, f"INVOICE #{invoice_number}", 0, 1, "C")
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 7, f"Issue Date: {issue_date}", 0, 1)
    pdf.cell(0, 7, f"Due Date: {due_date}", 0, 1)
    pdf.ln(10)

    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 7, "From:", 0, 1)
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 7, business_name_display, 0, 1)
    pdf.cell(0, 7, "[Your Business Address]", 0, 1) # Placeholder
    pdf.ln(5)

    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 7, "To:", 0, 1)
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 7, customer_name, 0, 1)
    pdf.cell(0, 7, customer_email, 0, 1)
    pdf.ln(10)

    pdf.set_font("Arial", "B", 10)
    col_widths = [90, 30, 30, 40]
    pdf.cell(col_widths[0], 8, "Item", 1, 0, "C")
    pdf.cell(col_widths[1], 8, "Price", 1, 0, "C")
    pdf.cell(col_widths[2], 8, "Qty", 1, 0, "C")
    pdf.cell(col_widths[3], 8, "Total", 1, 1, "C")
    
    pdf.set_font("Arial", "", 10)
    for item in items:
        pdf.cell(col_widths[0], 7, item['name'], 1)
        pdf.cell(col_widths[1], 7, f"${item['price']:.2f}", 1, 0, "R")
        pdf.cell(col_widths[2], 7, str(item['quantity']), 1, 0, "C")
        pdf.cell(col_widths[3], 7, f"${item['total']:.2f}", 1, 1, "R")
    pdf.ln(5)

    summary_col_width = sum(col_widths[:3])
    pdf.cell(summary_col_width, 7, "SUBTOTAL:", 0, 0, "R")
    pdf.cell(col_widths[3], 7, f"${subtotal:.2f}", 1, 1, "R")
    pdf.cell(summary_col_width, 7, "TAX (18%):", 0, 0, "R")
    pdf.cell(col_widths[3], 7, f"${tax_amount:.2f}", 1, 1, "R")
    pdf.set_font("Arial", "B", 10)
    pdf.cell(summary_col_width, 7, "TOTAL:", 0, 0, "R")
    pdf.cell(col_widths[3], 7, f"${total_amount:.2f}", 1, 1, "R")
    pdf.ln(10)

    pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 7, "Payment Terms: Due upon receipt.\nPayment Methods: [List your payment methods]")
    
    pdf_output_data = pdf.output(dest="S")

    if isinstance(pdf_output_data, str):
        # This is the expected path if FPDF.output(dest="S") returns a string
        return pdf_output_data.encode("latin-1")
    elif isinstance(pdf_output_data, bytearray):
        # This handles the case where it might return a bytearray
        return bytes(pdf_output_data)
    elif isinstance(pdf_output_data, bytes):
        # If it already returns bytes
        return pdf_output_data
    else:
        # Fallback for unexpected types, though unlikely for FPDF
        st.error(f"FPDF output returned an unexpected type: {type(pdf_output_data)}. Attempting to convert.")
        return str(pdf_output_data).encode("latin-1")



# HR Tools Module
def hr_module(business_id, ai_models):
    """Streamlit module for Human Resources management."""
    st.header("👥 HR Tools")
    conn = get_db_connection(); cur = conn.cursor()
    tabs_hr = ["Employee Directory", "Appraisals", "Attendance", "Work Reports", "Analytics"]
    tab1, tab2, tab3, tab4, tab5 = st.tabs(tabs_hr)

    with tab1: # Employee Directory
        cur.execute("SELECT id, name, email, position, department, salary, join_date, last_appraisal_date, performance_score, skills FROM employees WHERE business_id = %s ORDER BY name", (business_id,))
        employees = cur.fetchall()
        if employees:
            df_emp = pd.DataFrame(employees, columns=["ID", "Name", "Email", "Position", "Department", "Salary", "Join Date", "Last Appraisal", "Performance", "Skills"])
            st.dataframe(df_emp, hide_index=True)
            search_term = st.text_input("Search Employees")
            if search_term:
                df_emp_filtered = df_emp[df_emp.apply(lambda row: search_term.lower() in str(row).lower(), axis=1)]
                st.dataframe(df_emp_filtered, hide_index=True)
        else: st.info("No employees yet.")
        with st.expander("Add New Employee"):
            with st.form("add_employee"):
                name, email, pos, dept = st.text_input("Full Name"), st.text_input("Email"), st.text_input("Position"), st.text_input("Department")
                salary = st.number_input("Salary", min_value=0, step=1000)
                join_date = st.date_input("Join Date", value=date.today())
                skills_str = st.text_input("Skills (comma separated)")
                if st.form_submit_button("Add Employee"):
                    skills_list = [s.strip() for s in skills_str.split(",") if s.strip()]
                    try:
                        cur.execute("""INSERT INTO employees (business_id, name, email, position, department, salary, join_date, skills) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)""",
                                    (business_id, name, email, pos, dept, salary, join_date, skills_list))
                        conn.commit(); st.success("Employee added!"); st.rerun()
                    except psycopg2.IntegrityError: st.error("Email already exists.")
                    except Exception as e: st.error(f"Error: {e}")
    
    with tab2: # Appraisals
        st.subheader("Employee Appraisals")
        cur.execute("SELECT id, name FROM employees WHERE business_id = %s", (business_id,))
        employees_app = cur.fetchall()
        if employees_app:
            emp_options = {f"{e[1]} (ID: {e[0]})": e[0] for e in employees_app}
            selected_emp_key = st.selectbox("Select Employee", list(emp_options.keys()))
            if selected_emp_key:
                emp_id = emp_options[selected_emp_key]
                cur.execute("SELECT name, position, department, salary, last_appraisal_date, performance_score FROM employees WHERE id = %s AND business_id = %s", (emp_id, business_id))
                emp_data = cur.fetchone()
                # Display emp_data and form... (same as before)
                st.write(f"### Appraisal for {emp_data[0]} ({emp_data[1]})")
                st.write(f"**Position:** {emp_data[1]}")
                st.write(f"**Department:** {emp_data[2]}")
                st.write(f"**Current Salary:** ${emp_data[3]:,.2f}")
                st.write(f"**Last Appraisal:** {emp_data[4] if emp_data[4] else 'Never'}")
                st.write(f"**Performance Score:** {emp_data[5] if emp_data[5] is not None else 'Not rated'}/10")

                with st.form("appraisal_form"):
                    current_score = emp_data[5] if emp_data[5] is not None else 5
                    new_score = st.slider("New Performance Score (1-10)", 1, 10, int(current_score))
                    salary_adj_percent = st.number_input("Salary Adjustment (%)", -20.0, 50.0, 0.0, 0.5)
                    comments = st.text_area("Appraisal Comments")
                    appraisal_date = st.date_input("Appraisal Date", value=date.today())
                    if st.form_submit_button("Submit Appraisal"):
                        new_salary = float(emp_data[3]) * (1 + salary_adj_percent / 100)
                        cur.execute("""UPDATE employees SET performance_score=%s, salary=%s, last_appraisal_date=%s WHERE id=%s AND business_id=%s""",
                                    (new_score, new_salary, appraisal_date, emp_id, business_id))
                        # Add to documents (same as before)
                        doc_title = f"Appraisal for {emp_data[0]} - {appraisal_date}"
                        doc_content = f"""
                        Employee: {emp_data[0]}
                        Position: {emp_data[1]}
                        Department: {emp_data[2]}
                        
                        Previous Performance Score: {emp_data[5] if emp_data[5] is not None else 'N/A'} → {new_score}
                        
                        Previous Salary: ${emp_data[3]:,.2f}
                        New Salary: ${new_salary:,.2f}
                        Adjustment: {salary_adj_percent}%
                        
                        Comments:
                        {comments}
                        """
                        cur.execute(
                            """INSERT INTO documents 
                            (business_id, title, content, doc_type, created_by) 
                            VALUES (%s, %s, %s, %s, %s)""",
                            (business_id, doc_title, doc_content, "appraisal", emp_id)
                        )
                        conn.commit(); st.success("Appraisal submitted!"); st.rerun()
        else: st.info("No employees to appraise.")

    with tab3: # Attendance
        st.subheader("Attendance Tracking")
        view_option = st.radio("View Mode", ["Daily View", "Employee Summary"], key="att_view_mode")
        if view_option == "Daily View":
            att_date = st.date_input("Select Date", datetime.now().date(), key="att_daily_date")
            cur.execute("SELECT id, name FROM employees WHERE business_id = %s", (business_id,))
            employees_att = cur.fetchall()
            if employees_att:
                st.write("Mark Attendance:")
                for emp in employees_att:
                    cur.execute("SELECT status FROM attendance WHERE employee_id = %s AND date = %s AND business_id = %s", (emp[0], att_date, business_id))
                    existing_att = cur.fetchone()
                    att_status_options = ["Present", "Absent", "Late", "Leave"]
                    current_idx = att_status_options.index(existing_att[0]) if existing_att else 0
                    status = st.radio(f"{emp[1]}", att_status_options, index=current_idx, key=f"att_{emp[0]}_{att_date}", horizontal=True)
                    
                    # Note: This updates the DB on every radio button click. 
                    # A "Save All" button might be better for performance with many employees.
                    if existing_att:
                        if status != existing_att[0]: 
                            cur.execute("UPDATE attendance SET status = %s WHERE employee_id = %s AND date = %s AND business_id = %s", (status, emp[0], att_date, business_id))
                            conn.commit() # Commit immediately after update
                    else: 
                        cur.execute("INSERT INTO attendance (business_id, employee_id, date, status) VALUES (%s, %s, %s, %s)", (business_id, emp[0], att_date, status))
                        conn.commit() # Commit immediately after insert
                # Summary display...
                st.write("### Attendance Summary for Selected Date")
                cur.execute(
                    """SELECT status, COUNT(*) 
                    FROM attendance 
                    WHERE date = %s AND business_id = %s 
                    GROUP BY status""",
                    (att_date, business_id)
                )
                summary = cur.fetchall()
                if summary:
                    df_summary = pd.DataFrame(summary, columns=["Status", "Count"])
                    st.dataframe(df_summary, hide_index=True)
                    fig = px.pie(df_summary, values="Count", names="Status", title=f"Attendance for {att_date}")
                    st.plotly_chart(fig)
                else:
                    st.info("No attendance recorded for this date yet.")

            else: st.info("No employees for attendance.")
        else: # Employee Summary
            cur.execute("SELECT id, name FROM employees WHERE business_id = %s", (business_id,))
            employees = cur.fetchall()
            if employees:
                employee_options = {e[1]: e[0] for e in employees}
                selected_employee_name = st.selectbox("Select Employee", list(employee_options.keys()), key="att_emp_summary_select")
                if selected_employee_name:
                    selected_employee_id = employee_options[selected_employee_name]
                    start_date = st.date_input("Start Date", datetime.now().date() - timedelta(days=30), key="att_summary_start_date")
                    end_date = st.date_input("End Date", datetime.now().date(), key="att_summary_end_date")
                    
                    cur.execute(
                        """SELECT date, status 
                        FROM attendance 
                        WHERE employee_id = %s AND business_id = %s 
                        AND date BETWEEN %s AND %s 
                        ORDER BY date""",
                        (selected_employee_id, business_id, start_date, end_date)
                    )
                    records = cur.fetchall()
                    
                    if records:
                        df_records = pd.DataFrame(records, columns=["Date", "Status"])
                        st.write(f"### Attendance Records for {selected_employee_name}")
                        st.dataframe(df_records, hide_index=True)
                        
                        total_days_in_period = (end_date - start_date).days + 1
                        present_days = len([r for r in records if r[1] == "Present"])
                        attendance_percent = (present_days / total_days_in_period) * 100 if total_days_in_period > 0 else 0
                        
                        st.metric("Attendance Percentage (in period)", f"{attendance_percent:.1f}%")
                        
                        # Plot attendance trend
                        df_records['Present'] = df_records['Status'].apply(lambda x: 1 if x == "Present" else 0)
                        df_records.set_index('Date', inplace=True)
                        # Resample weekly, summing 'Present' days
                        df_weekly_present = df_records['Present'].resample('W').sum().reset_index()
                        df_weekly_present.columns = ['Week Start', 'Days Present']

                        if not df_weekly_present.empty:
                             fig = px.bar(df_weekly_present, x="Week Start", y="Days Present", title="Weekly Days Present")
                             st.plotly_chart(fig)
                        else:
                             st.info("Not enough attendance data in the period for a weekly trend.")

                    else:
                        st.info("No attendance records found for selected employee and period.")
            else:
                st.info("No employees to show attendance summary for.")


    with tab4: # Work Reports (Simulated)
        st.subheader("Work Reports (Simulated)")
        cur.execute("SELECT id, name FROM employees WHERE business_id = %s", (business_id,))
        employees = cur.fetchall()
        if employees:
            employee_options = {f"{e[1]} (ID: {e[0]})": e[0] for e in employees}
            selected_employee_key = st.selectbox("Select Employee", list(employee_options.keys()), key="work_report_emp_select")
            if selected_employee_key:
                employee_id = employee_options[selected_employee_key]
                report_period = st.selectbox("Report Period", ["Daily", "Weekly", "Monthly"], key="work_report_period")
                report_date = st.date_input("Report Date", datetime.now().date(), key="work_report_date")
                
                if st.button("Generate Report", key="generate_work_report"):
                    # In a real app, this would pull actual data based on tasks, project contributions, etc.
                    # For now, it's simulated.
                    report_data = {
                        "tasks_completed": np.random.randint(3, 10),
                        "hours_worked": np.random.randint(4, 9),
                        "meetings_attended": np.random.randint(1, 5),
                        "issues_resolved": np.random.randint(1, 4),
                        "feedback": "Good performance this period. Keep it up!"
                    }
                    
                    st.write(f"### {report_period} Work Report for {selected_employee_key.split(' (')[0]}")
                    st.write(f"- Tasks Completed: {report_data['tasks_completed']}")
                    st.write(f"- Hours Worked: {report_data['hours_worked']}")
                    st.write(f"- Meetings Attended: {report_data['meetings_attended']}")
                    st.write(f"- Issues Resolved: {report_data['issues_resolved']}")
                    st.write(f"- Feedback: {report_data['feedback']}")
                    
                    # Simulate saving to documents
                    doc_title = f"Work Report - {selected_employee_key.split(' (')[0]} - {report_date} ({report_period})"
                    doc_content = f"""
                    Employee: {selected_employee_key.split(' (')[0]}
                    Period: {report_period} ending {report_date}
                    
                    Tasks Completed: {report_data['tasks_completed']}
                    Hours Worked: {report_data['hours_worked']}
                    Meetings Attended: {report_data['meetings_attended']}
                    Issues Resolved: {report_data['issues_resolved']}
                    
                    Feedback:
                    {report_data['feedback']}
                    """
                    try:
                        cur.execute(
                            """INSERT INTO documents 
                            (business_id, title, content, doc_type, created_by) 
                            VALUES (%s, %s, %s, %s, %s)""",
                            (business_id, doc_title, doc_content, "work_report", employee_id)
                        )
                        conn.commit()
                        st.success("Report saved to documents!")
                    except Exception as e:
                        st.error(f"Error saving report to documents: {e}")

        else: st.info("No employees to generate reports for")

    with tab5: # Analytics
        st.subheader("HR Analytics Dashboard")
        
        # Employee distribution by department
        cur.execute("SELECT department, COUNT(*) FROM employees WHERE business_id = %s GROUP BY department", (business_id,))
        dept_data = cur.fetchall()
        if dept_data:
            df_dept = pd.DataFrame(dept_data, columns=["Department", "Count"])
            fig = px.pie(df_dept, values="Count", names="Department", title="Employees by Department")
            st.plotly_chart(fig)
        else: st.info("No department data available")
        
        # Salary distribution
        cur.execute("SELECT position, salary FROM employees WHERE business_id = %s", (business_id,))
        salary_data = cur.fetchall()
        if salary_data:
            df_salary = pd.DataFrame(salary_data, columns=["Position", "Salary"])
            fig = px.box(df_salary, y="Salary", title="Salary Distribution by Position")
            st.plotly_chart(fig)
        else: st.info("No salary data available")
        
        # Performance vs Salary
        cur.execute("SELECT performance_score, salary FROM employees WHERE business_id = %s AND performance_score IS NOT NULL", (business_id,))
        perf_data = cur.fetchall()
        if perf_data:
            df_perf = pd.DataFrame(perf_data, columns=["Performance", "Salary"])
            fig = px.scatter(df_perf, x="Performance", y="Salary", trendline="ols", 
                           title="Performance vs Salary (Rated Employees)")
            st.plotly_chart(fig)
        else: st.info("No performance data available for analysis.")
    
    cur.close(); conn.close()

# Project Manager Module
def project_module(business_id, ai_models):
    """Streamlit module for Project Management."""
    st.header("📊 Project Manager")
    conn = get_db_connection(); cur = conn.cursor()
    tab1, tab2, tab3, tab4 = st.tabs(["All Projects", "New Project", "Project Details", "Gantt Chart"])
    
    with tab1:
        st.subheader("All Projects")
        cur.execute(
            """SELECT p.id, p.name, p.client, p.start_date, p.end_date, p.status, p.progress, e.name as manager 
            FROM projects p LEFT JOIN employees e ON p.manager_id = e.id 
            WHERE p.business_id = %s ORDER BY p.start_date DESC""",
            (business_id,)
        )
        projects = cur.fetchall()
        
        if projects:
            df = pd.DataFrame(projects, columns=[
                "ID", "Name", "Client", "Start Date", "End Date", 
                "Status", "Progress", "Manager"
            ])
            st.dataframe(df, hide_index=True)
            
            # Filter options
            all_statuses = df["Status"].unique().tolist() if not df.empty else []
            status_filter = st.multiselect(
                "Filter by Status",
                options=all_statuses,
                default=all_statuses,
                key="project_status_filter"
            )
            
            if status_filter:
                filtered_df = df[df["Status"].isin(status_filter)]
                st.dataframe(filtered_df, hide_index=True)
            else:
                 st.info("Select statuses to filter projects.")
        else:
            st.info("No projects found.")
    
    with tab2:
        st.subheader("Create New Project")
        cur.execute("SELECT id, name FROM employees WHERE business_id = %s ORDER BY name", (business_id,))
        managers = cur.fetchall()
        manager_options = {m[1]: m[0] for m in managers}
        
        with st.form("new_project"):
            name = st.text_input("Project Name")
            client = st.text_input("Client Name")
            description = st.text_area("Description")
            start_date = st.date_input("Start Date", value=date.today())
            end_date = st.date_input("End Date", value=date.today() + timedelta(days=30))
            budget = st.number_input("Budget", min_value=0.0, step=1000.0)
            
            selected_manager_name = st.selectbox("Project Manager", options=["-- Select --"] + list(manager_options.keys()))
            
            if st.form_submit_button("Create Project"):
                if not name or not client or not description or not start_date or not end_date or selected_manager_name == "-- Select --":
                    st.error("Please fill in all required fields.")
                elif start_date > end_date:
                    st.error("End Date cannot be before Start Date.")
                else:
                    manager_id = manager_options[selected_manager_name]
                    cur.execute(
                        """INSERT INTO projects 
                        (business_id, name, client, description, start_date, end_date, budget, status, manager_id) 
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                        (business_id, name, client, description, start_date, end_date, budget, "Not Started", manager_id)
                    )
                    conn.commit()
                    st.success("Project created successfully!")
                    st.rerun()
    
    with tab3:
        st.subheader("Project Details")
        cur.execute("SELECT id, name FROM projects WHERE business_id = %s ORDER BY name", (business_id,))
        projects = cur.fetchall()
        
        if projects:
            project_options = {f"{p[1]} (ID: {p[0]})": p[0] for p in projects}
            selected_project_key = st.selectbox(
                "Select Project", 
                options=["-- Select --"] + list(project_options.keys()),
                key="select_project_details"
            )
            
            if selected_project_key != "-- Select --":
                project_id = project_options[selected_project_key]
                cur.execute(
                    """SELECT p.*, e.name as manager_name 
                    FROM projects p LEFT JOIN employees e ON p.manager_id = e.id 
                    WHERE p.id = %s AND p.business_id = %s""",
                    (project_id, business_id)
                )
                project = cur.fetchone()
                
                if project:
                    st.write(f"### {project[2]}") # Correct: name
                    st.write(f"**Client:** {project[4]}") # Correct: client
                    st.write(f"**Manager:** {project[11] or 'N/A'}") # Correct: manager_name
                    st.write(f"**Status:** {project[8]}") # Correct: status
                    st.write(f"**Progress:** {project[10]}%") # Correct: progress
                    st.write(f"**Description:** {project[3]}") # Corrected: description
                    
                    # Project timeline
                    today = datetime.now().date()
                    start_date = project[5] # Corrected: start_date
                    end_date = project[6] # Corrected: end_date
                    
                    st.write(f"**Start Date:** {start_date}")
                    st.write(f"**End Date:** {end_date}")

                    if start_date and end_date:
                        total_days = (end_date - start_date).days
                        days_passed = (today - start_date).days
                        progress_percent_timeline = min(100, max(0, (days_passed / total_days) * 100)) if total_days > 0 else 0
                        days_remaining = (end_date - today).days
                        st.write(f"**Days Remaining:** {days_remaining} days" if days_remaining >= 0 else "Project End Date Passed")
                        
                        # Ensure progress is treated as a number for the progress bar
                        current_progress_value = project[10] if project[10] is not None else 0

                        st.write("#### Progress")
                        st.progress(float(current_progress_value) / 100.0, text=f"Project Completion: {current_progress_value}%") # Corrected: use project[10]
                        st.progress(progress_percent_timeline / 100.0, text=f"Timeline Progress: {progress_percent_timeline:.1f}%")
                    elif start_date:
                         st.write("End date not set.")
                    elif end_date:
                         st.write("Start date not set.")
                    else:
                         st.write("Start and end dates not set.")


                    # Update project status
                    with st.expander("Update Project"):
                        cur.execute("SELECT id, name FROM employees WHERE business_id = %s ORDER BY name", (business_id,))
                        managers_update = cur.fetchall()
                        manager_options_update = {m[1]: m[0] for m in managers_update}
                        current_manager_name = project[11] # Corrected: manager_name
                        current_manager_id = project[9] # Corrected: manager_id
                        
                        default_manager_index = 0
                        if current_manager_name and current_manager_name in manager_options_update:
                             default_manager_index = list(manager_options_update.keys()).index(current_manager_name) + 1 # +1 for "-- Select --" if added

                        with st.form("update_project"):
                            new_status = st.selectbox(
                                "Status",
                                ["Not Started", "In Progress", "On Hold", "Completed", "Cancelled"],
                                index=["Not Started", "In Progress", "On Hold", "Completed", "Cancelled"].index(project[8]), # Corrected: status
                                key=f"update_status_{project_id}"
                            )
                            new_progress = st.slider("Progress (%)", 0, 100, current_progress_value, key=f"update_progress_{project_id}") # Corrected: use current_progress_value (from project[10])
                            new_manager_name = st.selectbox("Project Manager", options=["-- Select --"] + list(manager_options_update.keys()), index=default_manager_index, key=f"update_manager_{project_id}")
                            notes = st.text_area("Update Notes", key=f"update_notes_{project_id}")
                            
                            if st.form_submit_button("Update Project"):
                                updated_manager_id = manager_options_update.get(new_manager_name, None) if new_manager_name != "-- Select --" else None
                                cur.execute(
                                    """UPDATE projects 
                                    SET status = %s, progress = %s, manager_id = %s 
                                    WHERE id = %s AND business_id = %s""",
                                    (new_status, new_progress, updated_manager_id, project_id, business_id)
                                )
                                
                                # Add to project documents
                                doc_title = f"Project Update - {project[2]} - {datetime.now().date()}" # Correct: name
                                doc_content = f"""
                                Project: {project[2]} # Correct: name
                                Status Changed: {project[8]} → {new_status} # Corrected: status
                                Progress: {project[10]}% → {new_progress}% # Corrected: progress
                                Manager Changed: {current_manager_name or 'N/A'} → {new_manager_name or 'N/A'}

                                Notes:
                                {notes}
                                """
                                
                                cur.execute(
                                    """INSERT INTO documents 
                                    (business_id, title, content, doc_type) 
                                    VALUES (%s, %s, %s, %s)""",
                                    (business_id, doc_title, doc_content, "project_update")
                                )
                                
                                conn.commit()
                                st.success("Project updated successfully!")
                                st.rerun()
                    
                    # Project team 
                    st.write("### Project Team")
                    cur.execute(
                        """SELECT e.id, e.name, e.position 
                        FROM employees e 
                        JOIN project_assignments pa ON e.id = pa.employee_id 
                        WHERE pa.project_id = %s AND pa.business_id = %s""",
                        (project_id, business_id)
                    )
                    team_members = cur.fetchall()
                    
                    if team_members:
                        df_team = pd.DataFrame(team_members, columns=["ID", "Name", "Position"])
                        st.dataframe(df_team.drop(columns=["ID"]), hide_index=True)
                    else:
                        st.info("No team members assigned yet")
                        
                    # Assign team members
                    with st.expander("Assign/Remove Team Members"):
                        cur.execute("SELECT id, name FROM employees WHERE business_id = %s ORDER BY name", (business_id,))
                        all_employees = cur.fetchall()
                        
                        if all_employees:
                            current_assigned_ids = [m[0] for m in team_members]
                            employee_options_assign = {f"{e[1]} (ID: {e[0]})": e[0] for e in all_employees}
                            
                            selected_employees_keys = st.multiselect(
                                "Select Employees to Assign",
                                options=list(employee_options_assign.keys()),
                                default=[f"{e[1]} (ID: {e[0]})" for e in team_members], # Pre-select current team
                                key=f"assign_employees_{project_id}"
                            )
                            
                            if st.button("Update Team", key=f"update_team_btn_{project_id}"):
                                selected_employee_ids = [employee_options_assign[key] for key in selected_employees_keys]
                                
                                # Employees to remove
                                ids_to_remove = [id for id in current_assigned_ids if id not in selected_employee_ids]
                                if ids_to_remove:
                                    cur.execute(
                                        sql.SQL("DELETE FROM project_assignments WHERE project_id = %s AND business_id = %s AND employee_id IN ({})").format(sql.SQL(',').join(sql.Placeholder() * len(ids_to_remove))),
                                        [project_id, business_id] + ids_to_remove
                                    )

                                # Employees to add
                                ids_to_add = [id for id in selected_employee_ids if id not in current_assigned_ids]
                                if ids_to_add:
                                    insert_values = [(business_id, project_id, emp_id) for emp_id in ids_to_add]
                                    cur.executemany(
                                        """INSERT INTO project_assignments 
                                        (business_id, project_id, employee_id) 
                                        VALUES (%s, %s, %s) ON CONFLICT (business_id, project_id, employee_id) DO NOTHING""",
                                        insert_values
                                    )
                                
                                conn.commit()
                                st.success("Project team updated successfully!")
                                st.rerun()
                        else:
                            st.info("No employees available to assign.")
        else:
            st.info("Select a project from the dropdown to view details.")

    with tab4:
        st.subheader("Project Gantt Chart")
        
        cur.execute(
            "SELECT name, start_date, end_date, status FROM projects WHERE business_id = %s ORDER BY start_date",
            (business_id,)
        )
        projects = cur.fetchall()
        
        if projects:
            gantt_data = []
            for p in projects:
                # Ensure dates are not None for the chart
                if p[1] and p[2]:
                    gantt_data.append({
                        "Task": p[0],
                        "Start": p[1],
                        "Finish": p[2],
                        "Status": p[3]
                    })
            
            if gantt_data:
                df_gantt = pd.DataFrame(gantt_data)
                
                color_map = {
                    "Not Started": "#636EFA",
                    "In Progress": "#00CC96",
                    "On Hold": "#EF553B",
                    "Completed": "#AB63FA",
                    "Cancelled": "#FFA15A"
                }
                
                fig = px.timeline(
                    df_gantt, 
                    x_start="Start", 
                    x_end="Finish", 
                    y="Task",
                    color="Status",
                    color_discrete_map=color_map,
                    title="Project Timeline"
                )
                
                fig.update_yaxes(autorange="reversed") # Show tasks from top to bottom
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("No projects with valid start and end dates to display on the Gantt chart.")
        else:
            st.info("No projects to display.")
    
    cur.close(); conn.close()

# Document Generator Module
def document_module(business_id, ai_models):
    """Streamlit module for Document Generation and Management."""
    st.header("📝 Document Generator")
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    tab1, tab2, tab3, tab4 = st.tabs(["Templates", "Generate Document", "Document Library", "AI Assistant"])
    
    with tab1:
        st.subheader("Document Templates")
        st.info("Template generation is simulated. Click buttons to see examples.")
        
        template_cols = st.columns(3)
        
        with template_cols[0]:
            st.write("**Contract Templates**")
            if st.button("Employment Contract", key="tmpl_emp_contract"): st.write("Simulated Employment Contract Template Content...")
            if st.button("NDA Agreement", key="tmpl_nda"): st.write("Simulated NDA Agreement Template Content...")
            if st.button("Service Contract", key="tmpl_service_contract"): st.write("Simulated Service Contract Template Content...")
        
        with template_cols[1]:
            st.write("**Business Documents**")
            if st.button("Business Proposal", key="tmpl_biz_proposal"): st.write("Simulated Business Proposal Template Content...")
            if st.button("Invoice Template", key="tmpl_invoice"): st.write("Simulated Invoice Template Content...")
            if st.button("Meeting Minutes", key="tmpl_meeting_minutes"): st.write("Simulated Meeting Minutes Template Content...")
        
        with template_cols[2]:
            st.write("**Legal Documents**")
            if st.button("Privacy Policy", key="tmpl_privacy"): st.write("Simulated Privacy Policy Template Content...")
            if st.button("Terms of Service", key="tmpl_terms"): st.write("Simulated Terms of Service Template Content...")
            if st.button("Partnership Agreement", key="tmpl_partnership"): st.write("Simulated Partnership Agreement Template Content...")
    
    with tab2:
        st.subheader("Generate New Document")
        
        doc_type = st.selectbox("Document Type", [
            "Contract", "Letter", "Invoice", "Proposal", "Report", "Other"
        ], key="generate_doc_type")
        
        doc_title = st.text_input("Document Title", key="generate_doc_title")
        
        # Get employees for 'Created By' field
        cur.execute("SELECT id, name FROM employees WHERE business_id = %s ORDER BY name", (business_id,))
        employees_for_doc = cur.fetchall()
        employee_options_doc = {f"{e[1]} (ID: {e[0]})": e[0] for e in employees_for_doc}
        selected_employee_key_doc = st.selectbox("Created By (Employee)", options=["-- Select --"] + list(employee_options_doc.keys()), key="doc_created_by_employee")
        created_by_employee_id = employee_options_doc.get(selected_employee_key_doc) if selected_employee_key_doc != "-- Select --" else None

        # Specific fields for certain document types (simplified)
        doc_content = ""
        if doc_type == "Contract":
            st.write("#### Contract Details")
            parties = st.text_input("Parties Involved (comma separated)", key="contract_parties")
            terms = st.text_area("Key Terms", key="contract_terms")
            duration = st.text_input("Duration", key="contract_duration")
            termination = st.text_area("Termination Clause", key="contract_termination")
            
            if st.button("Generate Contract Content", key="generate_contract_btn"):
                 doc_content = f"""
CONTRACT AGREEMENT

This Agreement is made and entered into on {datetime.now().date()} by and between:

Parties: {parties}

1. TERMS
{terms}

2. DURATION
This agreement shall remain in effect for {duration}.

3. TERMINATION
{termination}

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

___________________________
Signature
"""
                 st.text_area("Generated Document Content", doc_content, height=400, key="generated_contract_content")

        elif doc_type == "Invoice":
             st.write("#### Invoice Details")
             st.info("Invoice generation with line items is handled in the 'Inventory & Billing' module. This is a simple text generator.")
             client_name = st.text_input("Client Name", key="invoice_client_name")
             invoice_details_text = st.text_area("Invoice Summary/Details", key="invoice_details_text")
             amount = st.number_input("Total Amount", min_value=0.0, step=0.01, key="invoice_amount")

             if st.button("Generate Simple Invoice Content", key="generate_simple_invoice_btn"):
                 invoice_number = f"INV-{datetime.now().strftime('%Y%m%d')}-{np.random.randint(1000, 9999)}"
                 doc_content = f"""
INVOICE #{invoice_number}
Date: {datetime.now().date()}

To: {client_name}

Details:
{invoice_details_text}

Total Amount Due: ${amount:,.2f}

Payment Terms: [Your Terms]
"""
                 st.text_area("Generated Document Content", doc_content, height=400, key="generated_simple_invoice_content")

        else:  # Generic document or other types
            prompt = st.text_area("Document Content Prompt (Use AI to generate)", 
                                "Create a professional business document about...", key="ai_doc_prompt")
            
            if st.button("Generate with AI", key="generate_ai_doc_btn"):
                if not prompt.strip():
                    st.warning("Please enter a prompt for the AI.")
                else:
                    with st.spinner("Generating document content with AI..."):
                        doc_content = ai_models.generate_text(
                            f"Create a {doc_type.lower()} document about: {prompt}",
                            max_length=1000
                        )
                        st.text_area("Generated Document Content", doc_content, height=400, key="generated_ai_doc_content")

        # Save and Download section (appears after content is generated/available)
        if doc_content:
            if st.button("Save Document to Library", key="save_generated_doc_btn"):
                if not doc_title.strip():
                    st.error("Please enter a Document Title to save.")
                else:
                    try:
                        cur.execute(
                            """INSERT INTO documents 
                            (business_id, title, content, doc_type, created_by) 
                            VALUES (%s, %s, %s, %s, %s)""",
                            (business_id, doc_title.strip(), doc_content, doc_type.lower(), created_by_employee_id)
                        )
                        conn.commit()
                        st.success("Document saved to library!")
                        # Clear inputs after saving (optional)
                        # st.session_state.generate_doc_title = ""
                        # st.session_state.ai_doc_prompt = "" # Or other specific inputs
                        st.rerun() # Rerun to clear form state if needed
                    except Exception as e:
                        st.error(f"Error saving document: {e}")

            st.write("#### Download Generated Content")
            col_dl_txt, col_dl_docx = st.columns(2)
            with col_dl_txt:
                 st.download_button(
                    "Download as TXT",
                    data=doc_content,
                    file_name=f"{doc_title.replace(' ', '_') or 'document'}.txt",
                    mime="text/plain",
                    key="download_generated_txt"
                )
            with col_dl_docx:
                 # DOCX generation requires python-docx, which might have limitations on complex formatting
                 try:
                    docx_file = BytesIO()
                    doc = docx.Document()
                    # Add content, handling potential formatting issues
                    doc.add_paragraph(doc_content) 
                    doc.save(docx_file)
                    docx_file.seek(0)
                    st.download_button(
                        "Download as DOCX",
                        data=docx_file,
                        file_name=f"{doc_title.replace(' ', '_') or 'document'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_generated_docx"
                    )
                 except Exception as e:
                     st.warning(f"Could not generate DOCX: {e}. Try downloading as TXT.")


    with tab3:
        st.subheader("Document Library")
        
        cur.execute(
            "SELECT d.id, d.title, d.content, d.doc_type, d.created_at, e.name FROM documents d LEFT JOIN employees e ON d.created_by = e.id WHERE d.business_id = %s ORDER BY d.created_at DESC",
            (business_id,)
        )
        documents = cur.fetchall()
        
        if documents:
            search_term = st.text_input("Search Documents", key="doc_library_search")
            
            df = pd.DataFrame(documents, columns=[
                "ID", "Title", "Content", "Type", "Created At", "Created By"
            ])

            df['Created At'] = pd.to_datetime(df['Created At'])  

            if search_term:
                df = df[
                    df["Title"].str.contains(search_term, case=False, na=False) |
                    df["Content"].str.contains(search_term, case=False, na=False) |
                    df["Type"].str.contains(search_term, case=False, na=False) |
                    df["Created By"].str.contains(search_term, case=False, na=False)
                ]
            
            if not df.empty:
                st.dataframe(df[["Title", "Type", "Created At", "Created By"]], hide_index=True)

                st.write("#### View/Download Documents")
                for _, row in df.iterrows():
                    with st.expander(f"{row['Title']} ({row['Type']}) - {row['Created At'].date()}"):
                        st.write(f"**Created By:** {row['Created By'] or 'N/A'}")
                        st.write("---")
                        st.text_area("Content Preview", row["Content"][:1000] + "..." if len(row["Content"]) > 1000 else row["Content"], height=300, key=f"preview_{row['ID']}")
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.download_button(
                                "Download as TXT",
                                data=row["Content"],
                                file_name=f"{row['Title'].replace(' ', '_')}.txt",
                                mime="text/plain",
                                key=f"txt_{row['ID']}"
                            )
                        with col2:
                            try:
                                docx_file = BytesIO()
                                doc = docx.Document()
                                doc.add_paragraph(row["Content"])
                                doc.save(docx_file)
                                docx_file.seek(0)
                                
                                st.download_button(
                                    "Download as DOCX",
                                    data=docx_file,
                                    file_name=f"{row['Title'].replace(' ', '_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"docx_{row['ID']}"
                                )
                            except Exception as e:
                                st.warning(f"Could not generate DOCX for this document: {e}")
            else:
                st.info("No documents found matching your search.")
        else:
            st.info("No documents in the library yet.")
    
    with tab4:
        st.subheader("Document AI Assistant")
        st.info("This AI assistant can help with general questions about documents, but does not currently analyze your specific document library content.")
        
        if "doc_chat_history" not in st.session_state:
            st.session_state.doc_chat_history = []
        
        user_input = st.text_input("Ask about documents or request general advice:", key="doc_ai_user_input")
        
        if user_input:
            with st.spinner("Generating response..."):
                # In a real app, this could potentially analyze uploaded docs or library content
                # For now, it uses the general chatbot model with a document-focused prompt
                prompt_for_ai = f"As a document assistant, respond to the user's query about documents: {user_input}"
                response = ai_models.generate_response(prompt_for_ai)
                
                st.session_state.doc_chat_history.append(("You", user_input))
                st.session_state.doc_chat_history.append(("AI Assistant", response))
        
        # Display chat history
        for speaker, text in st.session_state.doc_chat_history:
            if speaker == "You":
                st.markdown(f"**You**: {text}")
            else:
                st.markdown(f"**AI Assistant**: {text}")
                # Optional: Add TTS for AI responses
                if st.button("🔊 Play AI Response", key=f"tts_doc_ai_{hash(text)}"):
                    try:
                        tts = gTTS(text=text, lang='en'); audio_file = BytesIO()
                        tts.write_to_fp(audio_file); audio_file.seek(0)
                        st.audio(audio_file, format='audio/mp3')
                    except Exception as e:
                        st.error(f"Could not play audio: {e}")
                st.write("---")
    
    cur.close(); conn.close()

# Market Analysis Module
def market_analysis_module(business_id, ai_models):
    """Streamlit module for Market Analysis."""
    st.header("📈 Market Analysis Tool")
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "Industry Analysis", 
        "Trend Insights", 
        "Competitor Benchmark", 
        "Forecasting"
    ])
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    with tab1:
        st.subheader("Industry Analysis")
        industry = st.selectbox("Select Industry", [
            "Technology", "Healthcare", "Finance", "Retail", 
            "Manufacturing", "Education", "Real Estate", "Hospitality", "Other"
        ], key="industry_analysis_select")
        
        if st.button("Analyze Industry", key="analyze_industry_btn"):
            with st.spinner(f"Analyzing {industry} industry..."):
                # Simulate analysis or use AI models
                time.sleep(2)  # Simulate analysis
                
                # Generate analysis using AI
                prompt = f"""
                Provide a brief overview of the {industry} industry. 
                Include typical market size, growth rate, and 3 key trends.
                """
                analysis_text = ai_models.generate_text(prompt, max_length=500)
                
                # Attempt to parse key metrics from AI text or use simulation
                market_size = round(np.random.uniform(1, 500), 2) # Simulated
                growth_rate = round(np.random.uniform(1, 25), 2) # Simulated
                
                st.subheader(f"{industry} Industry Overview")
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Market Size (USD Billion)", f"${market_size}")
                with col2:
                    st.metric("Annual Growth Rate", f"{growth_rate}%")
                
                st.write("### Analysis")
                st.write(analysis_text)

                # Sentiment analysis example on a few simulated phrases
                st.write("### Market Sentiment (Simulated)")
                sample_reviews = [
                    f"Positive outlook for {industry} sector.",
                    f"Challenges in {industry} supply chain.",
                    f"Innovation driving growth in {industry}."
                ]
                try:
                    sentiments = ai_models.analyze_sentiment(sample_reviews)
                    for review, sentiment in zip(sample_reviews, sentiments):
                        label = sentiment['label']
                        score = sentiment['score']
                        st.write(f"- '{review}' - **{label}** ({score:.2f} confidence)")
                except Exception as e:
                    st.warning(f"Could not perform sentiment analysis: {e}")

                # Save analysis data to database (simplified)
                try:
                    cur.execute(
                        """INSERT INTO market_data 
                        (business_id, industry, metric, value, date, source) 
                        VALUES (%s, %s, %s, %s, %s, %s)""",
                        (business_id, industry, "market_size_est", market_size, datetime.now().date(), "AI Analysis")
                    )
                    cur.execute(
                        """INSERT INTO market_data 
                        (business_id, industry, metric, value, date, source) 
                        VALUES (%s, %s, %s, %s, %s, %s)""",
                        (business_id, industry, "growth_rate_est", growth_rate, datetime.now().date(), "AI Analysis")
                    )
                    conn.commit()
                    st.success("Analysis saved to market data.")
                except Exception as e:
                    st.error(f"Error saving market data: {e}")

    with tab2:
        st.subheader("Market Trend Insights")
        st.info("Insights are simulated or generated by AI based on general knowledge, not real-time market feeds.")
        
        trend_topic = st.text_input("Enter a market trend topic (e.g., 'AI in business', 'remote work')", key="trend_topic_input")
        
        if st.button("Get Trend Insights", key="get_trend_insights_btn"):
            if not trend_topic.strip():
                st.warning("Please enter a trend topic.")
            else:
                with st.spinner(f"Getting insights on '{trend_topic}'..."):
                    prompt = f"""
                    Provide insights on the market trend: "{trend_topic}".
                    Discuss its current impact, future potential, and implications for businesses.
                    """
                    insights = ai_models.generate_text(prompt, max_length=800)
                    
                    st.subheader(f"Insights on: {trend_topic}")
                    st.write(insights)

    with tab3:
        st.subheader("Competitor Benchmarking")
        st.info("Benchmarking is simulated. Enter competitor names to see placeholder data.")
        
        competitors_input = st.text_area("Enter your main competitors (one per line)", key="competitors_input")
        
        if st.button("Benchmark Competitors", key="benchmark_competitors_btn"):
            competitor_list = [c.strip() for c in competitors_input.split("\n") if c.strip()]
            
            if not competitor_list:
                st.warning("Please enter competitor names.")
            else:
                st.write("### Simulated Competitor Data")
                competitor_data = []
                for comp in competitor_list:
                    competitor_data.append({
                        "Competitor": comp,
                        "Market Share (%) (Est.)": round(np.random.uniform(1, 40), 1),
                        "Growth (%) (Est.)": round(np.random.uniform(-10, 30), 1),
                        "Key Strength (Sim.)": np.random.choice(["Innovation", "Pricing", "Brand", "Distribution"]),
                        "Key Weakness (Sim.)": np.random.choice(["High Cost", "Limited Reach", "Slow Adaptation", "Customer Service"])
                    })
                
                df_competitors = pd.DataFrame(competitor_data)
                st.dataframe(df_competitors, hide_index=True)

                # Market share chart (simulated)
                if not df_competitors.empty:
                    fig = px.bar(
                        df_competitors, 
                        x="Competitor", 
                        y="Market Share (%) (Est.)", 
                        title="Simulated Market Share Comparison"
                    )
                    st.plotly_chart(fig, use_container_width=True)

    with tab4:
        st.subheader("Market Forecasting")
        st.info("Forecasting uses AI based on general patterns, not specific market data feeds.")
        
        cur.execute(
            "SELECT name FROM products WHERE business_id = %s ORDER BY name",
            (business_id,)
        )
        products = [p[0] for p in cur.fetchall()]
        
        forecast_subject = st.selectbox("Forecast Subject", ["Overall Market", "Specific Product"] + products, key="forecast_subject_select")
        
        if forecast_subject == "Specific Product" and not products:
             st.warning("Add products in the Inventory module to forecast for a specific product.")
             selected_product_for_forecast = None
        elif forecast_subject in products:
             selected_product_for_forecast = forecast_subject
             forecast_subject_text = f"your product '{selected_product_for_forecast}'"
        elif forecast_subject == "Specific Product" and products:
             selected_product_for_forecast = st.selectbox("Select Product", products, key="select_product_for_forecast")
             forecast_subject_text = f"your product '{selected_product_for_forecast}'"
        else:
             selected_product_for_forecast = None
             forecast_subject_text = "the overall market relevant to your business" # General prompt

        forecast_period = st.selectbox("Forecast Period", ["3 months", "6 months", "1 year", "3 years"], key="forecast_period_select")
        
        if st.button("Generate Forecast", key="generate_forecast_btn"):
            with st.spinner("Generating market forecast..."):
                prompt = f"""
                Generate a {forecast_period} market forecast for {forecast_subject_text}.
                Include expected trends, potential challenges, and strategic recommendations.
                """
                forecast_text = ai_models.generate_text(prompt, max_length=1000)
                
                st.subheader(f"Market Forecast for {forecast_subject_text}")
                st.write(forecast_text)
                
                # Simulated forecast chart (simple line)
                try:
                    if forecast_period == "3 months": periods = pd.date_range(start=datetime.now(), periods=3, freq='MS')
                    elif forecast_period == "6 months": periods = pd.date_range(start=datetime.now(), periods=6, freq='MS')
                    elif forecast_period == "1 year": periods = pd.date_range(start=datetime.now(), periods=4, freq='QS') # Quarterly for a year
                    elif forecast_period == "3 years": periods = pd.date_range(start=datetime.now(), periods=3, freq='YS') # Yearly for 3 years
                    else: periods = pd.date_range(start=datetime.now(), periods=4, freq='QS') # Default

                    # Simulate growth based on a base value (e.g., last period's revenue or a base product value)
                    # This is highly simplified. A real forecast would use historical data and models.
                    base_value = 100 # Arbitrary base
                    if selected_product_for_forecast:
                         # Try to get average price or last sale value for the product
                         try:
                             cur.execute("SELECT AVG(price) FROM products WHERE name = %s AND business_id = %s", (selected_product_for_forecast, business_id))
                             avg_price = cur.fetchone()[0]
                             if avg_price is not None: base_value = float(avg_price) * 10 # Assume avg sale of 10 units
                         except: pass # Ignore if fetch fails

                    # Simulate values with a general upward trend and some noise
                    simulated_values = [base_value * (1 + (i * 0.05) + np.random.normal(0, 0.02)) for i in range(len(periods))]
                    
                    df_forecast = pd.DataFrame({"Period": periods, "Projected Value": simulated_values})
                    df_forecast['PeriodLabel'] = df_forecast['Period'].dt.strftime('%Y-%m') if forecast_period in ["3 months", "6 months"] else df_forecast['Period'].dt.strftime('%Y') + (df_forecast['Period'].dt.quarter.astype(str).apply(lambda x: f"-Q{x}") if forecast_period == "1 year" else "")

                    fig = px.line(
                        df_forecast, 
                        x="PeriodLabel", 
                        y="Projected Value",
                        title=f"{forecast_period} Projected Trend for {forecast_subject_text}",
                        labels={"PeriodLabel": "Period", "Projected Value": "Projected Value ($)"},
                        markers=True
                    )
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Save forecast data (simplified - saving just the final value)
                    try:
                        final_projected_value = simulated_values[-1] if simulated_values else 0.0
                        cur.execute(
                            """INSERT INTO market_data 
                            (business_id, industry, metric, value, date, source) 
                            VALUES (%s, %s, %s, %s, %s, %s)""",
                            (business_id, forecast_subject_text, f"{forecast_period}_forecast_est", float(final_projected_value), datetime.now().date(), "AI Forecast")
                        )
                        conn.commit()
                        st.success("Forecast data saved.")
                    except Exception as e:
                        st.error(f"Error saving forecast data: {e}")

                except Exception as e:
                    st.error(f"Error generating forecast chart: {e}")
        
    cur.close(); conn.close()

# AI Chatbot Module (for Business user, talking to AI)
def chatbot_module(business_id, ai_models): 
    """Streamlit module for the AI Market Doubt Assistant."""
    st.header("🤖 Market Doubt Assistant (AI General Helper)")
    if "ai_chat_history" not in st.session_state: st.session_state.ai_chat_history = []
    
    voice_input = st.checkbox("Use Voice Input for AI Assistant", key="ai_voice_input")
    user_input_ai = ""
    
    # Use a form to handle text input and voice input submission consistently
    with st.form("ai_chat_form", clear_on_submit=True):
        if voice_input:
            r = sr.Recognizer()
            audio_buffer = BytesIO()
            
            # Streamlit's audio recorder is better for this
            audio_bytes = st.audio_recorder("Click to record your question for the AI", key="ai_audio_recorder")
            
            if audio_bytes:
                with st.spinner("Transcribing audio..."):
                    try:
                        # Save bytes to a temporary file or buffer for speech_recognition
                        audio_file = BytesIO(audio_bytes)
                        with sr.AudioFile(audio_file) as source:
                            audio = r.record(source)
                        user_input_ai = r.recognize_google(audio)
                        st.text_area("You said:", user_input_ai, key="ai_user_said_voice_display", disabled=True) # Display transcribed text
                    except sr.UnknownValueError:
                        st.error("Could not understand audio. Please try speaking more clearly.")
                    except sr.RequestError as e:
                        st.error(f"Could not request results from speech recognition service; {e}")
                    except Exception as e:
                        st.error(f"An unexpected error occurred during transcription: {e}")
        else:
            user_input_ai = st.text_input("Ask the AI Assistant:", key="ai_user_text_input")

        send_button = st.form_submit_button("Send to AI")

    # Process input after form submission
    if send_button and user_input_ai.strip():
        with st.spinner("AI is thinking..."):
            try:
                response = ai_models.generate_response(user_input_ai) # Using the general chatbot model
                st.session_state.ai_chat_history.append(("You", user_input_ai))
                st.session_state.ai_chat_history.append(("AI Assistant", response))
            except Exception as e:
                st.error(f"Error generating AI response: {e}")
    elif send_button and not user_input_ai.strip():
        st.warning("Please enter a message or record your voice.")

    # Display chat history
    st.write("---")
    st.subheader("Conversation History")
    
    # Display in reverse order to show latest messages at the bottom
    for speaker, text in reversed(st.session_state.ai_chat_history):
        if speaker == "You": 
            st.markdown(f"**You**: {text}")
        else:
            st.markdown(f"**AI Assistant**: {text}")
            # Add Text-to-Speech for AI responses
            try:
                tts = gTTS(text=text, lang='en'); audio_file = BytesIO()
                tts.write_to_fp(audio_file); audio_file.seek(0)
                st.audio(audio_file, format='audio/mp3', key=f"tts_ai_{hash(text)}") # Unique key for each audio player
            except Exception as e:
                st.warning(f"Could not generate audio for this response: {e}")
            st.write("---")

# Investor & Agent Dashboards (Business view of their listed investors)
def investor_dashboard(business_id, ai_models):
    """Streamlit module for Business users to manage Investors."""
    st.header("💰 Investor & Agent Dashboards (Your Listed Contacts)")
    conn = get_db_connection(); cur = conn.cursor()
    tab1, tab2, tab3 = st.tabs(["Investor Directory", "Portfolio Analytics (Simulated)", "Deal Flow (Simulated)"])
    
    with tab1:
        st.subheader("Investor Directory")
        # Fetch investors linked to THIS business OR investors registered independently
        # For a business dashboard, we primarily show investors they have added or interacted with.
        # The current schema links investors to a business via business_id.
        # We'll show investors where business_id is THIS business's ID.
        cur.execute("SELECT id, name, firm, email, investment_focus, portfolio_companies, last_contact, profile_description, website_url, linkedin_profile FROM investors WHERE business_id = %s ORDER BY name", (business_id,))
        investors = cur.fetchall()
        
        if investors:
            df = pd.DataFrame(investors, columns=[
                "ID", "Name", "Firm", "Email", "Investment Focus", 
                "Portfolio Companies", "Last Contact", "Profile Description", "Website", "LinkedIn"
            ])
            st.dataframe(df, hide_index=True)
        else:
            st.info("No investors linked to your business yet.")
        
        # Add new investor linked to this business
        with st.expander("Add New Investor (Link to Your Business)"):
            with st.form("add_investor_to_biz"):
                name = st.text_input("Investor Name")
                firm = st.text_input("Firm")
                email = st.text_input("Email") # This email should ideally be unique across all investors
                focus = st.text_input("Investment Focus")
                companies = st.text_input("Portfolio Companies (comma separated)")
                last_contact = st.date_input("Last Contact Date", datetime.now().date())
                profile_description = st.text_area("Profile Description")
                website_url = st.text_input("Website URL")
                linkedin_profile = st.text_input("LinkedIn Profile URL")

                if st.form_submit_button("Add Investor"):
                    companies_list = [c.strip() for c in companies.split(",")] if companies else []
                    try:
                        # Check if email already exists for an investor (linked or unlinked)
                        cur.execute("SELECT id FROM investors WHERE email = %s", (email,))
                        existing_investor = cur.fetchone()

                        if existing_investor:
                            st.warning(f"An investor with this email ({email}) already exists (ID: {existing_investor[0]}). Linking this existing investor to your business.")
                            # Update the existing investor to link them to this business
                            cur.execute("UPDATE investors SET business_id = %s WHERE id = %s", (business_id, existing_investor[0]))
                            conn.commit()
                            st.success("Existing investor linked to your business!")
                        else:
                            # Insert a new investor linked to this business
                            cur.execute(
                                """INSERT INTO investors 
                                (business_id, name, firm, email, investment_focus, portfolio_companies, last_contact, profile_description, website_url, linkedin_profile) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, name, firm, email, focus, companies_list, last_contact, profile_description, website_url, linkedin_profile)
                            )
                            conn.commit()
                            st.success("New investor added and linked to your business!")
                        st.rerun()
                    except psycopg2.IntegrityError:
                        st.error("An investor with this email already exists and cannot be linked.")
                    except Exception as e:
                        st.error(f"Error adding/linking investor: {e}")

    with tab2:
        st.subheader("Portfolio Analytics (Simulated)")
        st.info("This section provides simulated analytics based on the investors you've added.")
        
        cur.execute("SELECT investment_focus, portfolio_companies FROM investors WHERE business_id = %s", (business_id,))
        investor_data_for_analytics = cur.fetchall()
        
        if investor_data_for_analytics:
            df_analytics = pd.DataFrame(investor_data_for_analytics, columns=["Investment Focus", "Portfolio Companies"])
            
            # Investment focus distribution
            st.write("### Investment Focus Distribution")
            focus_counts = df_analytics["Investment Focus"].value_counts().reset_index()
            focus_counts.columns = ["Focus Area", "Count"]
            if not focus_counts.empty:
                fig = px.bar(
                    focus_counts, 
                    x="Focus Area", 
                    y="Count", 
                    title="Investor Focus Areas (Your Network)"
                )
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("No investment focus data available from your linked investors.")

            # Portfolio composition (simulated based on listed companies)
            st.write("### Portfolio Company Sectors (Simulated)")
            all_companies = []
            for _, row in df_analytics.iterrows():
                portfolio_companies_list = row["Portfolio Companies"] 
                if portfolio_companies_list and isinstance(portfolio_companies_list, list):
                    all_companies.extend(portfolio_companies_list)
            
            if all_companies:
                df_companies = pd.DataFrame({"Company": all_companies})
                # Simple sector mapping (simulated)
                sector_map = {
                    "TechStart": "Technology", "DataAI": "Technology", "CloudScale": "Technology",
                    "HealthCo": "Healthcare", "FinCorp": "Finance", "RetailHub": "Retail",
                    "ManuFab": "Manufacturing", "EduTech": "Education", "PropDev": "Real Estate",
                    "HospGroup": "Hospitality"
                }
                df_companies['Sector'] = df_companies['Company'].apply(lambda x: sector_map.get(x, "Other"))

                sector_counts = df_companies["Sector"].value_counts().reset_index()
                sector_counts.columns = ["Sector", "Count"]
                
                if not sector_counts.empty:
                    fig = px.pie(
                        sector_counts, 
                        values="Count", 
                        names="Sector", 
                        title="Portfolio Companies by Sector (Simulated)"
                    )
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("No portfolio company data available from your linked investors.")
            else:
                st.info("No portfolio company data available from your linked investors.")

        else:
            st.info("Add investors to your directory to see simulated analytics.")
    
    with tab3:
        st.subheader("Deal Flow Management (Simulated)")
        st.info("This section is a placeholder for managing potential investment deals.")
        
        # Add new deal (simulated)
        with st.expander("Add New Deal (Simulated)"):
            with st.form("add_deal_simulated"):
                company = st.text_input("Company Name (Your Business)", value="Your Business Name", disabled=True)
                stage = st.selectbox("Deal Stage", [
                    "Initial Contact", "Pitch Meeting", "Due Diligence", 
                    "Term Sheet", "Closed", "Rejected"
                ], key="deal_stage_sim")
                amount = st.number_input("Potential Amount ($)", min_value=0, step=1000, key="deal_amount_sim")
                investor_contact = st.text_input("Investor Contact Name", key="deal_investor_contact_sim")
                next_step = st.text_input("Next Step", key="deal_next_step_sim")
                target_date = st.date_input("Target Date", datetime.now().date() + timedelta(days=30), key="deal_target_date_sim")
                notes = st.text_area("Notes", key="deal_notes_sim")

                if st.form_submit_button("Add Deal to Pipeline (Simulated)"):
                    # In a real app, this would save to a 'deals' or 'pipeline' table
                    st.success("Deal added to simulated pipeline!")
                    # You might store this in session state for the duration of the session if not using a DB table
                    if 'simulated_deals' not in st.session_state:
                        st.session_state.simulated_deals = []
                    st.session_state.simulated_deals.append({
                        "company": company, "stage": stage, "amount": amount,
                        "contact": investor_contact, "next_step": next_step,
                        "date": target_date.strftime("%Y-%m-%d"), "notes": notes
                    })

        # Simulated deal flow pipeline visualization
        st.write("### Simulated Deal Flow Pipeline")
        # Use session state if deals were added in the form, otherwise use hardcoded examples
        simulated_deals_list = st.session_state.get('simulated_deals', [])
        if not simulated_deals_list:
             # Hardcoded examples if no deals added in session
             simulated_deals_list = [
                {
                    "company": "Your Business",
                    "stage": "Initial Contact",
                    "amount": 500000,
                    "contact": "Sarah Johnson",
                    "next_step": "Schedule pitch meeting",
                    "date": (datetime.now() + timedelta(days=7)).strftime("%Y-%m-%d")
                },
                {
                    "company": "Your Business",
                    "stage": "Pitch Meeting",
                    "amount": 1000000,
                    "contact": "Michael Chen",
                    "next_step": "Send follow-up materials",
                    "date": (datetime.now() + timedelta(days=3)).strftime("%Y-%m-%d")
                },
                {
                    "company": "Your Business",
                    "stage": "Due Diligence",
                    "amount": 2000000,
                    "contact": "David Wilson",
                    "next_step": "Review financials",
                    "date": (datetime.now() + timedelta(days=1)).strftime("%Y-%m-%d")
                }
            ]


        if simulated_deals_list:
            df_deals = pd.DataFrame(simulated_deals_list)
            
            # Ensure stages are in a specific order for the funnel
            ordered_stages = ["Initial Contact", "Pitch Meeting", "Due Diligence", "Term Sheet", "Closed", "Rejected"]
            df_deals['stage'] = pd.Categorical(df_deals['stage'], categories=ordered_stages, ordered=True)
            df_deals = df_deals.sort_values('stage')

            # Aggregate amount by stage for the funnel chart
            df_funnel = df_deals.groupby('stage')['amount'].sum().reset_index()
            df_funnel.columns = ['Stage', 'Total Potential Amount']

            if not df_funnel.empty:
                fig = px.funnel(
                    df_funnel, 
                    x="Total Potential Amount", 
                    y="Stage", 
                    title="Simulated Deal Flow Pipeline Value"
                )
                st.plotly_chart(fig, use_container_width=True)
            else:
                 st.info("No deal data available for funnel visualization.")

            # Deal details
            st.write("### Simulated Deal Details")
            for deal in simulated_deals_list:
                with st.expander(f"{deal['company']} - {deal['stage']} (${deal['amount']:,})"):
                    st.write(f"**Investor Contact:** {deal['contact']}")
                    st.write(f"**Next Step:** {deal['next_step']}")
                    st.write(f"**Target Date:** {deal['date']}")
                    if deal.get('notes'): st.write(f"**Notes:** {deal['notes']}")

        else:
            st.info("No simulated deals in the pipeline yet.")

    cur.close(); conn.close()

# Govt/Private Schemes Module
def schemes_module(business_id, ai_models):
    """Streamlit module for discovering Govt/Private Schemes and News Alerts."""
    st.header("🏛️ Govt/Private Schemes & News Alerts")
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    tab1, tab2 = st.tabs(["Available Schemes", "News Alerts (Simulated)"])
    
    with tab1:
        st.subheader("Available Schemes")
        # Fetch schemes linked to this business or potentially public/relevant ones
        # Assuming schemes are added by the business for tracking, or could be fetched from an external source.
        # For now, we fetch schemes linked to this business.
        cur.execute("SELECT id, name, description, eligibility, benefits, deadline, sector, is_govt FROM schemes WHERE business_id = %s ORDER BY deadline", (business_id,))
        schemes = cur.fetchall()
        
        if schemes:
            df = pd.DataFrame(schemes, columns=[
                "ID", "Name", "Description", "Eligibility", 
                "Benefits", "Deadline", "Sector", "Is Govt"
            ])
            
            # Filter options
            all_sectors = df["Sector"].unique().tolist() if not df.empty else []
            sector_filter = st.multiselect(
                "Filter by Sector",
                options=all_sectors,
                default=all_sectors,
                key="scheme_sector_filter"
            )
            
            govt_filter = st.checkbox("Government Schemes Only", value=False, key="scheme_govt_filter")
            
            # Apply filters
            filtered_df = df[df["Sector"].isin(sector_filter)]
            if govt_filter:
                filtered_df = filtered_df[filtered_df["Is Govt"] == True]
            
            if not filtered_df.empty:
                st.dataframe(filtered_df.drop(columns=["ID"]), hide_index=True)

                # Display schemes with expanders
                st.write("#### Scheme Details")
                for _, row in filtered_df.iterrows():
                    days_left = (row['Deadline'] - datetime.now().date()).days if row['Deadline'] else None
                    deadline_status = ""
                    if days_left is not None:
                        if days_left > 0: deadline_status = f"⏰ {days_left} days left"
                        elif days_left == 0: deadline_status = "⚠️ Deadline TODAY!"
                        else: deadline_status = "❌ Deadline passed"
                    else:
                         deadline_status = "No deadline specified"

                    with st.expander(f"{row['Name']} ({'Govt' if row['Is Govt'] else 'Private'}) - Deadline: {row['Deadline'] or 'N/A'} {deadline_status}"):
                        st.write(f"**Sector:** {row['Sector'] or 'N/A'}")
                        st.write(f"**Description:** {row['Description'] or 'N/A'}")
                        st.write(f"**Eligibility:** {row['Eligibility'] or 'N/A'}")
                        st.write(f"**Benefits:** {row['Benefits'] or 'N/A'}")
                        
                        # Simulated Apply button
                        if st.button("Apply Now (Simulated)", key=f"apply_scheme_{row['ID']}"):
                            st.info("This would typically link to an external application page or an internal application process.")
            else:
                st.info("No schemes found matching your filters.")
        else:
            st.info("No schemes added to your database yet.")
        
        # Add new scheme
        with st.expander("Add New Scheme (for Tracking)"):
            with st.form("add_scheme"):
                name = st.text_input("Scheme Name")
                description = st.text_area("Description")
                eligibility = st.text_area("Eligibility Criteria")
                benefits = st.text_area("Benefits")
                deadline = st.date_input("Deadline", value=date.today() + timedelta(days=60))
                sector = st.text_input("Sector")
                is_govt = st.checkbox("Government Scheme", value=True)
                
                if st.form_submit_button("Add Scheme"):
                    if not name or not description or not deadline:
                        st.error("Name, Description, and Deadline are required.")
                    else:
                        try:
                            cur.execute(
                                """INSERT INTO schemes 
                                (business_id, name, description, eligibility, benefits, deadline, sector, is_govt) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, name, description, eligibility, benefits, deadline, sector, is_govt)
                            )
                            conn.commit()
                            st.success("Scheme added successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error adding scheme: {e}")
    
    with tab2:
        st.subheader("Latest Business News Alerts (Simulated)")
        st.info("These are simulated news alerts, not real-time feeds.")
        
        # Simulated news alerts
        alerts = [
            {
                "title": "New Export Promotion Scheme Announced",
                "summary": "Government launches scheme to boost exports in manufacturing sector",
                "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "priority": "High"
            },
            {
                "title": "Tax Incentives for R&D Investments",
                "summary": "New policy offers 150% deduction on R&D spending for eligible businesses",
                "date": (datetime.now() - timedelta(hours=2)).strftime("%Y-%m-%d %H:%M"),
                "priority": "Medium"
            },
            {
                "title": "Rural Business Grant Program",
                "summary": "Applications open for businesses operating in rural areas",
                "date": (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d %H:%M"),
                "priority": "Low"
            }
        ]
        
        for alert in alerts:
            with st.expander(f"{alert['date']}: {alert['title']} ({alert['priority']} Priority)"):
                st.write(alert['summary'])
                if st.button("Learn More (Simulated)", key=f"alert_{hash(alert['title'])}"): # Use hash for unique key
                    st.info("More details would appear here or link to an external news source.")
    
    cur.close(); conn.close()

# Opportunity Director Module
def opportunities_module(business_id, ai_models):
    """Streamlit module for tracking Business Leads, Grants, and Competitions."""
    st.header("🎯 Opportunity Director")
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    tab1, tab2, tab3 = st.tabs(["Business Leads", "Grants", "Competitions"])
    
    with tab1:
        st.subheader("Business Leads")
        
        cur.execute("SELECT id, title, description, category, deadline, reward, link FROM opportunities WHERE business_id = %s AND category = 'lead' ORDER BY deadline", (business_id,))
        leads = cur.fetchall()
        
        if leads:
            for lead in leads:
                days_left = (lead[4] - datetime.now().date()).days if lead[4] else None
                deadline_status = ""
                if days_left is not None:
                    if days_left > 0: deadline_status = f"⏰ {days_left} days left"
                    elif days_left == 0: deadline_status = "⚠️ Deadline TODAY!"
                    else: deadline_status = "❌ Expired"
                else:
                    deadline_status = "No deadline specified"

                with st.expander(f"{lead[1]} (Deadline: {lead[4] or 'N/A'}) {deadline_status}"):
                    st.write(f"**Description:** {lead[2] or 'N/A'}")
                    st.write(f"**Potential Reward:** {lead[5] or 'N/A'}")
                    if lead[6]: st.markdown(f"**Link:** [View Opportunity]({lead[6]})")
                    else: st.write("**Link:** N/A")

        else:
            st.info("No business leads added yet.")
        
        # Add new lead
        with st.expander("Add New Lead"):
            with st.form("add_lead"):
                title = st.text_input("Lead Title")
                description = st.text_area("Description")
                reward = st.text_input("Potential Reward")
                deadline = st.date_input("Deadline", value=date.today() + timedelta(days=30))
                link = st.text_input("Link (Optional)")
                
                if st.form_submit_button("Add Lead"):
                    if not title or not description or not deadline:
                        st.error("Title, Description, and Deadline are required.")
                    else:
                        try:
                            cur.execute(
                                """INSERT INTO opportunities 
                                (business_id, title, description, category, deadline, reward, link) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, title, description, "lead", deadline, reward, link)
                            )
                            conn.commit()
                            st.success("Lead added successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error adding lead: {e}")
    
    with tab2:
        st.subheader("Grant Opportunities")
        
        cur.execute("SELECT id, title, description, category, deadline, reward, link FROM opportunities WHERE business_id = %s AND category = 'grant' ORDER BY deadline", (business_id,))
        grants = cur.fetchall()
        
        if grants:
            for grant in grants:
                days_left = (grant[4] - datetime.now().date()).days if grant[4] else None
                deadline_status = ""
                if days_left is not None:
                    if days_left > 0: deadline_status = f"⏰ {days_left} days left"
                    elif days_left == 0: deadline_status = "⚠️ Deadline TODAY!"
                    else: deadline_status = "❌ Expired"
                else:
                    deadline_status = "No deadline specified"

                with st.expander(f"{grant[1]} (Deadline: {grant[4] or 'N/A'}) {deadline_status}"):
                    st.write(f"**Description:** {grant[2] or 'N/A'}")
                    st.write(f"**Amount:** {grant[5] or 'N/A'}")
                    if grant[6]: st.markdown(f"**Application Link:** [Apply Here]({grant[6]})")
                    else: st.write("**Application Link:** N/A")

        else:
            st.info("No grant opportunities added yet.")
        
        # Add new grant
        with st.expander("Add New Grant"):
            with st.form("add_grant"):
                title = st.text_input("Grant Name")
                description = st.text_area("Description")
                amount = st.text_input("Grant Amount")
                deadline = st.date_input("Deadline", value=date.today() + timedelta(days=60))
                link = st.text_input("Application Link (Optional)")
                
                if st.form_submit_button("Add Grant"):
                    if not title or not description or not deadline:
                        st.error("Name, Description, and Deadline are required.")
                    else:
                        try:
                            cur.execute(
                                """INSERT INTO opportunities 
                                (business_id, title, description, category, deadline, reward, link) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, title, description, "grant", deadline, amount, link)
                            )
                            conn.commit()
                            st.success("Grant opportunity added successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error adding grant: {e}")
    
    with tab3:
        st.subheader("Business Competitions")
        
        cur.execute("SELECT id, title, description, category, deadline, reward, link FROM opportunities WHERE business_id = %s AND category = 'competition' ORDER BY deadline", (business_id,))
        competitions = cur.fetchall()
        
        if competitions:
            for comp in competitions:
                days_left = (comp[4] - datetime.now().date()).days if comp[4] else None
                deadline_status = ""
                if days_left is not None:
                    if days_left > 0: deadline_status = f"⏰ {days_left} days left"
                    elif days_left == 0: deadline_status = "⚠️ Deadline TODAY!"
                    else: deadline_status = "❌ Expired"
                else:
                    deadline_status = "No deadline specified"

                with st.expander(f"{comp[1]} (Deadline: {comp[4] or 'N/A'}) {deadline_status}"):
                    st.write(f"**Description:** {comp[2] or 'N/A'}")
                    st.write(f"**Prize:** {comp[5] or 'N/A'}")
                    if comp[6]: st.markdown(f"**Registration Link:** [Register Here]({comp[6]})")
                    else: st.write("**Registration Link:** N/A")

        else:
            st.info("No competitions added yet.")
        
        # Add new competition
        with st.expander("Add New Competition"):
            with st.form("add_competition"):
                title = st.text_input("Competition Name")
                description = st.text_area("Description")
                prize = st.text_input("Prize")
                deadline = st.date_input("Deadline", value=date.today() + timedelta(days=45))
                link = st.text_input("Registration Link (Optional)")
                
                if st.form_submit_button("Add Competition"):
                    if not title or not description or not deadline:
                        st.error("Name, Description, and Deadline are required.")
                    else:
                        try:
                            cur.execute(
                                """INSERT INTO opportunities 
                                (business_id, title, description, category, deadline, reward, link) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, title, description, "competition", deadline, prize, link)
                            )
                            conn.commit()
                            st.success("Competition added successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error adding competition: {e}")
    
    cur.close(); conn.close()

# Voice Navigation Module
class AudioRecorderProcessor(AudioProcessorBase):
    """
    A custom audio processor for streamlit-webrtc.
    It collects audio frames and makes them available as a single byte buffer.
    """
    def __init__(self):
        super().__init__()
        self.frames_bytes = []  # Buffer to store raw bytes of audio frames
        self.sample_rate = None
        self.sample_width = None # Bytes per sample (e.g., 2 for int16)
        self._is_recording = True # Internal flag to control frame accumulation

    def recv(self, frame: av.AudioFrame) -> av.AudioFrame:
        """Receives audio frames from WebRTC, processes, and stores them."""
        if not self._is_recording:
            return frame # Pass through if not actively recording

        if not self.frames_bytes:  # On the first frame received
            self.sample_rate = frame.sample_rate
            # Determine sample_width from frame format for speech_recognition.AudioData
            if frame.format.name == 's16': # Signed 16-bit PCM
                self.sample_width = 2
            elif frame.format.name in ['fltp', 'flt']: # Float PCM (planar or non-planar)
                self.sample_width = 2 # We will convert this to s16 for speech_recognition
            else:
                # Fallback for other formats, assuming conversion to s16
                st.warning(f"Audio frame format is {frame.format.name}, attempting s16 conversion. Sample width set to 2.")
                self.sample_width = 2
        
        # Convert PyAV frame to a NumPy array
        np_frame_data = frame.to_ndarray() # Shape is (channels, samples) or (samples,) for mono

        # Ensure mono audio (speech_recognition usually prefers mono)
        if np_frame_data.ndim > 1 and np_frame_data.shape[0] > 1: # If stereo
            # Convert to mono by averaging channels or taking the first one
            np_frame_mono = np.mean(np_frame_data, axis=0).astype(np_frame_data.dtype)
            # Alternatively: np_frame_mono = np_frame_data[0, :] 
        else:
            np_frame_mono = np_frame_data.flatten()

        # Convert to 16-bit PCM if it's float (speech_recognition.AudioData expects PCM)
        if np_frame_mono.dtype in [np.float32, np.float64]:
            np_frame_mono_int16 = (np_frame_mono * 32767).astype(np.int16)
        elif np_frame_mono.dtype == np.int16:
            np_frame_mono_int16 = np_frame_mono
        else:
            # Attempt conversion for other unexpected types
            st.warning(f"Unexpected audio frame dtype: {np_frame_mono.dtype}. Attempting int16 conversion.")
            np_frame_mono_int16 = np_frame_mono.astype(np.int16)
        
        self.frames_bytes.append(np_frame_mono_int16.tobytes())
        
        return frame # Echo the frame back (standard practice)

    def get_recorded_data(self) -> tuple[bytes | None, int | None, int | None]:
        """Returns the accumulated audio data and resets the buffer."""
        if not self.frames_bytes:
            return None, None, None
        
        audio_data = b"".join(self.frames_bytes)
        # Clear buffer for next potential recording
        self.frames_bytes = [] 
        # Keep sample_rate and sample_width as they were detected for the last recording session
        return audio_data, self.sample_rate, self.sample_width
    
    def stop_recording(self):
        """Signals that recording should stop accumulating frames."""
        self._is_recording = False

    def start_recording(self):
        """Resets the processor state for a new recording."""
        self._is_recording = True
        self.frames_bytes = []
        self.sample_rate = None
        self.sample_width = None


def voice_navigation(business_id, ai_models): # ai_models might not be used directly here
    """Streamlit module for Voice Navigation using streamlit-webrtc."""
    st.header("🎙️ Voice Navigation")
    st.info("Click 'Start' to record your command using your microphone. Click 'Stop' when finished to process.")
    st.caption("Ensure your browser has microphone permissions enabled for this site.")

    r = sr.Recognizer()
    webrtc_ctx = webrtc_streamer(
        key="voice_nav_webrtc",
        mode=WebRtcMode.SENDONLY,
        audio_processor_factory=AudioRecorderProcessor, # Pass the class itself
        media_stream_constraints={"audio": True, "video": False},
        rtc_configuration={"iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]},
        # desired_playing_state can be used to control start/stop programmatically if needed
    )

    if webrtc_ctx.state.playing:
        st.info("🎤 Recording... Click 'Stop' on the recorder above when finished.")
    
    # Check if the stream has been stopped by the user and if the processor is available
    if not webrtc_ctx.state.playing and webrtc_ctx.audio_processor:
        # Access the processor instance that was used for the stream
        audio_processor_instance = webrtc_ctx.audio_processor
        
        if 'voice_nav_processed' not in st.session_state:
            st.session_state.voice_nav_processed = False

        if not st.session_state.voice_nav_processed and audio_processor_instance.frames_bytes: # Check internal buffer before calling get
            audio_bytes_data, sample_rate, sample_width = audio_processor_instance.get_recorded_data()
            st.session_state.voice_nav_processed = True # Mark as processed for this stop event

            if audio_bytes_data and sample_rate and sample_width:
                st.success("Audio captured via WebRTC. Transcribing...")
                with st.spinner("Transcribing command..."):
                    try:
                        audio_data_sr = sr.AudioData(audio_bytes_data, sample_rate, sample_width)
                        command = r.recognize_google(audio_data_sr)
                        st.success(f"You said: \"{command}\"")
                        
                        command_lower = command.lower()
                        nav_target = None

                        if "inventory" in command_lower or "billing" in command_lower or "product" in command_lower:
                            nav_target = "Inventory & Billing"
                        elif "hr" in command_lower or "human resources" in command_lower or "employee" in command_lower:
                            nav_target = "HR Tools"
                        elif "project" in command_lower or "gantt" in command_lower:
                            nav_target = "Project Manager"
                        elif "document" in command_lower or "template" in command_lower or "library" in command_lower:
                            nav_target = "Document Generator"
                        elif "market analysis" in command_lower or "trend" in command_lower or "competitor" in command_lower:
                            nav_target = "Market Analysis Tool"
                        elif "chat" in command_lower or "assistant" in command_lower or "ai" in command_lower:
                            nav_target = "Market Doubt Assistant (AI Chatbot)"
                        elif "investor" in command_lower or "funding" in command_lower or "deal" in command_lower:
                            nav_target = "Investor & Agent Dashboards"
                        elif "scheme" in command_lower or "grant" in command_lower or "news" in command_lower:
                            nav_target = "Govt/Private Schemes & News Alerts"
                        elif "opportunity" in command_lower or "lead" in command_lower or "competition" in command_lower:
                            nav_target = "Opportunity Director"
                        elif "pitch" in command_lower or "deck" in command_lower or "script" in command_lower:
                            nav_target = "Pitching Helper"
                        elif "strategy" in command_lower or "playbook" in command_lower:
                            nav_target = "Strategy Generator"
                        elif "hiring" in command_lower or "job" in command_lower or "onboarding" in command_lower:
                            nav_target = "Hiring Helper"
                        elif "tax" in command_lower or "gst" in command_lower or "filing" in command_lower:
                            nav_target = "Tax & GST Filing"
                        elif "ipo" in command_lower or "cap table" in command_lower:
                            nav_target = "IPO & Cap Table Management"
                        elif "legal" in command_lower or "ca" in command_lower or "insurance" in command_lower or "marketplace" in command_lower:
                            nav_target = "Legal, CA & Insurance Marketplace"
                        elif "intelligence" in command_lower or "analytics" in command_lower or "report" in command_lower:
                            nav_target = "Enterprise Intelligence Dashboards"
                        elif "forecasting" in command_lower or "predictive" in command_lower:
                            nav_target = "AI Market Forecasting"
                        elif "dashboard" in command_lower or "overview" in command_lower:
                            nav_target = "Dashboard"
                        elif "message" in command_lower or "chat with" in command_lower:
                             nav_target = "Messaging"

                        if nav_target:
                            st.info(f"Simulating navigation to: **{nav_target}**")
                            # st.rerun() 
                        else:
                            st.info("Command not recognized or no specific navigation target found.")

                    except sr.UnknownValueError:
                        st.error("Google Speech Recognition could not understand the audio.")
                    except sr.RequestError as e:
                        st.error(f"Could not request results from Google Speech Recognition service; {e}")
                    except Exception as e:
                        st.error(f"An unexpected error occurred during transcription: {e}")
            
    if webrtc_ctx.state.playing and st.session_state.get('voice_nav_processed', False):
        st.session_state.voice_nav_processed = False
        if webrtc_ctx.audio_processor: # Call start on the new/reused processor instance
            webrtc_ctx.audio_processor.start_recording()

# Pitching Helper Module
def pitching_helper(business_id, ai_models):
    """Streamlit module for Pitching Assistance."""
    st.header("📢 Pitching Helper")
    
    tab1, tab2, tab3 = st.tabs(["Pitch Deck Generator", "Funding Scripts", "Investor Prep"])
    
    with tab1:
        st.subheader("AI-Crafted Pitch Decks")
        st.info("Generates a text outline for a pitch deck. Not a visual deck creator.")
        
        with st.form("pitch_deck_input"):
            company_name = st.text_input("Company Name")
            business_description = st.text_area("Business Description")
            problem = st.text_area("Problem Statement")
            solution = st.text_area("Your Solution")
            market_size = st.text_input("Market Size")
            business_model = st.text_input("Business Model")
            funding_amount = st.text_input("Funding Amount Sought")
            target_audience = st.text_input("Target Audience (e.g., Angel Investors, VC Firms)")
            
            if st.form_submit_button("Generate Pitch Deck Outline"):
                if not company_name or not business_description or not problem or not solution or not funding_amount:
                    st.error("Please fill in required fields.")
                else:
                    with st.spinner("Creating your pitch deck outline..."):
                        # Generate slides content
                        prompt = f"""
                        Create a pitch deck outline for {company_name} targeting {target_audience or 'investors'}.
                        
                        Business: {business_description}
                        Problem: {problem}
                        Solution: {solution}
                        Market: {market_size or 'Not specified'}
                        Model: {business_model or 'Not specified'}
                        Funding Ask: {funding_amount}
                        
                        Include standard pitch deck sections (e.g., Problem, Solution, Market, Business Model, Team, Financials, Ask) with titles and key bullet points for each slide. Aim for about 10-15 slides.
                        """
                        
                        deck_content = ai_models.generate_text(prompt, max_length=1500)
                        
                        st.subheader("Generated Pitch Deck Outline")
                        st.text_area("Outline Content", deck_content, height=500, key="generated_pitch_deck_content")
                        
                        # Create downloadable doc
                        try:
                            doc = docx.Document()
                            doc.add_heading(f"{company_name} Pitch Deck Outline", 0)
                            doc.add_paragraph(f"Generated on: {datetime.now().date()}")
                            doc.add_paragraph(f"Target Audience: {target_audience or 'Investors'}")
                            doc.add_paragraph("\n---\n")
                            
                            # Add generated content
                            doc.add_paragraph(deck_content)
                            
                            # Save to buffer
                            docx_file = BytesIO()
                            doc.save(docx_file)
                            docx_file.seek(0)
                            
                            st.download_button(
                                "Download Pitch Deck Outline (DOCX)",
                                data=docx_file,
                                file_name=f"{company_name.replace(' ', '_')}_Pitch_Deck_Outline.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e:
                            st.error(f"Error creating DOCX file: {e}")

    with tab2:
        st.subheader("Funding Script Generator")
        st.info("Generates a text script for investor conversations.")
        
        with st.form("funding_script_input"):
            investor_type = st.selectbox("Investor Type", [
                "Angel Investor", "VC Firm", "Corporate Investor", "Crowdfunding", "Other"
            ], key="script_investor_type")
            meeting_type = st.selectbox("Meeting Type", [
                "Initial Pitch", "Follow-up", "Due Diligence Call", "Term Negotiation", "Casual Chat"
            ], key="script_meeting_type")
            key_points = st.text_area("Key Points to Cover (e.g., 'traction metrics', 'team background', 'use of funds')", key="script_key_points")
            
            if st.form_submit_button("Generate Script"):
                if not key_points.strip():
                    st.warning("Please enter key points to cover.")
                else:
                    with st.spinner("Creating your funding script..."):
                        prompt = f"""
                        Create a funding conversation script for a {meeting_type} meeting with a {investor_type}.
                        
                        Key points to cover:
                        {key_points}
                        
                        Include sections for introduction, discussing key points, anticipating likely questions and suggesting responses, and closing remarks. Make it sound professional and engaging.
                        """
                        
                        script = ai_models.generate_text(prompt, max_length=1000)
                        
                        st.subheader("Generated Funding Script")
                        st.text_area("Script Content", script, height=500, key="generated_script_content")
                        
                        st.download_button(
                            "Download Script (TXT)",
                            data=script,
                            file_name=f"Funding_Script_{investor_type.replace(' ', '_')}_{meeting_type.replace(' ', '_')}.txt",
                            mime="text/plain"
                        )
    
    with tab3:
        st.subheader("Investor Preparation")
        st.info("Provides AI-generated answers to common investor questions and simulated practice.")
        
        st.write("### Common Investor Questions & Suggested Answers")
        questions = [
            "What problem are you solving?",
            "How big is the market opportunity?",
            "What makes your solution unique?",
            "What's your business model?",
            "What's your customer acquisition strategy?",
            "Who are your competitors?",
            "What are the key risks?",
            "What's your funding ask and how will you use it?",
            "What's your exit strategy?"
        ]
        
        for i, q in enumerate(questions):
            with st.expander(q):
                try:
                    answer_prompt = f"Provide a concise, effective answer for a startup pitching to investors for the question: '{q}'"
                    answer = ai_models.generate_text(answer_prompt, max_length=300)
                    st.write(answer)
                except Exception as e:
                    st.warning(f"Could not generate answer for this question: {e}")

        st.write("### Practice Pitch Session (Simulated)")
        st.warning("This is a simulated feature. In a real app, this would involve recording your pitch and getting AI feedback.")
        if st.button("Start Mock Pitch Session (Simulated)", key="start_mock_pitch_btn"):
            st.info("Simulating a mock pitch session... Imagine you are pitching now!")
            st.write("AI Feedback (Simulated): Your opening was strong, but clarify your market size. Good job on the team slide!")

# Strategy Generator Module
def strategy_generator(business_id, ai_models):
    d=st.download_button
    """Streamlit module for generating business strategies."""
    st.header("♟️ Strategy Generator")
    st.info("Generates a text outline for a business strategy based on your inputs.")
    
    with st.form("strategy_input"):
        business_type = st.text_input("Your Business Type (e.g., SaaS, E-commerce, Consulting)", key="strategy_biz_type")
        business_stage = st.selectbox("Business Stage", [
            "Ideation", "Early-stage", "Growth", "Mature", "Scaling", "Other"
        ], key="strategy_biz_stage")
        challenges = st.text_area("Key Challenges You Are Facing (e.g., 'customer acquisition cost is too high', 'difficulty scaling operations')", key="strategy_challenges")
        goals = st.text_area("Short-term Goals (e.g., 'increase monthly revenue by 15% in 6 months', 'launch new product feature')", key="strategy_short_goals")
        long_term_goals = st.text_area("Long-term Goals (e.g., 'become market leader in 3 years', 'expand to international markets')", key="strategy_long_goals")
        
        if st.form_submit_button("Generate Growth Strategy"):
            if not business_type or not business_stage or not challenges or not goals or not long_term_goals:
                st.error("Please fill in all required fields.")
            else:
                with st.spinner("Creating your personalized growth playbook..."):
                    prompt = f"""
                    Create a detailed growth strategy playbook for a {business_stage} stage {business_type} business.
                    
                    Key Challenges:
                    {challenges}
                    
                    Short-term Goals (3-6 months):
                    {goals}
                    
                    Long-term Goals (1-3 years):
                    {long_term_goals}
                    
                    The playbook should include:
                    1. Executive Summary
                    2. Analysis of Challenges and Opportunities
                    3. Strategic Pillars/Themes
                    4. Key Initiatives for each Strategic Pillar (with brief descriptions)
                    5. Suggested Timeline (e.g., Phase 1, Phase 2)
                    6. Key Performance Indicators (KPIs) to track progress
                    7. Potential Risks and Mitigation Strategies
                    
                    Format the output clearly with headings and bullet points.
                    """
                    
                    strategy = ai_models.generate_text(prompt, max_length=2000)
                    
                    st.subheader("Your Growth Playbook")
                    st.text_area("Strategy Content", strategy, height=600, key="generated_strategy_content")
                    
                    # Create sections with expanders for better readability
                    st.write("#### Playbook Sections")
                    sections = strategy.split("\n\n") # Simple split, might need refinement based on AI output format
                    for section in sections:
                        if section.strip():
                            # Use the first line as the expander title
                            lines = section.strip().split("\n")
                            expander_title = lines[0][:100] + "..." if len(lines[0]) > 100 else lines[0]
                            with st.expander(expander_title):
                                st.write(section) # Display the full section content
                    filename_business_type = business_type if business_type and business_type.strip() else "Unnamed_Strategy"
                    file_name_str = f"{filename_business_type.replace(' ', '_')}_Growth_Playbook.txt"
                    

                    d(
                        "Download Playbook (TXT)",
                        data=strategy,
                        file_name=f"{business_type.replace(' ', '_')}_Growth_Playbook.txt",
                        mime="text/plain"
                      )

# Hiring Helper Module
def hiring_helper(business_id, ai_models):
    """Streamlit module for Hiring Assistance."""
    st.header("👔 Hiring Helper")
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    tab1, tab2, tab3 = st.tabs([
        "Job Openings", 
        "JD Generator", 
        "Onboarding Toolkit (Simulated)"
    ])
    
    with tab1:
        st.subheader("Manage Job Openings")
        
        cur.execute("SELECT id, title, department, description, requirements, experience_needed, posted_date, status FROM job_openings WHERE business_id = %s ORDER BY posted_date DESC", (business_id,))
        jobs = cur.fetchall()
        
        if jobs:
            df = pd.DataFrame(jobs, columns=[
                "ID", "Title", "Department", "Description", 
                "Requirements", "Experience", "Posted Date", "Status"
            ])
            st.dataframe(df.drop(columns=["ID"]), hide_index=True) # Hide ID in display
        else:
            st.info("No job openings posted yet.")
        
        # Add new job opening
        with st.expander("Post New Job Opening"):
            with st.form("add_job"):
                title = st.text_input("Job Title")
                department = st.text_input("Department")
                description = st.text_area("Job Description")
                requirements = st.text_input("Key Requirements (comma separated)")
                experience = st.text_input("Experience Needed (e.g., '3+ years', 'Entry-level')")
                status = st.selectbox("Status", ["Active", "Closed", "On Hold"])
                
                if st.form_submit_button("Post Job"):
                    if not title or not department or not description or not requirements or not experience:
                        st.error("Please fill in all required fields.")
                    else:
                        req_list = [r.strip() for r in requirements.split(",") if r.strip()]
                        try:
                            cur.execute(
                                """INSERT INTO job_openings 
                                (business_id, title, department, description, requirements, experience_needed, posted_date, status) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, title, department, description, req_list, experience, datetime.now().date(), status)
                            )
                            conn.commit()
                            st.success("Job posted successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error posting job: {e}")
    
    with tab2:
        st.subheader("Job Description Generator")
        st.info("Generates a professional job description using AI.")
        
        with st.form("jd_generator"):
            job_title = st.text_input("Job Title")
            department = st.text_input("Department")
            key_responsibilities = st.text_area("Key Responsibilities (list them)")
            qualifications = st.text_area("Required Qualifications (list them)")
            preferred_skills = st.text_area("Preferred Skills (list them, optional)")
            
            if st.form_submit_button("Generate JD"):
                if not job_title or not department or not key_responsibilities or not qualifications:
                    st.error("Please fill in required fields.")
                else:
                    with st.spinner("Creating professional job description..."):
                        prompt = f"""
                        Create a professional job description for a {job_title} position in the {department} department.
                        
                        Key Responsibilities:
                        {key_responsibilities}
                        
                        Required Qualifications:
                        {qualifications}
                        
                        Preferred Skills (Optional):
                        {preferred_skills}

                        Include a brief company overview (placeholder), benefits (placeholder), and equal opportunity employer statement.
                        """
                        
                        jd = ai_models.generate_text(prompt, max_length=1000)
                        
                        st.subheader("Generated Job Description")
                        st.text_area("Job Description Content", jd, height=500, key="generated_jd_content")
                        
                        st.download_button(
                            "Download JD (TXT)",
                            data=jd,
                            file_name=f"JD_{job_title.replace(' ', '_')}.txt",
                            mime="text/plain"
                        )
    
    with tab3:
        st.subheader("Onboarding Toolkit (Simulated)")
        st.info("Provides a checklist and simulated document package for new hires.")
        
        st.write("### New Hire Checklist")
        checklist_items = [
            "Complete HR paperwork",
            "Set up email and system accounts",
            "Provide necessary equipment (laptop, phone, etc.)",
            "Schedule company and team orientation sessions",
            "Assign a mentor or buddy",
            "Plan 30-60-90 day goals with manager",
            "Schedule initial training sessions",
            "Introduce to team members"
        ]
        
        st.write("Mark items as completed:")
        for i, item in enumerate(checklist_items):
            st.checkbox(item, key=f"onboarding_checklist_{i}")
        
        st.write("### Onboarding Documents (Simulated Package)")
        doc_options = [
            "Employee Handbook",
            "Benefits Guide",
            "Company Policies",
            "Team Directory",
            "Project Overview (Relevant Projects)"
        ]
        
        selected_docs = st.multiselect("Select documents to include in the package (Simulated)", doc_options, key="onboarding_doc_select")
        
        if st.button("Generate Onboarding Package (Simulated)", key="generate_onboarding_package_btn"):
            if not selected_docs:
                st.warning("Please select documents to include.")
            else:
                # In a real app, this would compile actual documents.
                # Here, we just simulate the creation and offer a dummy download.
                package_content = f"Simulated Onboarding Package Contents:\n\nIncludes:\n- " + "\n- ".join(selected_docs) + "\n\n(This is a simulated file. Actual document content is not included.)"
                
                st.success("Simulated onboarding package generated!")
                st.download_button(
                    "Download Simulated Package (TXT)",
                    data=package_content,
                    file_name="Simulated_Onboarding_Package.txt",
                    mime="text/plain" # Use text/plain for the simulated content
                )
    
    cur.close(); conn.close()

# Tax & GST Module
def tax_module(business_id, ai_models):
    """Streamlit module for Tax and GST Filing assistance."""
    st.header("🧾 Automated Tax & GST Filing")
    st.info("This module provides tools for tracking tax records and calculating GST. It does NOT perform actual tax filing.")
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    tab1, tab2, tab3 = st.tabs(["Tax Records", "GST Calculator", "Filing Status (Simulated)"])
    
    with tab1:
        st.subheader("Tax Records")
        
        cur.execute("SELECT id, financial_year, total_income, tax_paid, filing_date, status, notes FROM tax_records WHERE business_id = %s ORDER BY financial_year DESC", (business_id,))
        records = cur.fetchall()
        
        if records:
            df = pd.DataFrame(records, columns=[
                "ID", "Financial Year", "Total Income", "Tax Paid", 
                "Filing Date", "Status", "Notes"
            ])
            st.dataframe(df.drop(columns=["ID"]), hide_index=True)
            
            # Tax summary
            st.write("### Tax Summary")
            total_tax = df["Tax Paid"].sum() if not df["Tax Paid"].empty else 0.0
            total_income = df["Total Income"].sum() if not df["Total Income"].empty else 0.0
            avg_rate = (total_tax / total_income) * 100 if total_income > 0 else 0.0
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Total Tax Paid (Recorded)", f"${total_tax:,.2f}")
            with col2:
                st.metric("Average Tax Rate (Recorded)", f"{avg_rate:.1f}%")
        else:
            st.info("No tax records found.")
        
        # Add new record
        with st.expander("Add Tax Record"):
            with st.form("add_tax_record"):
                financial_year = st.text_input("Financial Year (e.g., 2023-24)")
                total_income = st.number_input("Total Income", min_value=0.0, step=1000.0)
                tax_paid = st.number_input("Tax Paid", min_value=0.0, step=1000.0)
                filing_date = st.date_input("Filing Date")
                status = st.selectbox("Status", ["Filed", "Pending", "Revised", "Extension"])
                notes = st.text_area("Notes (Optional)")
                
                if st.form_submit_button("Add Record"):
                    if not financial_year or total_income is None or tax_paid is None or not filing_date or not status:
                        st.error("Please fill in all required fields.")
                    else:
                        try:
                            cur.execute(
                                """INSERT INTO tax_records 
                                (business_id, financial_year, total_income, tax_paid, filing_date, status, notes) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, financial_year, total_income, tax_paid, filing_date, status, notes)
                            )
                            conn.commit()
                            st.success("Tax record added successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error adding tax record: {e}")
    
    with tab2:
        st.subheader("GST Calculator")
        
        col1, col2 = st.columns(2)
        with col1:
            amount = st.number_input("Amount (before GST)", min_value=0.0, step=100.0)
        with col2:
            gst_rate_percent = st.selectbox("GST Rate", ["5%", "12%", "18%", "28%"])
        
        if amount > 0:
            try:
                rate = float(gst_rate_percent.replace("%", "")) / 100
                gst_amount = amount * rate
                total = amount + gst_amount
                
                st.write("### Calculation Results")
                col_base, col_gst, col_total = st.columns(3)
                with col_base:
                    st.metric("Base Amount", f"₹{amount:,.2f}")
                with col_gst:
                    st.metric(f"GST ({gst_rate_percent})", f"₹{gst_amount:,.2f}")
                with col_total:
                    st.metric("Total Amount", f"₹{total:,.2f}")
            except ValueError:
                st.error("Invalid GST rate selected.")
            except Exception as e:
                st.error(f"Error calculating GST: {e}")
            
            # GST filing due dates (Simulated)
            st.write("### Upcoming GST Filing Dates (Simulated)")
            st.info("These are general reminders, not linked to your actual filing status.")
            today = datetime.now().date()
            
            # GSTR-3B due date (typically 20th of next month)
            next_month_start = today.replace(day=1) + timedelta(days=32) # Go to next month
            gstr3b_due_date = next_month_start.replace(day=20)

            # GSTR-1 due date (typically 11th of next month for monthly filers)
            gstr1_due_date = next_month_start.replace(day=11)

            st.write(f"- GSTR-1 for {next_month_start.strftime('%B %Y')}: **{gstr1_due_date.strftime('%d %B %Y')}**")
            days_left_gstr1 = (gstr1_due_date - today).days
            if days_left_gstr1 >= 0: st.warning(f"⏰ {days_left_gstr1} days remaining")
            else: st.error(f"❌ Overdue by {abs(days_left_gstr1)} days")

            st.write(f"- GSTR-3B for {next_month_start.strftime('%B %Y')}: **{gstr3b_due_date.strftime('%d %B %Y')}**")
            days_left_gstr3b = (gstr3b_due_date - today).days
            if days_left_gstr3b >= 0: st.warning(f"⏰ {days_left_gstr3b} days remaining")
            else: st.error(f"❌ Overdue by {abs(days_left_gstr3b)} days")


    with tab3:
        st.subheader("Filing Status Tracker (Simulated)")
        st.info("This is a simulated tracker. Update statuses manually.")
        
        # Simulated filing tracker (can be stored in session state or a simple DB table if needed)
        if 'simulated_filings' not in st.session_state:
            st.session_state.simulated_filings = [
                {
                    "id": 1,
                    "form": "GSTR-1",
                    "period": "July 2023",
                    "status": "Filed",
                    "date": (datetime.now() - timedelta(days=10)).date(),
                    "due_date": (datetime.now() - timedelta(days=5)).date()
                },
                {
                    "id": 2,
                    "form": "GSTR-3B",
                    "period": "July 2023",
                    "status": "Filed",
                    "date": (datetime.now() - timedelta(days=5)).date(),
                    "due_date": (datetime.now() - timedelta(days=2)).date()
                },
                {
                    "id": 3,
                    "form": "GSTR-1",
                    "period": "August 2023",
                    "status": "Pending",
                    "date": None,
                    "due_date": (datetime.now() + timedelta(days=5)).date()
                }
            ]
        
        st.write("#### Your Filings")
        filing_statuses = ["Pending", "Filed", "Revised", "Extension"]

        for filing in st.session_state.simulated_filings:
            days_left = (filing['due_date'] - datetime.now().date()).days if filing['due_date'] else None
            deadline_status = ""
            if days_left is not None:
                if days_left > 0: deadline_status = f"⏰ {days_left} days left"
                elif days_left == 0: deadline_status = "⚠️ Due TODAY!"
                else: deadline_status = f"❌ Overdue by {abs(days_left)} days"
            else:
                deadline_status = "No due date"

            with st.expander(f"{filing['form']} - {filing['period']} ({filing['status']}) - Due: {filing['due_date'] or 'N/A'} {deadline_status}"):
                st.write(f"**Status:** {filing['status']}")
                if filing['date']:
                    st.write(f"**Filed On:** {filing['date']}")
                
                # Allow updating status
                new_status = st.selectbox(
                    "Update Status",
                    filing_statuses,
                    index=filing_statuses.index(filing['status']),
                    key=f"update_filing_status_{filing['id']}"
                )
                if new_status != filing['status']:
                    filing['status'] = new_status
                    if new_status == "Filed":
                        filing['date'] = datetime.now().date() # Auto-set filed date
                    st.success(f"Status for {filing['form']} - {filing['period']} updated to {new_status}!")
                    st.rerun() # Rerun to update display

                if st.button("Simulate Filing Portal Link", key=f"file_{filing['id']}"):
                    st.info("This would typically open a link to the official tax filing portal.")
    
    cur.close(); conn.close()

# IPO & Cap Table Module
def ipo_module(business_id, ai_models):
    """Streamlit module for IPO and Cap Table management."""
    st.header("📊 IPO & Cap Table Management")
    st.info("This module provides tools for tracking IPO data and managing a simulated Cap Table.")
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    tab1, tab2, tab3 = st.tabs(["IPO Tracker", "Cap Table (Simulated)", "Investor Relations (Simulated)"])
    
    with tab1:
        st.subheader("IPO Tracker")
        
        cur.execute("SELECT id, company_name, issue_size, price_range, open_date, close_date, status, allotment_date, listing_date FROM ipo_data WHERE business_id = %s ORDER BY open_date DESC", (business_id,))
        ipos = cur.fetchall()
        
        if ipos:
            df = pd.DataFrame(ipos, columns=[
                "ID", "Company", "Issue Size", "Price Range", 
                "Open Date", "Close Date", "Status", "Allotment Date", "Listing Date"
            ])
            st.dataframe(df.drop(columns=["ID"]), hide_index=True)
            
            # Filter by status
            all_statuses = df["Status"].unique().tolist() if not df.empty else []
            status_filter = st.multiselect(
                "Filter by Status",
                options=all_statuses,
                default=all_statuses,
                key="ipo_status_filter"
            )
            
            if status_filter:
                filtered_df = df[df["Status"].isin(status_filter)]
                st.dataframe(filtered_df.drop(columns=["ID"]), hide_index=True)
            else:
                st.info("Select statuses to filter IPOs.")

        else:
            st.info("No IPO data added yet.")
        
        # Add new IPO
        with st.expander("Add IPO Details"):
            with st.form("add_ipo"):
                company_name = st.text_input("Company Name")
                issue_size = st.number_input("Issue Size (₹)", min_value=0.0, step=1000000.0)
                price_range = st.text_input("Price Range (₹ e.g., 100-110)")
                open_date = st.date_input("Open Date")
                close_date = st.date_input("Close Date")
                status = st.selectbox("Status", [
                    "Upcoming", "Open", "Closed", "Allotted", "Listed", "Cancelled"
                ])
                allotment_date = st.date_input("Allotment Date (Optional)")
                listing_date = st.date_input("Listing Date (Optional)")
                
                if st.form_submit_button("Add IPO"):
                    if not company_name or issue_size is None or not price_range or not open_date or not close_date or not status:
                         st.error("Please fill in required fields (Company Name, Issue Size, Price Range, Open/Close Dates, Status).")
                    elif open_date > close_date:
                         st.error("Close Date cannot be before Open Date.")
                    elif allotment_date and allotment_date < close_date:
                         st.error("Allotment Date cannot be before Close Date.")
                    elif listing_date and listing_date < (allotment_date or close_date):
                         st.error("Listing Date cannot be before Allotment Date or Close Date.")
                    else:
                        try:
                            cur.execute(
                                """INSERT INTO ipo_data 
                                (business_id, company_name, issue_size, price_range, open_date, close_date, status, allotment_date, listing_date) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                (business_id, company_name, issue_size, price_range, open_date, close_date, status, allotment_date, listing_date)
                            )
                            conn.commit()
                            st.success("IPO details added successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error adding IPO details: {e}")
    
    with tab2:
        st.subheader("Cap Table Management (Simulated)")
        st.info("This is a simulated Capitalization Table. It does not reflect real equity ownership.")
        
        # Simulated cap table data (can be stored in session state or a simple DB table)
        if 'simulated_cap_table' not in st.session_state:
            st.session_state.simulated_cap_table = [
                {"id": 1, "name": "Founders", "shares": 5000000, "percentage": 50.0, "type": "Common"},
                {"id": 2, "name": "Seed Investors", "shares": 2000000, "percentage": 20.0, "type": "Preferred"},
                {"id": 3, "name": "Series A Investors", "shares": 2000000, "percentage": 20.0, "type": "Preferred"},
                {"id": 4, "name": "Employee Pool", "shares": 1000000, "percentage": 10.0, "type": "Options"}
            ]
        
        df_cap = pd.DataFrame(st.session_state.simulated_cap_table)
        
        if not df_cap.empty:
            # Visualization
            st.write("### Cap Table Ownership Distribution")
            fig = px.pie(
                df_cap, 
                values="percentage", 
                names="name", 
                title="Simulated Cap Table Ownership"
            )
            st.plotly_chart(fig, use_container_width=True)
            
            # Detailed view
            st.write("### Detailed Cap Table")
            st.dataframe(df_cap.drop(columns=["id"]), hide_index=True)

            # Add/Edit Stakeholder (Simulated)
            with st.expander("Add/Edit Stakeholder (Simulated)"):
                 with st.form("add_edit_stakeholder_sim"):
                      stakeholder_id_edit = st.number_input("Enter ID to Edit (leave 0 for New)", min_value=0, step=1, key="stakeholder_id_edit")
                      
                      current_stakeholder = None
                      if stakeholder_id_edit > 0:
                           current_stakeholder_list = [s for s in st.session_state.simulated_cap_table if s['id'] == stakeholder_id_edit]
                           if current_stakeholder_list:
                                current_stakeholder = current_stakeholder_list[0]
                                st.info(f"Editing Stakeholder ID: {stakeholder_id_edit}")
                           else:
                                st.warning(f"Stakeholder with ID {stakeholder_id_edit} not found. Adding new.")
                                stakeholder_id_edit = 0 # Reset to add new

                      name = st.text_input("Name", value=current_stakeholder.get('name', '') if current_stakeholder else '', key="stakeholder_name")
                      shares = st.number_input("Shares", min_value=0, step=1000, value=current_stakeholder.get('shares', 0) if current_stakeholder else 0, key="stakeholder_shares")
                      percentage = st.number_input("Percentage (%)", min_value=0.0, max_value=100.0, step=0.1, value=current_stakeholder.get('percentage', 0.0) if current_stakeholder else 0.0, key="stakeholder_percentage")
                      stakeholder_type = st.selectbox("Type", ["Common", "Preferred", "Options", "Other"], index=["Common", "Preferred", "Options", "Other"].index(current_stakeholder.get('type', 'Common')) if current_stakeholder else 0, key="stakeholder_type")

                      if st.form_submit_button("Save Stakeholder (Simulated)"):
                           if not name or shares is None or percentage is None:
                                st.error("Name, Shares, and Percentage are required.")
                           else:
                                if stakeholder_id_edit > 0 and current_stakeholder:
                                     # Update existing
                                     current_stakeholder['name'] = name
                                     current_stakeholder['shares'] = shares
                                     current_stakeholder['percentage'] = percentage
                                     current_stakeholder['type'] = stakeholder_type
                                     st.success(f"Stakeholder ID {stakeholder_id_edit} updated!")
                                else:
                                     # Add new
                                     new_id = max([s['id'] for s in st.session_state.simulated_cap_table]) + 1 if st.session_state.simulated_cap_table else 1
                                     st.session_state.simulated_cap_table.append({
                                          "id": new_id, "name": name, "shares": shares, "percentage": percentage, "type": stakeholder_type
                                     })
                                     st.success(f"New stakeholder '{name}' added!")
                                st.rerun() # Rerun to update the table and chart

            # Waterfall analysis (Simulated)
            st.write("### Waterfall Analysis (Pre-IPO - Simulated)")
            st.info("Simulated valuation growth stages.")
            waterfall_data = [
                {"stage": "Pre-Seed", "valuation": 5000000},
                {"stage": "Seed", "valuation": 20000000},
                {"stage": "Series A", "valuation": 50000000},
                {"stage": "Series B", "valuation": 120000000},
                {"stage": "IPO Projection", "valuation": 500000000}
            ]
            
            fig = px.funnel(
                pd.DataFrame(waterfall_data), 
                x="valuation", 
                y="stage", 
                title="Simulated Valuation Growth"
            )
            st.plotly_chart(fig, use_container_width=True)

        else:
            st.info("No stakeholders in the simulated cap table yet. Add one above.")
    
    with tab3:
        st.subheader("Investor Relations (Simulated)")
        st.info("This is a simulated tracker for investor communications.")
        
        # Simulated investor communications (can be stored in session state or a simple DB table)
        if 'simulated_investor_comms' not in st.session_state:
            st.session_state.simulated_investor_comms = [
                {
                    "id": 1,
                    "date": (datetime.now() - timedelta(days=30)).date(),
                    "type": "Quarterly Report",
                    "recipients": "All Investors",
                    "status": "Sent"
                },
                {
                    "id": 2,
                    "date": (datetime.now() - timedelta(days=15)).date(),
                    "type": "Board Meeting",
                    "recipients": "Board Members",
                    "status": "Completed"
                },
                {
                    "id": 3,
                    "date": (datetime.now() + timedelta(days=10)).date(),
                    "type": "Roadshow",
                    "recipients": "Institutional Investors",
                    "status": "Scheduled"
                }
            ]

        st.write("### Recent Communications")
        if st.session_state.simulated_investor_comms:
            df_comms = pd.DataFrame(st.session_state.simulated_investor_comms)
            st.dataframe(df_comms.drop(columns=["id"]), hide_index=True)

            st.write("#### Communication Details")
            for comm in st.session_state.simulated_investor_comms:
                days_left = (comm['date'] - datetime.now().date()).days if comm['date'] else None
                date_status = ""
                if days_left is not None:
                    if days_left > 0: date_status = f"({days_left} days away)"
                    elif days_left == 0: date_status = "(Today)"
                    else: date_status = f"({abs(days_left)} days ago)"

                with st.expander(f"{comm['date']}: {comm['type']} - {comm['status']} {date_status}"):
                    st.write(f"**Recipients:** {comm['recipients']}")
                    st.write(f"**Status:** {comm['status']}")
                    # Add notes if available (assuming notes were added in the form)
                    if comm.get('notes'): st.write(f"**Notes:** {comm['notes']}")

                    # Allow updating status (Simulated)
                    # Ensure all statuses from initial data are in this list
                    comm_statuses = ["Scheduled", "Completed", "Cancelled", "Postponed", "Sent"] # Added "Sent"
                    
                    current_status_index = 0 # Default to first option
                    try:
                        current_status_index = comm_statuses.index(comm['status'])
                    except ValueError:
                        st.warning(f"Status '{comm['status']}' not in predefined list for communication ID {comm['id']}. Defaulting selection.")
                        # Optionally, add the unknown status to the list dynamically if desired, or log this.
                        # For now, it will default to "Scheduled" if the status is unknown.

                    new_status = st.selectbox(
                        "Update Status",
                        comm_statuses,
                        index=current_status_index, # Use the found or default index
                        key=f"update_comm_status_{comm['id']}"
                    )
                    if new_status != comm['status']:
                        comm['status'] = new_status
                        st.success(f"Status for {comm['type']} on {comm['date']} updated to {new_status}!")
                        st.rerun()

        else:
            st.info("No simulated investor communications recorded yet.")

        # Schedule New Communication (Simulated)
        with st.expander("Schedule New Communication (Simulated)"):
            with st.form("new_communication_sim"):
                comm_type = st.selectbox("Type", [
                    "Investor Update", "Board Meeting", "Roadshow", "Earnings Call", "Other"
                ])
                recipients = st.text_input("Recipients (e.g., 'All Investors', 'Board Members', 'Specific Investor Name')")
                scheduled_date = st.date_input("Date")
                notes = st.text_area("Notes (Optional)")
                
                if st.form_submit_button("Schedule (Simulated)"):
                    if not comm_type or not recipients or not scheduled_date:
                        st.error("Type, Recipients, and Date are required.")
                    else:
                        # Add to simulated list
                        new_id = max([c['id'] for c in st.session_state.simulated_investor_comms]) + 1 if st.session_state.simulated_investor_comms else 1
                        st.session_state.simulated_investor_comms.append({
                            "id": new_id,
                            "date": scheduled_date,
                            "type": comm_type,
                            "recipients": recipients,
                            "status": "Scheduled", # Default status for new
                            "notes": notes
                        })
                        st.success("Simulated communication scheduled successfully!")
                        st.rerun() # Rerun to update the list

    cur.close(); conn.close()

# Legal Marketplace Module (Business view of their listed providers)
def legal_marketplace(business_id, ai_models):
    """Streamlit module for Business users to manage Service Providers."""
    st.header("⚖️ Legal, CA & Insurance Marketplace (Your Listed Contacts)")
    st.info("Manage your contacts for Legal, CA, and Insurance services. You can also find new providers (simulated search).")
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    tab1, tab2, tab3, tab4 = st.tabs(["Legal Providers", "Chartered Accountants", "Insurance Providers", "Find Providers (Simulated)"])
    
    with tab1:
        st.subheader("Legal Service Providers")
        # Fetch legal providers linked to THIS business
        cur.execute("SELECT id, name, contact_email, rating, experience_years, pricing, availability, profile_description, website_url, linkedin_profile FROM service_providers WHERE business_id = %s AND service_type = 'legal' ORDER BY name", (business_id,))
        providers = cur.fetchall()
        
        if providers:
            df = pd.DataFrame(providers, columns=[
                "ID", "Name", "Contact Email", "Rating", 
                "Experience (Years)", "Pricing", "Available", "Profile Description", "Website", "LinkedIn"
            ])
            st.dataframe(df.drop(columns=["ID"]), hide_index=True)
        else:
            st.info("No legal service providers linked to your business yet.")
        
        # Add new legal provider linked to this business
        with st.expander("Add New Legal Provider (Link to Your Business)"):
            with st.form("add_legal_provider_to_biz"):
                name = st.text_input("Provider Name")
                email = st.text_input("Login Email (Optional - if they have a GrowBis account)") # Optional login email
                contact_email = st.text_input("Public Contact Email")
                rating = st.slider("Your Rating (1-5)", 1.0, 5.0, 4.0, step=0.1)
                experience = st.number_input("Years of Experience", min_value=0, step=1)
                pricing = st.text_input("Pricing (e.g., '$150/hr', 'Project based')")
                availability = st.checkbox("Currently Available", value=True)
                profile_description = st.text_area("Profile Description (Optional)")
                specializations_str = st.text_input("Specializations (comma-separated, Optional)")
                website_url = st.text_input("Website URL (Optional)")
                linkedin_profile = st.text_input("LinkedIn Profile URL (Optional)")

                if st.form_submit_button("Add Legal Provider"):
                    if not name or not contact_email:
                        st.error("Name and Public Contact Email are required.")
                    else:
                        specializations_list = [s.strip() for s in specializations_str.split(',') if s.strip()]
                        try:
                            # Check if a service provider with this login email already exists
                            existing_provider_id = None
                            if email:
                                cur.execute("SELECT id FROM service_providers WHERE email = %s", (email,))
                                existing_provider = cur.fetchone()
                                if existing_provider:
                                    existing_provider_id = existing_provider[0]

                            if existing_provider_id:
                                st.warning(f"A service provider with this login email ({email}) already exists (ID: {existing_provider_id}). Linking this existing provider to your business.")
                                # Update the existing provider to link them to this business
                                cur.execute("UPDATE service_providers SET business_id = %s WHERE id = %s", (business_id, existing_provider_id))
                                conn.commit()
                                st.success("Existing service provider linked to your business!")
                            else:
                                # Insert a new service provider linked to this business
                                cur.execute(
                                    """INSERT INTO service_providers 
                                    (business_id, name, service_type, email, password_hash, contact_email, rating, experience_years, pricing, availability, profile_description, specializations, website_url, linkedin_profile) 
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                    (business_id, name, 'legal', email if email else None, None, contact_email, rating, experience, pricing, availability, profile_description, specializations_list, website_url, linkedin_profile)
                                )
                                conn.commit()
                                st.success("New legal provider added and linked to your business!")
                            st.rerun()
                        except psycopg2.IntegrityError:
                             st.error("A service provider with this login email already exists and cannot be linked.")
                        except Exception as e:
                            st.error(f"Error adding/linking legal provider: {e}")

    with tab2:
        st.subheader("Chartered Accountants")
        # Fetch CA providers linked to THIS business
        cur.execute("SELECT id, name, contact_email, rating, experience_years, pricing, availability, profile_description, website_url, linkedin_profile FROM service_providers WHERE business_id = %s AND service_type = 'ca' ORDER BY name", (business_id,))
        providers = cur.fetchall()
        
        if providers:
            df = pd.DataFrame(providers, columns=[
                "ID", "Name", "Contact Email", "Rating", 
                "Experience (Years)", "Pricing", "Available", "Profile Description", "Website", "LinkedIn"
            ])
            st.dataframe(df.drop(columns=["ID"]), hide_index=True)
        else:
            st.info("No Chartered Accountants linked to your business yet.")
        
        # Add new CA linked to this business
        with st.expander("Add New Chartered Accountant (Link to Your Business)"):
            with st.form("add_ca_provider_to_biz"):
                name = st.text_input("CA Name")
                email = st.text_input("Login Email (Optional - if they have a GrowBis account)", key="ca_login_email")
                contact_email = st.text_input("Public Contact Email", key="ca_contact_email")
                rating = st.slider("Your Rating (1-5)", 1.0, 5.0, 4.0, step=0.1, key="ca_rating")
                experience = st.number_input("Years of Experience", min_value=0, step=1, key="ca_experience")
                pricing = st.text_input("Pricing", key="ca_pricing")
                availability = st.checkbox("Currently Available", value=True, key="ca_availability")
                profile_description = st.text_area("Profile Description (Optional)", key="ca_profile_desc")
                specializations_str = st.text_input("Specializations (comma-separated, Optional)", key="ca_specializations")
                website_url = st.text_input("Website URL (Optional)", key="ca_website")
                linkedin_profile = st.text_input("LinkedIn Profile URL (Optional)", key="ca_linkedin")

                if st.form_submit_button("Add Chartered Accountant"):
                    if not name or not contact_email:
                        st.error("Name and Public Contact Email are required.")
                    else:
                        specializations_list = [s.strip() for s in specializations_str.split(',') if s.strip()]
                        try:
                            existing_provider_id = None
                            if email:
                                cur.execute("SELECT id FROM service_providers WHERE email = %s", (email,))
                                existing_provider = cur.fetchone()
                                if existing_provider:
                                    existing_provider_id = existing_provider[0]

                            if existing_provider_id:
                                st.warning(f"A service provider with this login email ({email}) already exists (ID: {existing_provider_id}). Linking this existing provider to your business.")
                                cur.execute("UPDATE service_providers SET business_id = %s WHERE id = %s", (business_id, existing_provider_id))
                                conn.commit()
                                st.success("Existing service provider linked to your business!")
                            else:
                                cur.execute(
                                    """INSERT INTO service_providers 
                                    (business_id, name, service_type, email, password_hash, contact_email, rating, experience_years, pricing, availability, profile_description, specializations, website_url, linkedin_profile) 
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                    (business_id, name, 'ca', email if email else None, None, contact_email, rating, experience, pricing, availability, profile_description, specializations_list, website_url, linkedin_profile)
                                )
                                conn.commit()
                                st.success("New CA added and linked to your business!")
                            st.rerun()
                        except psycopg2.IntegrityError:
                             st.error("A service provider with this login email already exists and cannot be linked.")
                        except Exception as e:
                            st.error(f"Error adding/linking CA: {e}")

    with tab3:
        st.subheader("Insurance Providers")
        # Fetch Insurance providers linked to THIS business
        cur.execute("SELECT id, name, contact_email, rating, experience_years, pricing, availability, profile_description, website_url, linkedin_profile FROM service_providers WHERE business_id = %s AND service_type = 'insurance' ORDER BY name", (business_id,))
        providers = cur.fetchall()
        
        if providers:
            df = pd.DataFrame(providers, columns=[
                "ID", "Name", "Contact Email", "Rating", 
                "Experience (Years)", "Pricing", "Available", "Profile Description", "Website", "LinkedIn"
            ])
            st.dataframe(df.drop(columns=["ID"]), hide_index=True)
        else:
            st.info("No insurance providers linked to your business yet.")
        
        # Add new Insurance provider linked to this business
        with st.expander("Add New Insurance Provider (Link to Your Business)"):
            with st.form("add_insurance_provider_to_biz"):
                name = st.text_input("Provider Name")
                email = st.text_input("Login Email (Optional - if they have a GrowBis account)", key="ins_login_email")
                contact_email = st.text_input("Public Contact Email", key="ins_contact_email")
                rating = st.slider("Your Rating (1-5)", 1.0, 5.0, 4.0, step=0.1, key="ins_rating")
                experience = st.number_input("Years of Experience", min_value=0, step=1, key="ins_experience")
                pricing = st.text_input("Pricing", key="ins_pricing")
                availability = st.checkbox("Currently Available", value=True, key="ins_availability")
                profile_description = st.text_area("Profile Description (Optional)", key="ins_profile_desc")
                specializations_str = st.text_input("Specializations (comma-separated, Optional)", key="ins_specializations")
                website_url = st.text_input("Website URL (Optional)", key="ins_website")
                linkedin_profile = st.text_input("LinkedIn Profile URL (Optional)", key="ins_linkedin")

                if st.form_submit_button("Add Insurance Provider"):
                    if not name or not contact_email:
                        st.error("Name and Public Contact Email are required.")
                    else:
                        specializations_list = [s.strip() for s in specializations_str.split(',') if s.strip()]
                        try:
                            existing_provider_id = None
                            if email:
                                cur.execute("SELECT id FROM service_providers WHERE email = %s", (email,))
                                existing_provider = cur.fetchone()
                                if existing_provider:
                                    existing_provider_id = existing_provider[0]

                            if existing_provider_id:
                                st.warning(f"A service provider with this login email ({email}) already exists (ID: {existing_provider_id}). Linking this existing provider to your business.")
                                cur.execute("UPDATE service_providers SET business_id = %s WHERE id = %s", (business_id, existing_provider_id))
                                conn.commit()
                                st.success("Existing service provider linked to your business!")
                            else:
                                cur.execute(
                                    """INSERT INTO service_providers 
                                    (business_id, name, service_type, email, password_hash, contact_email, rating, experience_years, pricing, availability, profile_description, specializations, website_url, linkedin_profile) 
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                    (business_id, name, 'insurance', email if email else None, None, contact_email, rating, experience, pricing, availability, profile_description, specializations_list, website_url, linkedin_profile)
                                )
                                conn.commit()
                                st.success("New Insurance provider added and linked to your business!")
                            st.rerun()
                        except psycopg2.IntegrityError:
                             st.error("A service provider with this login email already exists and cannot be linked.")
                        except Exception as e:
                            st.error(f"Error adding/linking Insurance provider: {e}")

    with tab4:
        st.subheader("Find Service Providers (Simulated Search)")
        st.info("This simulates searching for service providers on the GrowBis platform.")
        
        service_type_search = st.selectbox("Service Needed", [
            "Legal", "Chartered Accountant", "Insurance", "Consulting", "Other"
        ], key="service_search_type")
        min_rating_search = st.slider("Minimum Rating", 1.0, 5.0, 4.0, step=0.1, key="service_search_rating")
        # max_price_search = st.text_input("Maximum Budget (leave blank for any)", key="service_search_price") # Price search is complex without structured data

        if st.button("Search Providers (Simulated)", key="search_providers_btn_sim"):
            st.write(f"#### Simulated Results for {service_type_search} Providers with Rating ≥ {min_rating_search}★")
            
            # Simulate fetching providers based on type and rating
            # In a real app, this would query the service_providers table (potentially across all businesses or a public listing)
            # For simulation, we'll generate some fake results.
            simulated_results = []
            num_results = np.random.randint(2, 6)
            for i in range(num_results):
                 simulated_results.append({
                      "Name": f"{service_type_search} Expert {i+1}",
                      "Rating": round(np.random.uniform(min_rating_search, 5.0), 1),
                      "Experience (Years)": np.random.randint(5, 20),
                      "Pricing (Est.)": "$100-$300/hr" if service_type_search == "Legal" else "₹5,000-₹15,000/month",
                      "Availability": np.random.choice(["Available", "Limited", "Busy"]),
                      "Profile Snippet": f"Experienced {service_type_search} provider.",
                      "Contact Info": "contact@example.com" # Simulated contact
                 })

            if simulated_results:
                df_results = pd.DataFrame(simulated_results)
                st.dataframe(df_results, hide_index=True)

                st.write("#### Contact Simulated Provider")
                selected_sim_provider_name = st.selectbox("Select Provider to Contact", ["-- Select --"] + df_results["Name"].tolist(), key="select_sim_provider")
                if selected_sim_provider_name != "-- Select --":
                     provider_contact_info = next(item['Contact Info'] for item in simulated_results if item['Name'] == selected_sim_provider_name)
                     st.info(f"Simulated Contact Info for {selected_sim_provider_name}: {provider_contact_info}")
                     # In a real app, this could initiate a chat via the messaging module or show a contact form.
                     if st.button(f"Initiate Chat with {selected_sim_provider_name} (Simulated)", key=f"chat_sim_provider_{selected_sim_provider_name}"):
                          st.warning("Chat initiation with simulated providers is not implemented.")

            else:
                st.info("No simulated providers found matching your criteria.")

    cur.close(); conn.close()
def get_financial_performance_data(business_id, num_periods=12, period_type='M'):
    """
    Fetches financial performance data (Revenue, Expenses, Profit) for the last num_periods.
    period_type can be 'M' for Monthly or 'Q' for Quarterly.
    """
    conn = get_db_connection()
    cur = conn.cursor()
    
    end_date = datetime.now().date()
    
    if period_type == 'M':
        # Use Month Start frequency for easier grouping with DATE_TRUNC
        dates = pd.date_range(end=end_date.replace(day=1), periods=num_periods, freq='MS')
        period_label_format = '%Y-%m'
        sql_trunc_period = 'month'
    elif period_type == 'Q':
        # Use Quarter Start frequency
        current_quarter_start = pd.Timestamp(end_date).to_period('Q').start_time.date()
        dates = pd.date_range(end=current_quarter_start, periods=num_periods, freq='QS')
        # For label, we'll calculate quarter number manually
        sql_trunc_period = 'quarter'
    else: # Default to monthly
        dates = pd.date_range(end=end_date.replace(day=1), periods=num_periods, freq='MS')
        period_label_format = '%Y-%m'
        sql_trunc_period = 'month'

    df_data = pd.DataFrame({'PeriodStart': dates})
    
    if period_type == 'Q':
        df_data['PeriodLabel'] = df_data['PeriodStart'].apply(lambda x: f"{x.year}-Q{((x.month-1)//3)+1}")
    else:
        df_data['PeriodLabel'] = df_data['PeriodStart'].dt.strftime(period_label_format)

    # Fetch Revenue
    try:
        query_revenue = sql.SQL("""
            SELECT DATE_TRUNC(%s, issue_date) as period_start_db, SUM(total_amount) as revenue
            FROM invoices
            WHERE business_id = %s AND issue_date >= %s AND issue_date <= %s
            GROUP BY period_start_db
            ORDER BY period_start_db;
        """)
        # Ensure dates.min() is a date object
        min_date_for_query = dates.min().date() if hasattr(dates.min(), 'date') else dates.min()
        cur.execute(query_revenue, (sql_trunc_period, business_id, min_date_for_query, end_date))
        revenue_data = cur.fetchall()
        df_revenue = pd.DataFrame(revenue_data, columns=['PeriodStart', 'Revenue'])
        if not df_revenue.empty:
            df_revenue['PeriodStart'] = pd.to_datetime(df_revenue['PeriodStart'])
            df_data = pd.merge(df_data, df_revenue, on='PeriodStart', how='left')
        else:
            df_data['Revenue'] = 0.0
    except Exception as e:
        st.error(f"Error fetching revenue data: {e}")
        df_data['Revenue'] = 0.0
    df_data['Revenue'] = df_data['Revenue'].fillna(0.0).astype(float)

    # Fetch Expenses (Simplified: using total monthly salaries)
    # For quarterly, this will be 3x monthly salary.
    try:
        cur.execute("SELECT SUM(salary) FROM employees WHERE business_id = %s", (business_id,))
        total_monthly_salary_tuple = cur.fetchone()
        total_monthly_salary = float(total_monthly_salary_tuple[0]) if total_monthly_salary_tuple and total_monthly_salary_tuple[0] is not None else 0.0
        
        expense_per_period = total_monthly_salary * 3 if period_type == 'Q' else total_monthly_salary
        df_data['Expenses'] = expense_per_period
    except Exception as e:
        st.error(f"Error fetching salary data for expenses: {e}")
        df_data['Expenses'] = 0.0
    df_data['Expenses'] = df_data['Expenses'].fillna(0.0).astype(float)
    
    df_data['Profit'] = df_data['Revenue'] - df_data['Expenses']
    
    cur.close()
    conn.close()
    return df_data[['PeriodLabel', 'Revenue', 'Expenses', 'Profit']].rename(columns={'PeriodLabel': 'Period'})
def calculate_delta(current, previous):
    if previous == 0: return "N/A" if current == 0 else "New"
    return f"{(current - previous) / previous * 100:.1f}%"
def get_inventory_turnover_data(business_id, lookback_days=365):
    """Calculates Inventory Turnover."""
    conn = get_db_connection()
    cur = conn.cursor()
    
    cogs = 0.0
    current_inventory_value_at_cost = 0.0
    
    end_date = datetime.now().date()
    start_date_cogs = end_date - timedelta(days=lookback_days)

    try:
        # Fetch all products to get their prices (and estimate cost)
        cur.execute("SELECT id, price, quantity FROM products WHERE business_id = %s", (business_id,))
        products_details = cur.fetchall()
        
        product_info = {} # {product_id: {'price': selling_price, 'quantity': current_quantity}}
        for prod_id, price, quantity in products_details:
            product_info[prod_id] = {'price': float(price) if price else 0.0, 
                                     'quantity': int(quantity) if quantity else 0}

        # Calculate COGS
        cur.execute(
            """SELECT items FROM invoices 
               WHERE business_id = %s AND issue_date BETWEEN %s AND %s""",
            (business_id, start_date_cogs, end_date)
        )
        invoices_items_json = cur.fetchall()
        
        for items_json_tuple in invoices_items_json:
            if items_json_tuple and items_json_tuple[0]:
                try:
                    # items column is JSONB, psycopg2 might return it as dict/list already
                    if isinstance(items_json_tuple[0], str):
                        items_list = json.loads(items_json_tuple[0])
                    else:
                        items_list = items_json_tuple[0] # Assuming it's already a list of dicts
                        
                    for item_detail in items_list:
                        product_id = item_detail.get('product_id')
                        quantity_sold = item_detail.get('quantity', 0)
                        
                        if product_id in product_info:
                            selling_price = product_info[product_id]['price']
                            cost_price_estimate = selling_price * 0.6 # Assume 60% cost margin
                            cogs += quantity_sold * cost_price_estimate
                except json.JSONDecodeError as e:
                    st.warning(f"Could not decode items JSON: {e} - Data: {items_json_tuple[0][:100]}")
                except TypeError as e:
                    st.warning(f"Type error processing items: {e} - Data: {items_json_tuple[0]}")


        # Calculate Current Inventory Value at Cost
        for prod_id, info in product_info.items():
            cost_price_estimate = info['price'] * 0.6
            current_inventory_value_at_cost += info['quantity'] * cost_price_estimate
                
        if current_inventory_value_at_cost > 0:
            inventory_turnover_raw = cogs / current_inventory_value_at_cost
        else:
            inventory_turnover_raw = 0.0

        # Example target and trend logic (can be refined)
        target_turnover = 6.0 
        trend = "neutral" # Default
        if inventory_turnover_raw > target_turnover * 0.8: trend = "up"
        elif inventory_turnover_raw < target_turnover * 0.5: trend = "critical_low"
        else: trend = "down" # Could be neutral if within a range, but simplified here


    except Exception as e:
        st.error(f"Error calculating inventory turnover: {e}")
        inventory_turnover_raw = 0.0 
        target_turnover = 6.0
        trend = "error"
    finally:
        cur.close()
        conn.close()
        
    return {
        "name": "Inventory Turnover",
        "value": round(inventory_turnover_raw, 1),
        "target": target_turnover,
        "trend": trend 
    }
def get_customer_lifetime_value_data(business_id):
    """Calculates average Customer Lifetime Value (CLV)."""
    conn = get_db_connection()
    cur = conn.cursor()
    clv_avg = 0.0
    num_customers = 0
    try:
        # Using customer_email to identify unique customers and their total spending
        cur.execute(
            """
            SELECT customer_email, SUM(total_amount) as total_spent
            FROM invoices
            WHERE business_id = %s AND customer_email IS NOT NULL AND customer_email != ''
            GROUP BY customer_email;
            """, (business_id,)
        )
        customer_spending = cur.fetchall()
        
        if customer_spending:
            df_spending = pd.DataFrame(customer_spending, columns=['customer_email', 'total_spent'])
            df_spending['total_spent'] = df_spending['total_spent'].astype(float)
            clv_avg = df_spending['total_spent'].mean()
            num_customers = df_spending['customer_email'].nunique()
            
    except Exception as e:
        st.error(f"Error calculating CLV: {e}")
    finally:
        cur.close()
        conn.close()
    
    # Target CLV can be a business goal, placeholder for now
    target_clv = 1000  # Example target
    trend = "neutral" # Default
    if clv_avg > 0: # Only calculate trend if there's a CLV
        if clv_avg > target_clv * 0.8: trend = "up"
        elif clv_avg < target_clv * 0.5: trend = "critical_low" # More distinct from just "down"
        elif clv_avg < target_clv: trend = "down"
    
    return {
        "name": "Avg. Customer Lifetime Value (CLV)",
        "value": round(clv_avg, 0), # CLV is typically a currency value
        "target": target_clv,
        "trend": trend,
        "unit": "$", # Assuming currency is USD
        "note": f"Based on {num_customers} unique customers." if num_customers > 0 else "No customer data for CLV."
    }
def get_operational_efficiency_trends_data(business_id, num_months=6):
    """
    Fetches average project progress for the last num_months to show operational efficiency.
    """
    conn = get_db_connection()
    cur = conn.cursor()
    df_efficiency = pd.DataFrame(columns=["Month", "Efficiency"])
    
    try:
        # Calculate the start date for the lookback period
        end_date_for_query = datetime.now().date()
        # To get 'num_months' of data, we need to go back num_months-1 full months, 
        # and then to the start of that month.
        # Example: if num_months = 6 and today is 2023-08-15
        # We want data for Aug, Jul, Jun, May, Apr, Mar
        # So, start_date_lookback should be 2023-03-01
        
        # First day of the current month
        first_day_current_month = end_date_for_query.replace(day=1)
        
        # Go back (num_months - 1) full months
        start_date_lookback = first_day_current_month
        for _ in range(num_months - 1):
            # Go to the last day of the previous month, then the first day of that month
            last_day_prev_month = start_date_lookback - timedelta(days=1)
            start_date_lookback = last_day_prev_month.replace(day=1)

        query = sql.SQL("""
            SELECT 
                DATE_TRUNC('month', COALESCE(p.end_date, p.start_date)) as project_month, 
                AVG(p.progress) as avg_progress
            FROM projects p
            WHERE p.business_id = %s 
              AND p.status IN ('In Progress', 'Completed') -- Consider relevant statuses
              AND COALESCE(p.end_date, p.start_date) >= %s -- Projects active or ending in the period
              AND p.start_date <= %s -- Projects started before or during the period end
            GROUP BY project_month
            HAVING COUNT(p.id) > 0 -- Ensure there are projects in the month to average
            ORDER BY project_month DESC
            LIMIT %s; 
        """)
        # The LIMIT might be too restrictive if some months have no projects.
        # A better approach is to generate all months and left-join.
        # For simplicity with current structure, we'll proceed but note this.

        cur.execute(query, (business_id, start_date_lookback, end_date_for_query, num_months))
        data = cur.fetchall()
        
        if data:
            df_raw_efficiency = pd.DataFrame(data, columns=["Month", "Efficiency"])
            df_raw_efficiency["Month"] = pd.to_datetime(df_raw_efficiency["Month"])
            df_raw_efficiency["Efficiency"] = df_raw_efficiency["Efficiency"].astype(float).round(1)
            
            # Create a full date range for the last num_months to ensure all months are present
            month_series = pd.date_range(end=first_day_current_month, periods=num_months, freq='MS')
            df_template_months = pd.DataFrame({'Month': month_series})

            # Merge and fill missing values
            df_efficiency = pd.merge(df_template_months, df_raw_efficiency, on="Month", how="left")
            df_efficiency["Efficiency"] = df_efficiency["Efficiency"].fillna(0) # Fill months with no projects with 0 efficiency
            df_efficiency["Month"] = df_efficiency["Month"].dt.strftime('%Y-%m')
            df_efficiency = df_efficiency.sort_values(by="Month").reset_index(drop=True)

    except Exception as e:
        st.error(f"Error fetching operational efficiency trends: {e}")
    finally:
        cur.close()
        conn.close()
    return df_efficiency

def get_employee_productivity_score_data(business_id):
    """Calculates average Employee Productivity Score."""
    conn = get_db_connection()
    cur = conn.cursor()
    avg_score = 0.0
    num_employees_rated = 0
    try:
        cur.execute(
            "SELECT AVG(performance_score), COUNT(id) FROM employees WHERE business_id = %s AND performance_score IS NOT NULL",
            (business_id,)
        )
        result = cur.fetchone()
        if result and result[0] is not None:
            avg_score = float(result[0])
            num_employees_rated = int(result[1] if result[1] is not None else 0) # Ensure count is int
    except Exception as e:
        st.error(f"Error fetching employee productivity score: {e}")
    finally:
        cur.close()
        conn.close()

    target_score = 8.0 # Assuming score is out of 10
    trend = "neutral" # Default
    if num_employees_rated > 0: # Only calculate trend if there are rated employees
        if avg_score >= target_score * 0.9: trend = "up"
        elif avg_score < target_score * 0.7: trend = "down" 
        # Could add more nuanced trends like "critical_low" if avg_score is very low
    
    return {
        "name": "Avg. Employee Productivity Score",
        "value": round(avg_score, 1) if num_employees_rated > 0 else "N/A",
        "target": target_score,
        "trend": trend,
        "unit": "/10",
        "note": f"Based on {num_employees_rated} rated employees." if num_employees_rated > 0 else "No rated employees."
    }
def get_sales_performance_report_data(business_id, start_date, end_date):
    """Fetches data for the Sales Performance custom report."""
    conn = get_db_connection()
    cur = conn.cursor()
    
    report = {
        "total_revenue": 0.0,
        "top_product_name": "N/A",
        "top_product_revenue": 0.0,
        "num_invoices": 0,
        "avg_invoice_value": 0.0
    }

    try:
        # Total Revenue and Number of Invoices
        cur.execute(
            """SELECT SUM(total_amount), COUNT(id) FROM invoices
               WHERE business_id = %s AND issue_date BETWEEN %s AND %s""",
            (business_id, start_date, end_date)
        )
        result = cur.fetchone()
        if result:
            report["total_revenue"] = float(result[0]) if result[0] is not None else 0.0
            report["num_invoices"] = int(result[1]) if result[1] is not None else 0
        
        if report["num_invoices"] > 0:
            report["avg_invoice_value"] = report["total_revenue"] / report["num_invoices"]

        # Top Product
        cur.execute(
            """SELECT items FROM invoices
               WHERE business_id = %s AND issue_date BETWEEN %s AND %s""",
            (business_id, start_date, end_date)
        )
        all_items_json_tuples = cur.fetchall()
        
        product_sales = {} # {product_name: total_sales_value}
        for items_json_tuple in all_items_json_tuples:
            if items_json_tuple and items_json_tuple[0]:
                try:
                    if isinstance(items_json_tuple[0], str):
                        items_list = json.loads(items_json_tuple[0])
                    else:
                        items_list = items_json_tuple[0]

                    for item_detail in items_list:
                        name = item_detail.get("name")
                        total = item_detail.get("total", 0.0)
                        if name:
                            product_sales[name] = product_sales.get(name, 0.0) + float(total)
                except json.JSONDecodeError as e:
                    st.warning(f"Could not decode items JSON for top product: {e} - Data: {items_json_tuple[0][:100]}")
                except TypeError as e:
                    st.warning(f"Type error processing items for top product: {e} - Data: {items_json_tuple[0]}")

        if product_sales:
            # Sort products by sales value in descending order
            sorted_products = sorted(product_sales.items(), key=lambda x: x[1], reverse=True)
            if sorted_products:
                report["top_product_name"] = sorted_products[0][0]
                report["top_product_revenue"] = sorted_products[0][1]
                # You could extend this to top N products
                # report["top_products_list"] = sorted_products[:3] 

    except Exception as e:
        st.error(f"Error generating sales performance report data: {e}")
    finally:
        cur.close()
        conn.close()
    return report
def get_employee_productivity_report_data(business_id, start_date_filter, end_date_filter):
    """Fetches data for the Employee Productivity custom report."""
    # Note: start_date_filter and end_date_filter are for the report's conceptual period.
    # Employee performance_score is typically a current or recent snapshot.
    # If you want to filter employees based on join_date or last_appraisal_date within this period,
    # that logic would need to be added to the SQL query.
    conn = get_db_connection()
    cur = conn.cursor()
    
    report = {
        "avg_score_overall": 0.0,
        "top_performer_name": "N/A",
        "top_performer_score": 0.0,
        "department_scores": [], # List of dicts: {"department": "X", "avg_score": Y}
        "num_employees_total": 0,
        "num_employees_rated": 0,
    }

    try:
        # Get all employees for total count
        cur.execute(
            "SELECT COUNT(id) FROM employees WHERE business_id = %s",
            (business_id,)
        )
        total_emp_result = cur.fetchone()
        if total_emp_result:
            report["num_employees_total"] = int(total_emp_result[0])

        # Get rated employees for detailed stats
        cur.execute(
            """SELECT name, department, performance_score 
               FROM employees 
               WHERE business_id = %s AND performance_score IS NOT NULL""",
            (business_id,)
        )
        employees_data = cur.fetchall()
        report["num_employees_rated"] = len(employees_data)

        if employees_data:
            df_employees = pd.DataFrame(employees_data, columns=["name", "department", "performance_score"])
            df_employees["performance_score"] = pd.to_numeric(df_employees["performance_score"], errors='coerce')
            
            # Filter out any rows where performance_score became NaN after coercion (if any)
            rated_employees_df = df_employees.dropna(subset=['performance_score'])
            report["num_employees_rated"] = len(rated_employees_df) # Update count after potential dropna

            if not rated_employees_df.empty:
                report["avg_score_overall"] = round(rated_employees_df["performance_score"].mean(), 1)
                
                # Find top performer
                top_perf_idx = rated_employees_df["performance_score"].idxmax()
                top_performer = rated_employees_df.loc[top_perf_idx]
                report["top_performer_name"] = top_performer["name"]
                report["top_performer_score"] = top_performer["performance_score"]
                
                # Calculate average scores by department
                dept_scores_series = rated_employees_df.groupby("department")["performance_score"].mean().round(1)
                report["department_scores"] = [{"department": idx, "avg_score": val} for idx, val in dept_scores_series.items()]
            else: # Handle case where all performance_scores were invalid
                report["avg_score_overall"] = 0.0
                report["top_performer_name"] = "N/A"
                report["top_performer_score"] = 0.0
                report["department_scores"] = []


    except Exception as e:
        st.error(f"Error generating employee productivity report data: {e}")
    finally:
        cur.close()
        conn.close()
    return report
def get_inventory_analysis_report_data(business_id, start_date_filter, end_date_filter):
    """Fetches data for the Inventory Analysis custom report."""
    conn = get_db_connection()
    cur = conn.cursor()
    report = {
        "total_inventory_value_at_price": 0.0,
        "total_inventory_value_at_cost_est": 0.0,
        "num_products": 0,
        "slow_moving_items_count": 0,
        "top_slow_moving_items": [], # List of dicts {"name": X, "quantity": Y, "days_since_last_sale_est": Z}
        "inventory_turnover": 0.0 # Default, will be updated
    }
    cost_margin = 0.6 # Assumption for cost estimation

    try:
        # Get inventory turnover first
        turnover_data = get_inventory_turnover_data(business_id) # Assumes lookback_days default is fine for this context
        report["inventory_turnover"] = turnover_data.get('value', 0.0)

        # Fetch all products for current value
        cur.execute(
            "SELECT id, name, price, quantity FROM products WHERE business_id = %s",
            (business_id,)
        )
        products_data = cur.fetchall()
        report["num_products"] = len(products_data)

        if products_data:
            df_products = pd.DataFrame(products_data, columns=["id", "name", "price", "quantity"])
            df_products["price"] = pd.to_numeric(df_products["price"], errors='coerce').fillna(0)
            df_products["quantity"] = pd.to_numeric(df_products["quantity"], errors='coerce').fillna(0)
            
            df_products["value_at_price"] = df_products["price"] * df_products["quantity"]
            df_products["value_at_cost_est"] = df_products["value_at_price"] * cost_margin
            
            report["total_inventory_value_at_price"] = df_products["value_at_price"].sum()
            report["total_inventory_value_at_cost_est"] = df_products["value_at_cost_est"].sum()

            # Identify slow-moving items: items with high quantity not sold within the filtered period
            # This requires sales data from the invoices table within the report's date range.
            cur.execute(
                """SELECT items, issue_date FROM invoices 
                   WHERE business_id = %s AND issue_date BETWEEN %s AND %s""",
                (business_id, start_date_filter, end_date_filter)
            )
            invoices_items_json = cur.fetchall()
            
            product_sales_dates = {} # {product_id: last_sale_date_in_period}
            for items_json_tuple, issue_date_val in invoices_items_json:
                if items_json_tuple and items_json_tuple[0]:
                    try:
                        items_list = json.loads(items_json_tuple[0]) if isinstance(items_json_tuple[0], str) else items_json_tuple[0]
                        for item_detail in items_list:
                            product_id = item_detail.get('product_id')
                            if product_id:
                                # We only care about sales within the filtered period for this specific "slow-moving" definition
                                current_last_sale_in_period = product_sales_dates.get(product_id)
                                if current_last_sale_in_period is None or issue_date_val > current_last_sale_in_period:
                                    product_sales_dates[product_id] = issue_date_val
                    except (json.JSONDecodeError, TypeError) as e_json:
                        st.warning(f"Error processing invoice items for slow-moving check: {e_json}")
            
            today = datetime.now().date() # Or use end_date_filter as the "today" for consistency
            
            slow_items_list = []
            # Define "slow-moving" criteria: e.g., quantity > 20 and no sales in the report period
            quantity_threshold_slow = 20 
            
            for index, row in df_products.iterrows():
                # If a product_id is NOT in product_sales_dates, it means it wasn't sold in the report period.
                if row['quantity'] > quantity_threshold_slow and row['id'] not in product_sales_dates:
                    # To estimate days_since_last_sale, we'd need global last sale date, not just in period.
                    # For this report, "Never in period" is more accurate if it wasn't sold in start_date_filter to end_date_filter
                    slow_items_list.append({
                        "name": row["name"], 
                        "quantity": row["quantity"],
                        "days_since_last_sale_est": "Never in period" 
                    })
            
            report["slow_moving_items_count"] = len(slow_items_list)
            # Sort by quantity descending for "top" slow-moving
            report["top_slow_moving_items"] = sorted(slow_items_list, key=lambda x: x["quantity"], reverse=True)[:5]

    except Exception as e:
        st.error(f"Error generating inventory analysis report data: {e}")
        # Ensure all keys are present even on error
        report.setdefault("total_inventory_value_at_price", 0.0)
        report.setdefault("total_inventory_value_at_cost_est", 0.0)
        report.setdefault("num_products", 0)
        report.setdefault("slow_moving_items_count", 0)
        report.setdefault("top_slow_moving_items", [])
        report.setdefault("inventory_turnover", 0.0)
    finally:
        cur.close()
        conn.close()
    return report

# Enterprise Intelligence Module
def enterprise_intelligence(business_id, ai_models): 
    """Streamlit module for Enterprise Intelligence Dashboards."""
    st.header("📊 Enterprise Intelligence Dashboards")
    st.info("Provides key business metrics and reports based on your data.")
    
    tab1, tab2, tab3 = st.tabs([
        "Financial Performance", 
        "Operational Metrics", 
        "Custom Reports"
    ])
    
    with tab1:
        st.subheader("Financial Performance")
        
        fin_period_type = st.selectbox("View Financials By:", ["Monthly", "Quarterly"], key="fin_period_type_ei")
        num_fin_periods = st.slider("Number of Past Periods:", min_value=3, max_value=24, 
                                    value=12 if fin_period_type == "Monthly" else 8, key="num_fin_periods_ei")

        with st.spinner("Fetching financial data..."):
            df_finance = get_financial_performance_data(business_id, num_periods=num_fin_periods, period_type=fin_period_type[0])
        
        if not df_finance.empty:
            fig_finance = px.line(
                df_finance, 
                x="Period", 
                y=["Revenue", "Expenses", "Profit"],
                title=f"{fin_period_type} Financial Performance ({num_fin_periods} Periods)"
            )
            st.plotly_chart(fig_finance, use_container_width=True)
            
            st.write(f"### Key Financial Metrics (Latest {fin_period_type[:-2]})")
            latest_period_data = df_finance.iloc[-1]
            revenue_latest = latest_period_data['Revenue']
            expenses_latest = latest_period_data['Expenses']
            profit_latest = latest_period_data['Profit']
            
            revenue_prev, expenses_prev, profit_prev = 0.0, 0.0, 0.0
            delta_label = f"vs Prev. {fin_period_type[:-2]}"
            if len(df_finance) >= 2:
                prev_period_data = df_finance.iloc[-2]
                revenue_prev = prev_period_data['Revenue']
                expenses_prev = prev_period_data['Expenses']
                profit_prev = prev_period_data['Profit']
            else:
                delta_label = "(No Previous Period Data)"

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric(f"Total Revenue", f"${revenue_latest:,.0f}", 
                         f"{calculate_delta(revenue_latest, revenue_prev)} {delta_label}")
            with col2:
                st.metric(f"Total Expenses (Est.)", f"${expenses_latest:,.0f}", 
                         f"{calculate_delta(expenses_latest, expenses_prev)} {delta_label}")
                st.caption("Expenses estimated based on total salaries.")
            with col3:
                st.metric(f"Net Profit (Est.)", f"${profit_latest:,.0f}", 
                         f"{calculate_delta(profit_latest, profit_prev)} {delta_label}")
        else:
            st.info("No financial data available to display for the selected period.")
    
    with tab2:
        st.subheader("Operational Metrics")
        
        operational_metrics_data = []
        with st.spinner("Fetching operational metrics..."):
            operational_metrics_data.append(get_inventory_turnover_data(business_id))
            operational_metrics_data.append(get_customer_lifetime_value_data(business_id))
            operational_metrics_data.append(get_employee_productivity_score_data(business_id))

            # Metrics that require more data/setup
            operational_metrics_data.append({"name": "Customer Acquisition Cost (CAC)", "value": "N/A", "target": 120, "trend": "neutral", "unit": "$", "note": "Requires marketing spend data."})
            operational_metrics_data.append({"name": "Conversion Rate", "value": "N/A", "target": 4.0, "trend": "neutral", "unit": "%", "note": "Requires lead tracking data."})
            operational_metrics_data.append({"name": "Churn Rate", "value": "N/A", "target": 4.0, "trend": "neutral", "unit": "%", "note": "Requires customer lifecycle data."})

        cols_op = st.columns(3)
        for i, metric in enumerate(operational_metrics_data):
            with cols_op[i % 3]:
                delta_display = ""
                if isinstance(metric.get('value'), (int, float)) and isinstance(metric.get('target'), (int, float)):
                    delta_val = metric['value'] - metric['target']
                    delta_display = f"{delta_val:+.1f} vs target"
                elif metric.get('trend') and metric['trend'] != "neutral":
                     delta_display = f"Trend: {metric['trend'].capitalize()}"

                value_display = f"{metric.get('unit', '')}{metric['value']}" if metric.get('unit') == "$" else f"{metric['value']}{metric.get('unit', '')}"
                if metric['value'] == "N/A": value_display = "N/A"


                delta_color_logic = "normal"
                if metric.get('value') != "N/A" and metric.get('target') is not None:
                    is_better = False
                    if metric['name'] in ["Churn Rate", "Customer Acquisition Cost (CAC)"]: # Lower is better
                        is_better = metric['value'] <= metric['target']
                    else: # Higher is better
                        is_better = metric['value'] >= metric['target']
                    delta_color_logic = "normal" if is_better else "inverse"
                
                st.metric(
                    metric["name"],
                    value_display,
                    delta_display,
                    delta_color=delta_color_logic
                )
                if metric.get("note"):
                    st.caption(metric["note"])
        
        st.write("### Operational Efficiency Trends (Project Progress)")
        num_efficiency_months = st.slider("Months for Efficiency Trend:", min_value=3, max_value=12, value=6, key="num_efficiency_months")
        with st.spinner("Fetching efficiency trends..."):
            df_efficiency = get_operational_efficiency_trends_data(business_id, num_months=num_efficiency_months)
        
        if not df_efficiency.empty:
            fig_efficiency = px.line(
                df_efficiency, 
                x="Month", 
                y="Efficiency",
                title=f"Avg. Project Progress (%) Over Last {num_efficiency_months} Months",
                markers=True
            )
            fig_efficiency.update_yaxes(range=[0, 100])
            st.plotly_chart(fig_efficiency, use_container_width=True)
        else:
            st.info("No project data available for efficiency trends.")
    
    with tab3:
        st.subheader("Custom Reports")
        
        report_type = st.selectbox("Select Report Type", [
            "Sales Performance", 
            "Employee Productivity", 
            "Inventory Analysis",
            "Marketing ROI (Simulated - Requires Data)" 
        ], key="custom_report_type_ei")
        
        # Common time period selection for reports
        st.write("#### Select Report Period")
        report_time_period = st.selectbox("Time Period", [
            "Last 7 Days", "Last 30 Days", "Last 90 Days",
            "Last Month", "Last Quarter", "Last Year", "Custom Range"
        ], key="report_time_period_ei")
        
        report_start_date, report_end_date = datetime.now().date(), datetime.now().date()

        if report_time_period == "Custom Range":
            col_sd, col_ed = st.columns(2)
            with col_sd:
                report_start_date = st.date_input("Start Date", datetime.now().date() - timedelta(days=30), key="report_sd_ei")
            with col_ed:
                report_end_date = st.date_input("End Date", datetime.now().date(), key="report_ed_ei")
        else:
            today = datetime.now().date()
            if report_time_period == "Last 7 Days": report_start_date = today - timedelta(days=6)
            elif report_time_period == "Last 30 Days": report_start_date = today - timedelta(days=29)
            elif report_time_period == "Last 90 Days": report_start_date = today - timedelta(days=89)
            elif report_time_period == "Last Month":
                end_of_last_month = today.replace(day=1) - timedelta(days=1)
                report_start_date, report_end_date = end_of_last_month.replace(day=1), end_of_last_month
            elif report_time_period == "Last Quarter":
                current_q_start = pd.Timestamp(today).to_period('Q').start_time.date()
                report_end_date = current_q_start - timedelta(days=1)
                report_start_date = pd.Timestamp(report_end_date).to_period('Q').start_time.date()
            elif report_time_period == "Last Year":
                report_start_date = today.replace(year=today.year - 1, month=1, day=1)
                report_end_date = today.replace(year=today.year - 1, month=12, day=31)
            
            if report_time_period not in ["Last Month", "Last Quarter", "Last Year", "Custom Range"]:
                 report_end_date = today
        
        st.info(f"Report period: {report_start_date.strftime('%Y-%m-%d')} to {report_end_date.strftime('%Y-%m-%d')}")

        if st.button("Generate Report", key="generate_custom_report_ei"):
            with st.spinner(f"Generating {report_type} report..."):
                report_content_for_download = f"Report Type: {report_type}\nPeriod: {report_start_date} to {report_end_date}\n\n"
                file_name_prefix = report_type.replace(" ", "_").split("(")[0].strip()

                if report_type == "Sales Performance":
                    sales_data = get_sales_performance_report_data(business_id, report_start_date, report_end_date)
                    st.success("Sales Performance Report Generated!")
                    
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Total Revenue", f"${sales_data['total_revenue']:,.2f}")
                    col2.metric("Number of Invoices", f"{sales_data['num_invoices']}")
                    col3.metric("Avg. Invoice Value", f"${sales_data['avg_invoice_value']:,.2f}")
                    st.write(f"**Top Product:** {sales_data['top_product_name']} (Revenue: ${sales_data['top_product_revenue']:,.2f})")
                    
                    report_content_for_download += (
                        f"Total Revenue: ${sales_data['total_revenue']:,.2f}\n"
                        f"Number of Invoices: {sales_data['num_invoices']}\n"
                        f"Average Invoice Value: ${sales_data['avg_invoice_value']:,.2f}\n"
                        f"Top Product: {sales_data['top_product_name']} (Revenue: ${sales_data['top_product_revenue']:,.2f})\n"
                    )

                elif report_type == "Employee Productivity":
                    emp_prod_data = get_employee_productivity_report_data(business_id, report_start_date, report_end_date)
                    st.success("Employee Productivity Report Generated!")
                    st.metric("Overall Avg. Performance Score", f"{emp_prod_data['avg_score_overall']}/10", 
                              help=f"Based on {emp_prod_data['num_employees_rated']} of {emp_prod_data['num_employees_total']} employees.")
                    st.write(f"**Top Performer:** {emp_prod_data['top_performer_name']} (Score: {emp_prod_data['top_performer_score']}/10)")
                    
                    report_content_for_download += (
                        f"Overall Average Performance Score: {emp_prod_data['avg_score_overall']}/10\n"
                        f"Total Employees: {emp_prod_data['num_employees_total']}, Rated: {emp_prod_data['num_employees_rated']}\n"
                        f"Top Performer: {emp_prod_data['top_performer_name']} (Score: {emp_prod_data['top_performer_score']}/10)\n\n"
                        "Department Average Scores:\n"
                    )
                    if emp_prod_data['department_scores']:
                        st.write("##### Department Average Scores:")
                        for dept in emp_prod_data['department_scores']:
                            st.write(f"- {dept['department']}: {dept['avg_score']}/10")
                            report_content_for_download += f"- {dept['department']}: {dept['avg_score']}/10\n"
                    else:
                        st.write("No department scores available.")
                        report_content_for_download += "No department scores available.\n"

                elif report_type == "Inventory Analysis":
                    inv_analysis_data = get_inventory_analysis_report_data(business_id, report_start_date, report_end_date)
                    st.success("Inventory Analysis Report Generated!")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Total Inventory Value (Selling Price)", f"${inv_analysis_data['total_inventory_value_at_price']:,.2f}")
                    col2.metric("Est. Inventory Value (Cost)", f"${inv_analysis_data['total_inventory_value_at_cost_est']:,.2f}", help="Estimated at 60% of selling price.")
                    col3.metric("Inventory Turnover", f"{inv_analysis_data['inventory_turnover']}x")
                    
                    st.metric("Number of Products", inv_analysis_data['num_products'])
                    st.metric("Slow-Moving Items Count", inv_analysis_data['slow_moving_items_count'], help=f"Items with >20 quantity and no sales in the last 90 days within the period {report_start_date} to {report_end_date}.")

                    report_content_for_download += (
                        f"Total Inventory Value (Selling Price): ${inv_analysis_data['total_inventory_value_at_price']:,.2f}\n"
                        f"Estimated Inventory Value (Cost): ${inv_analysis_data['total_inventory_value_at_cost_est']:,.2f}\n"
                        f"Inventory Turnover: {inv_analysis_data['inventory_turnover']}x\n"
                        f"Number of Products: {inv_analysis_data['num_products']}\n"
                        f"Slow-Moving Items Count: {inv_analysis_data['slow_moving_items_count']}\n\n"
                        "Top Slow-Moving Items (Name, Current Quantity, Days Since Last Sale in Period):\n"
                    )
                    if inv_analysis_data['top_slow_moving_items']:
                        st.write("##### Top Slow-Moving Items (Max 5 Shown):")
                        for item in inv_analysis_data['top_slow_moving_items']:
                            st.write(f"- {item['name']} (Qty: {item['quantity']}, Last Sale: {item['days_since_last_sale_est']} days ago)")
                            report_content_for_download += f"- {item['name']} (Qty: {item['quantity']}, Last Sale: {item['days_since_last_sale_est']} days ago)\n"
                    else:
                        st.write("No significant slow-moving items identified based on current criteria.")
                        report_content_for_download += "No significant slow-moving items identified.\n"
                
                elif report_type == "Marketing ROI (Simulated - Requires Data)":
                    st.warning("Marketing ROI calculation requires dedicated marketing spend and campaign data, which is not yet implemented in the database.")
                    st.info("Displaying simulated data as a placeholder:")
                    st.write("### Marketing ROI Report (Simulated)")
                    st.write("- Total Marketing Spend: $25,000")
                    st.write("- Revenue Attributed to Marketing: $125,000")
                    st.write("- ROI: 5.0x")
                    st.write("- Top Performing Channel: Social Media Ads")
                    report_content_for_download += (
                        "Marketing ROI Report (Simulated - Requires Data Integration)\n"
                        "Total Marketing Spend: $25,000\n"
                        "Revenue Attributed to Marketing: $125,000\n"
                        "ROI: 5.0x\n"
                    )
                
                st.download_button(
                    "Download Report (TXT)",
                    data=report_content_for_download,
                    file_name=f"{file_name_prefix}_Report_{report_start_date}_{report_end_date}.txt",
                    mime="text/plain",
                    key=f"download_{file_name_prefix}_ei"
                )

# Enterprise Intelligence Module
def enterprise_intelligence(business_id, ai_models): 
    """Streamlit module for Enterprise Intelligence Dashboards."""
    st.header("📊 Enterprise Intelligence Dashboards")
    st.info("Provides key business metrics and reports based on your data.")
    
    tab1, tab2, tab3 = st.tabs([
        "Financial Performance", 
        "Operational Metrics", 
        "Custom Reports"
    ])
    
    conn = get_db_connection()  # Initialize connection here
    cur = conn.cursor()    
    with tab1:
        st.subheader("Financial Performance")
        
        fin_period_type = st.selectbox("View Financials By:", ["Monthly", "Quarterly"], key="fin_period_type_ei")
        num_fin_periods = st.slider("Number of Past Periods:", min_value=3, max_value=24, 
                                    value=12 if fin_period_type == "Monthly" else 8, key="num_fin_periods_ei")

        with st.spinner("Fetching financial data..."):
            df_finance = get_financial_performance_data(business_id, num_periods=num_fin_periods, period_type=fin_period_type[0])
        
        if not df_finance.empty:
            fig_finance = px.line(
                df_finance, 
                x="Period", 
                y=["Revenue", "Expenses", "Profit"],
                title=f"{fin_period_type} Financial Performance ({num_fin_periods} Periods)"
            )
            st.plotly_chart(fig_finance, use_container_width=True)
            
            st.write(f"### Key Financial Metrics (Latest {fin_period_type[:-2]})")
            latest_period_data = df_finance.iloc[-1]
            revenue_latest = latest_period_data['Revenue']
            expenses_latest = latest_period_data['Expenses']
            profit_latest = latest_period_data['Profit']
            
            revenue_prev, expenses_prev, profit_prev = 0.0, 0.0, 0.0
            delta_label = f"vs Prev. {fin_period_type[:-2]}"
            if len(df_finance) >= 2:
                prev_period_data = df_finance.iloc[-2]
                revenue_prev = prev_period_data['Revenue']
                expenses_prev = prev_period_data['Expenses']
                profit_prev = prev_period_data['Profit']
            else:
                delta_label = "(No Previous Period Data)"

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric(f"Total Revenue", f"${revenue_latest:,.0f}", 
                         f"{calculate_delta(revenue_latest, revenue_prev)} {delta_label}")
            with col2:
                st.metric(f"Total Expenses (Est.)", f"${expenses_latest:,.0f}", 
                         f"{calculate_delta(expenses_latest, expenses_prev)} {delta_label}")
                st.caption("Expenses estimated based on total salaries.")
            with col3:
                st.metric(f"Net Profit (Est.)", f"${profit_latest:,.0f}", 
                         f"{calculate_delta(profit_latest, profit_prev)} {delta_label}")
        else:
            st.info("No financial data available to display for the selected period.")
    
    with tab2:
        st.subheader("Operational Metrics")
        
        operational_metrics_data = []
        with st.spinner("Fetching operational metrics..."):
            operational_metrics_data.append(get_inventory_turnover_data(business_id))
            operational_metrics_data.append(get_customer_lifetime_value_data(business_id))
            operational_metrics_data.append(get_employee_productivity_score_data(business_id))

            # Metrics that require more data/setup
            operational_metrics_data.append({"name": "Customer Acquisition Cost (CAC)", "value": "N/A", "target": 120, "trend": "neutral", "unit": "$", "note": "Requires marketing spend data."})
            operational_metrics_data.append({"name": "Conversion Rate", "value": "N/A", "target": 4.0, "trend": "neutral", "unit": "%", "note": "Requires lead tracking data."})
            operational_metrics_data.append({"name": "Churn Rate", "value": "N/A", "target": 4.0, "trend": "neutral", "unit": "%", "note": "Requires customer lifecycle data."})

        cols_op = st.columns(3)
        for i, metric in enumerate(operational_metrics_data):
            with cols_op[i % 3]:
                delta_display = ""
                if isinstance(metric.get('value'), (int, float)) and isinstance(metric.get('target'), (int, float)):
                    delta_val = metric['value'] - metric['target']
                    delta_display = f"{delta_val:+.1f} vs target"
                elif metric.get('trend') and metric['trend'] != "neutral":
                     delta_display = f"Trend: {metric['trend'].capitalize()}"

                value_display = f"{metric.get('unit', '')}{metric['value']}" if metric.get('unit') == "$" else f"{metric['value']}{metric.get('unit', '')}"
                if metric['value'] == "N/A": value_display = "N/A"


                delta_color_logic = "normal"
                if metric.get('value') != "N/A" and metric.get('target') is not None:
                    is_better = False
                    if metric['name'] in ["Churn Rate", "Customer Acquisition Cost (CAC)"]: # Lower is better
                        is_better = metric['value'] <= metric['target']
                    else: # Higher is better
                        is_better = metric['value'] >= metric['target']
                    delta_color_logic = "normal" if is_better else "inverse"
                
                st.metric(
                    metric["name"],
                    value_display,
                    delta_display,
                    delta_color=delta_color_logic
                )
                if metric.get("note"):
                    st.caption(metric["note"])
        
        st.write("### Operational Efficiency Trends (Project Progress)")
        num_efficiency_months = st.slider("Months for Efficiency Trend:", min_value=3, max_value=12, value=6, key="num_efficiency_months")
        with st.spinner("Fetching efficiency trends..."):
            df_efficiency = get_operational_efficiency_trends_data(business_id, num_months=num_efficiency_months)
        
        if not df_efficiency.empty:
            fig_efficiency = px.line(
                df_efficiency, 
                x="Month", 
                y="Efficiency",
                title=f"Avg. Project Progress (%) Over Last {num_efficiency_months} Months",
                markers=True
            )
            fig_efficiency.update_yaxes(range=[0, 100])
            st.plotly_chart(fig_efficiency, use_container_width=True)
        else:
            st.info("No project data available for efficiency trends.")
    
    with tab3:
        st.subheader("Custom Reports")
        
        report_type = st.selectbox("Select Report Type", [
            "Sales Performance", 
            "Employee Productivity", 
            "Inventory Analysis",
            "Marketing ROI (Simulated - Requires Data)" 
        ], key="custom_report_type_ei")
        
        # Common time period selection for reports
        st.write("#### Select Report Period")
        report_time_period = st.selectbox("Time Period", [
            "Last 7 Days", "Last 30 Days", "Last 90 Days",
            "Last Month", "Last Quarter", "Last Year", "Custom Range"
        ], key="report_time_period_ei")
        
        report_start_date, report_end_date = datetime.now().date(), datetime.now().date()

        if report_time_period == "Custom Range":
            col_sd, col_ed = st.columns(2)
            with col_sd:
                report_start_date = st.date_input("Start Date", datetime.now().date() - timedelta(days=30), key="report_sd_ei")
            with col_ed:
                report_end_date = st.date_input("End Date", datetime.now().date(), key="report_ed_ei")
        else:
            today = datetime.now().date()
            if report_time_period == "Last 7 Days": report_start_date = today - timedelta(days=6)
            elif report_time_period == "Last 30 Days": report_start_date = today - timedelta(days=29)
            elif report_time_period == "Last 90 Days": report_start_date = today - timedelta(days=89)
            elif report_time_period == "Last Month":
                end_of_last_month = today.replace(day=1) - timedelta(days=1)
                report_start_date, report_end_date = end_of_last_month.replace(day=1), end_of_last_month
            elif report_time_period == "Last Quarter":
                current_q_start = pd.Timestamp(today).to_period('Q').start_time.date()
                report_end_date = current_q_start - timedelta(days=1)
                report_start_date = pd.Timestamp(report_end_date).to_period('Q').start_time.date()
            elif report_time_period == "Last Year":
                report_start_date = today.replace(year=today.year - 1, month=1, day=1)
                report_end_date = today.replace(year=today.year - 1, month=12, day=31)
            
            if report_time_period not in ["Last Month", "Last Quarter", "Last Year", "Custom Range"]:
                 report_end_date = today
        
        st.info(f"Report period: {report_start_date.strftime('%Y-%m-%d')} to {report_end_date.strftime('%Y-%m-%d')}")

        if st.button("Generate Report", key="generate_custom_report_ei"):
            with st.spinner(f"Generating {report_type} report..."):
                report_content_for_download = f"Report Type: {report_type}\nPeriod: {report_start_date} to {report_end_date}\n\n"
                file_name_prefix = report_type.replace(" ", "_").split("(")[0].strip()

                if report_type == "Sales Performance":
                    sales_data = get_sales_performance_report_data(business_id, report_start_date, report_end_date)
                    st.success("Sales Performance Report Generated!")
                    
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Total Revenue", f"${sales_data['total_revenue']:,.2f}")
                    col2.metric("Number of Invoices", f"{sales_data['num_invoices']}")
                    col3.metric("Avg. Invoice Value", f"${sales_data['avg_invoice_value']:,.2f}")
                    st.write(f"**Top Product:** {sales_data['top_product_name']} (Revenue: ${sales_data['top_product_revenue']:,.2f})")
                    
                    report_content_for_download += (
                        f"Total Revenue: ${sales_data['total_revenue']:,.2f}\n"
                        f"Number of Invoices: {sales_data['num_invoices']}\n"
                        f"Average Invoice Value: ${sales_data['avg_invoice_value']:,.2f}\n"
                        f"Top Product: {sales_data['top_product_name']} (Revenue: ${sales_data['top_product_revenue']:,.2f})\n"
                    )

                elif report_type == "Employee Productivity":
                    emp_prod_data = get_employee_productivity_report_data(business_id, report_start_date, report_end_date)
                    st.success("Employee Productivity Report Generated!")
                    st.metric("Overall Avg. Performance Score", f"{emp_prod_data['avg_score_overall']}/10", 
                              help=f"Based on {emp_prod_data['num_employees_rated']} of {emp_prod_data['num_employees_total']} employees.")
                    st.write(f"**Top Performer:** {emp_prod_data['top_performer_name']} (Score: {emp_prod_data['top_performer_score']}/10)")
                    
                    report_content_for_download += (
                        f"Overall Average Performance Score: {emp_prod_data['avg_score_overall']}/10\n"
                        f"Total Employees: {emp_prod_data['num_employees_total']}, Rated: {emp_prod_data['num_employees_rated']}\n"
                        f"Top Performer: {emp_prod_data['top_performer_name']} (Score: {emp_prod_data['top_performer_score']}/10)\n\n"
                        "Department Average Scores:\n"
                    )
                    if emp_prod_data['department_scores']:
                        st.write("##### Department Average Scores:")
                        for dept in emp_prod_data['department_scores']:
                            st.write(f"- {dept['department']}: {dept['avg_score']}/10")
                            report_content_for_download += f"- {dept['department']}: {dept['avg_score']}/10\n"
                    else:
                        st.write("No department scores available.")
                        report_content_for_download += "No department scores available.\n"

                elif report_type == "Inventory Analysis":
                    inv_analysis_data = get_inventory_analysis_report_data(business_id, report_start_date, report_end_date)
                    st.success("Inventory Analysis Report Generated!")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Total Inventory Value (Selling Price)", f"${inv_analysis_data['total_inventory_value_at_price']:,.2f}")
                    col2.metric("Est. Inventory Value (Cost)", f"${inv_analysis_data['total_inventory_value_at_cost_est']:,.2f}", help="Estimated at 60% of selling price.")
                    col3.metric("Inventory Turnover", f"{inv_analysis_data['inventory_turnover']}x")
                    
                    st.metric("Number of Products", inv_analysis_data['num_products'])
                    st.metric("Slow-Moving Items Count", inv_analysis_data['slow_moving_items_count'], help=f"Items with >20 quantity and no sales in the last 90 days within the period {report_start_date} to {report_end_date}.")

                    report_content_for_download += (
                        f"Total Inventory Value (Selling Price): ${inv_analysis_data['total_inventory_value_at_price']:,.2f}\n"
                        f"Estimated Inventory Value (Cost): ${inv_analysis_data['total_inventory_value_at_cost_est']:,.2f}\n"
                        f"Inventory Turnover: {inv_analysis_data['inventory_turnover']}x\n"
                        f"Number of Products: {inv_analysis_data['num_products']}\n"
                        f"Slow-Moving Items Count: {inv_analysis_data['slow_moving_items_count']}\n\n"
                        "Top Slow-Moving Items (Name, Current Quantity, Days Since Last Sale in Period):\n"
                    )
                    if inv_analysis_data['top_slow_moving_items']:
                        st.write("##### Top Slow-Moving Items (Max 5 Shown):")
                        for item in inv_analysis_data['top_slow_moving_items']:
                            st.write(f"- {item['name']} (Qty: {item['quantity']}, Last Sale: {item['days_since_last_sale_est']} days ago)")
                            report_content_for_download += f"- {item['name']} (Qty: {item['quantity']}, Last Sale: {item['days_since_last_sale_est']} days ago)\n"
                    else:
                        st.write("No significant slow-moving items identified based on current criteria.")
                        report_content_for_download += "No significant slow-moving items identified.\n"
                
                elif report_type == "Marketing ROI (Simulated - Requires Data)":
                    st.warning("Marketing ROI calculation requires dedicated marketing spend and campaign data, which is not yet implemented in the database.")
                    st.info("Displaying simulated data as a placeholder:")
                    st.write("### Marketing ROI Report (Simulated)")
                    st.write("- Total Marketing Spend: $25,000")
                    st.write("- Revenue Attributed to Marketing: $125,000")
                    st.write("- ROI: 5.0x")
                    st.write("- Top Performing Channel: Social Media Ads")
                    report_content_for_download += (
                        "Marketing ROI Report (Simulated - Requires Data Integration)\n"
                        "Total Marketing Spend: $25,000\n"
                        "Revenue Attributed to Marketing: $125,000\n"
                        "ROI: 5.0x\n"
                    )
                
                st.download_button(
                    "Download Report (TXT)",
                    data=report_content_for_download,
                    file_name=f"{file_name_prefix}_Report_{report_start_date}_{report_end_date}.txt",
                    mime="text/plain",
                    key=f"download_{file_name_prefix}_ei"
                )

# AI Market Forecasting Module
def market_forecasting(business_id, ai_models):
    """Streamlit module for AI Market Forecasting."""
    st.header("🔮 AI Market Forecasting")
    st.info("Generates market forecasts using AI based on general trends and your inputs.")
    
    tab1, tab2 = st.tabs(["Trend Analysis", "Predictive Insights"])

    conn = get_db_connection()  # Initialize connection here
    cur = conn.cursor()  
    with tab1:
        st.subheader("Market Trend Analysis")
        
        industry = st.selectbox("Select Industry for Analysis", [
            "Technology", "Retail", "Healthcare", "Finance", 
            "Manufacturing", "Energy", "Transportation", "Other"
        ], key="trend_analysis_industry")
        
        metric = st.selectbox("Select Metric", [
            "Market Size", "Growth Rate", "Adoption Rate", 
            "Investment Activity", "Regulatory Changes", "Consumer Behavior"
        ], key="trend_analysis_metric")
        
        if st.button("Analyze Trends", key="analyze_trends_btn"):
            with st.spinner("Analyzing market trends..."):
                # Generate trend analysis using AI
                prompt = f"""
                Provide a detailed analysis of {metric} trends in the {industry} industry. 
                Discuss recent changes, key drivers, challenges, and potential future direction.
                """
                analysis = ai_models.generate_text(prompt, max_length=1000)
                
                st.subheader(f"{industry} Industry {metric} Analysis")
                st.write(analysis)
                
                # Simulated trend chart (simple line)
                try:
                    years = list(range(datetime.now().year - 4, datetime.now().year + 3)) # Past 5, next 3
                    # Simulate values with some trend and noise
                    base_value = 10 if "Rate" in metric else 100
                    simulated_values = [base_value * (1 + (i * 0.05) + np.random.normal(0, 0.02)) for i in range(len(years))]
                    
                    df_trend = pd.DataFrame({"Year": years, metric: simulated_values})
                    fig = px.line(
                        df_trend, 
                        x="Year", 
                        y=metric,
                        title=f"{metric} Trend for {industry} Industry (Simulated)",
                        markers=True
                    )
                    st.plotly_chart(fig, use_container_width=True)
                except Exception as e:
                    st.warning(f"Could not generate simulated trend chart: {e}")

    with tab2:
        st.subheader("Predictive Insights")
        st.info("Generates a predictive forecast using AI based on your inputs.")
        
        cur.execute(
            "SELECT name FROM products WHERE business_id = %s ORDER BY name",
            (business_id,)
        )
        products = [p[0] for p in cur.fetchall()]
        
        forecast_subject = st.selectbox("Forecast Subject", ["Overall Market", "Specific Product"] + products, key="predictive_forecast_subject")
        
        if forecast_subject == "Specific Product" and not products:
             st.warning("Add products in the Inventory module to forecast for a specific product.")
             selected_product_for_forecast = None
        elif forecast_subject in products:
             selected_product_for_forecast = forecast_subject
             forecast_subject_text = f"your product '{selected_product_for_forecast}'"
        elif forecast_subject == "Specific Product" and products:
             selected_product_for_forecast = st.selectbox("Select Product", products, key="select_product_for_predictive_forecast")
             forecast_subject_text = f"your product '{selected_product_for_forecast}'"
        else:
             selected_product_for_forecast = None
             forecast_subject_text = "the overall market relevant to your business" # General prompt

        forecast_period = st.selectbox("Forecast Period", ["3 months", "6 months", "1 year", "3 years"], key="predictive_forecast_period")
        
        if st.button("Generate Predictive Forecast", key="generate_predictive_forecast_btn"):
            with st.spinner("Generating predictive forecast..."):
                prompt = f"""
                Generate a predictive market forecast for {forecast_subject_text} over the next {forecast_period}.
                Include expected growth, potential risks, and strategic recommendations to capitalize on or mitigate these predictions.
                """
                forecast_text = ai_models.generate_text(prompt, max_length=1000)
                
                st.subheader(f"Predictive Forecast for {forecast_subject_text}")
                st.write(forecast_text)
                
                # Simulated forecast chart (simple line)
                try:
                    if forecast_period == "3 months": periods = pd.date_range(start=datetime.now(), periods=3, freq='MS')
                    elif forecast_period == "6 months": periods = pd.date_range(start=datetime.now(), periods=6, freq='MS')
                    elif forecast_period == "1 year": periods = pd.date_range(start=datetime.now(), periods=4, freq='QS') # Quarterly for a year
                    elif forecast_period == "3 years": periods = pd.date_range(start=datetime.now(), periods=3, freq='YS') # Yearly for 3 years
                    else: periods = pd.date_range(start=datetime.now(), periods=4, freq='QS') # Default

                    # Simulate growth based on a base value (e.g., last period's revenue or a base product value)
                    base_value = 100 # Arbitrary base
                    if selected_product_for_forecast:
                         try:
                             cur.execute("SELECT AVG(price) FROM products WHERE name = %s AND business_id = %s", (selected_product_for_forecast, business_id))
                             avg_price = cur.fetchone()[0]
                             if avg_price is not None: base_value = float(avg_price) * 10 # Assume avg sale of 10 units
                         except: pass

                    # Simulate values with a general upward trend and some noise
                    simulated_values = [base_value * (1 + (i * 0.07) + np.random.normal(0, 0.03)) for i in range(len(periods))] # Slightly higher growth/noise for predictive

                    df_forecast = pd.DataFrame({"Period": periods, "Projected Value": simulated_values})
                    df_forecast['PeriodLabel'] = df_forecast['Period'].dt.strftime('%Y-%m') if forecast_period in ["3 months", "6 months"] else df_forecast['Period'].dt.strftime('%Y') + (df_forecast['Period'].dt.quarter.astype(str).apply(lambda x: f"-Q{x}") if forecast_period == "1 year" else "")

                    fig = px.line(
                        df_forecast, 
                        x="PeriodLabel", 
                        y="Projected Value",
                        title=f"{forecast_period} Projected Trend for {forecast_subject_text} (Simulated)",
                        labels={"PeriodLabel": "Period", "Projected Value": "Projected Value ($)"},
                        markers=True
                    )
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Save forecast data (simplified - saving just the final value)
                    try:
                        final_projected_value = simulated_values[-1] if simulated_values else 0.0
                        cur.execute(
                            """INSERT INTO market_data 
                            (business_id, industry, metric, value, date, source) 
                            VALUES (%s, %s, %s, %s, %s, %s)""",
                            (business_id, forecast_subject_text, f"{forecast_period}_predictive_forecast_est", float(final_projected_value), datetime.now().date(), "AI Predictive Forecast")
                        )
                        conn.commit()
                        st.success("Predictive forecast data saved.")
                    except Exception as e:
                        st.error(f"Error saving predictive forecast data: {e}")

                except Exception as e:
                    st.error(f"Error generating predictive forecast chart: {e}")
        
    cur.close(); conn.close()

# User-to-User Chat Module (for Business)
def chat_module(business_id, ai_models): # ai_models not used for user-to-user
    """Streamlit module for Business users to chat with Investors and Service Providers."""
    st.header("💬 Messaging")
    conn = get_db_connection(); cur = conn.cursor()

    TYPE_MAP_DISPLAY_TO_DB = {"Businesses": "business", "Investors": "investor", "Service Providers": "service_provider"}
    TYPE_MAP_DB_TO_DISPLAY = {v: k for k, v in TYPE_MAP_DISPLAY_TO_DB.items()}

    # Initialize session state for chat partner details if not present
    if 'chat_partner_db_type' not in st.session_state: st.session_state.chat_partner_db_type = None
    if 'chat_partner_id' not in st.session_state: st.session_state.chat_partner_id = None
    if 'chat_partner_name' not in st.session_state: st.session_state.chat_partner_name = None

    # --- Sidebar: Select Chat Partner ---
    st.sidebar.subheader("Start or Continue Chat")
    
    st.sidebar.write("**Recent Chats:**")
    try:
        query_recent_chats = """
        SELECT DISTINCT
            CASE WHEN sender_id = %(current_business_id)s THEN receiver_type ELSE sender_type END as partner_type,
            CASE WHEN sender_id = %(current_business_id)s THEN receiver_id ELSE sender_id END as partner_id
        FROM messages
        WHERE (sender_type = 'business' AND sender_id = %(current_business_id)s)
           OR (receiver_type = 'business' AND receiver_id = %(current_business_id)s);
        """
        cur.execute(query_recent_chats, {'current_business_id': business_id})
        recent_partners_db = cur.fetchall()
        
        recent_chat_options = {} 
        for p_type_db, p_id_db in recent_partners_db: 
            p_id_int = int(p_id_db) # Ensure p_id is int
            partner_table_map = {'business': 'businesses', 'investor': 'investors', 'service_provider': 'service_providers'}
            table_name = partner_table_map.get(p_type_db)
            if table_name:
                cur.execute(sql.SQL("SELECT name FROM {} WHERE id = %s").format(sql.Identifier(table_name)), (p_id_int,))
                name_res = cur.fetchone()
                if name_res:
                    display_name = f"{name_res[0]} ({TYPE_MAP_DB_TO_DISPLAY.get(p_type_db, p_type_db.capitalize())})"
                    recent_chat_options[display_name] = (p_type_db, p_id_int)
        
        selected_recent_chat_display = st.sidebar.selectbox("Select from recent chats:", ["-- New Chat --"] + list(recent_chat_options.keys()), key="select_recent_chat_business")
        if selected_recent_chat_display != "-- New Chat --":
            p_type, p_id = recent_chat_options[selected_recent_chat_display]
            if st.session_state.chat_partner_id != p_id or st.session_state.chat_partner_db_type != p_type:
                st.session_state.chat_partner_db_type, st.session_state.chat_partner_id, st.session_state.chat_partner_name = p_type, p_id, selected_recent_chat_display.split(" (")[0]
                st.rerun()
    except Exception as e: st.sidebar.error(f"Error loading recent chats: {e}")

    st.sidebar.markdown("---"); st.sidebar.subheader("Find New Chat Partner")
    selected_chat_type_display = st.sidebar.selectbox("Chat with:", list(TYPE_MAP_DISPLAY_TO_DB.keys()), key="select_chat_type_display_business")
    
    partner_options_new = {}
    target_db_type = TYPE_MAP_DISPLAY_TO_DB.get(selected_chat_type_display)
    try:
        if target_db_type == "business": cur.execute("SELECT id, name FROM businesses WHERE id != %s ORDER BY name", (business_id,))
        elif target_db_type == "investor": cur.execute("SELECT id, name FROM investors ORDER BY name")
        elif target_db_type == "service_provider": cur.execute("SELECT id, name, service_type FROM service_providers ORDER BY name")
        
        partners_list = cur.fetchall()
        for partner_data in partners_list:
            pid, pname = partner_data[0], partner_data[1]
            label_suffix = f" ({selected_chat_type_display.rstrip('s')})"
            if target_db_type == "service_provider" and len(partner_data) > 2: label_suffix = f" ({partner_data[2].capitalize()} - Service Provider)"
            partner_options_new[f"{pname}{label_suffix}"] = (target_db_type, pid)
    except Exception as e: st.sidebar.error(f"Error fetching new partners: {e}")

    selected_partner_display_new = st.sidebar.selectbox(f"Select {selected_chat_type_display.rstrip('s')}:", ["-- Select --"] + list(partner_options_new.keys()), key="select_new_chat_partner_business")
    if selected_partner_display_new != "-- Select --":
        p_type, p_id = partner_options_new[selected_partner_display_new]
        if st.session_state.chat_partner_id != p_id or st.session_state.chat_partner_db_type != p_type:
            st.session_state.chat_partner_db_type, st.session_state.chat_partner_id, st.session_state.chat_partner_name = p_type, p_id, selected_partner_display_new.split(" (")[0]
            st.rerun()

    # --- Main Chat Area ---
    if st.session_state.chat_partner_db_type and st.session_state.chat_partner_id:
        partner_db_type, partner_id, partner_name = st.session_state.chat_partner_db_type, st.session_state.chat_partner_id, st.session_state.chat_partner_name
        st.subheader(f"Chat with {partner_name} ({TYPE_MAP_DB_TO_DISPLAY.get(partner_db_type, partner_db_type).rstrip('s')})")

        # --- Display Profile Snippet ---
        with st.expander(f"View {partner_name}'s Profile Snippet", expanded=False):
            try:
                profile_info = "Not available."
                if partner_db_type == 'business':
                    cur.execute("SELECT email, created_at FROM businesses WHERE id = %s", (partner_id,))
                    p_data = cur.fetchone()
                    if p_data: profile_info = f"Email: {p_data[0]}\nJoined: {p_data[1].strftime('%Y-%m-%d')}"
                elif partner_db_type == 'investor':
                    cur.execute("SELECT firm, email, investment_focus, profile_description FROM investors WHERE id = %s", (partner_id,))
                    p_data = cur.fetchone()
                    if p_data: profile_info = f"Firm: {p_data[0] or 'N/A'}\nEmail: {p_data[1]}\nFocus: {p_data[2] or 'N/A'}\nAbout: {p_data[3] or 'N/A'}"
                elif partner_db_type == 'service_provider':
                    cur.execute("SELECT service_type, contact_email, rating, experience_years, profile_description, specializations FROM service_providers WHERE id = %s", (partner_id,))
                    p_data = cur.fetchone()
                    if p_data: profile_info = f"Service: {p_data[0].capitalize()}\nEmail: {p_data[1]}\nRating: {p_data[2]}/5\nExperience: {p_data[3]} years\nAbout: {p_data[4] or 'N/A'}\nSpecializations: {', '.join(p_data[5]) if p_data[5] else 'N/A'}"
                st.text(profile_info)
            except Exception as e_prof:
                st.warning(f"Could not load profile snippet: {e_prof}")
        
        # --- Fetch and Display Messages ---
        current_user_db_type = "business" 
        current_user_id = business_id

        try:
            query_messages = """SELECT sender_type, sender_id, content, created_at, read_at FROM messages WHERE
                                (sender_type = %(current_type)s AND sender_id = %(current_id)s AND receiver_type = %(partner_type)s AND receiver_id = %(partner_id)s) OR
                                (sender_type = %(partner_type)s AND sender_id = %(partner_id)s AND receiver_type = %(current_type)s AND receiver_id = %(current_id)s)
                                ORDER BY created_at ASC;"""
            cur.execute(query_messages, {
                'current_type': current_user_db_type, 'current_id': current_user_id,
                'partner_type': partner_db_type, 'partner_id': partner_id
            })
            messages = cur.fetchall()

            chat_container = st.container() 
            with chat_container:
                if not messages: st.info(f"No messages with {partner_name}. Start chatting!")
                for s_type, s_id, content, created_at, read_at in messages:
                    ts = created_at.strftime('%Y-%m-%d %H:%M')
                    if s_type == current_user_db_type and s_id == current_user_id: # Message by current business user
                        read_status = " (Read)" if read_at else " (Delivered)"
                        st.markdown(f"<div style='text-align: right; margin-left: 20%; margin-bottom: 5px; padding: 10px; background-color: #DCF8C6; border-radius: 10px;'><b>You</b> ({ts}){read_status}:<br>{content}</div>", unsafe_allow_html=True)
                    else: # Message by partner
                        st.markdown(f"<div style='text-align: left; margin-right: 20%; margin-bottom: 5px; padding: 10px; background-color: #FFFFFF; border-radius: 10px; border: 1px solid #E0E0E0;'><b>{partner_name}</b> ({ts}):<br>{content}</div>", unsafe_allow_html=True)
            
            # Auto-scroll (basic attempt)
            if messages: st.markdown("<script>window.scrollTo(0,document.body.scrollHeight);</script>", unsafe_allow_html=True)

        except Exception as e: st.error(f"Error fetching messages: {e}")

        # --- Message Input Form ---
        with st.form("new_message_form_business", clear_on_submit=True):
            new_msg_content = st.text_area("Your message:", key=f"new_msg_business_{partner_id}")
            send_button = st.form_submit_button("Send")

            if send_button and new_msg_content.strip():
                try:
                    insert_query = """
                        INSERT INTO messages (sender_type, sender_id, receiver_type, receiver_id, content)
                        VALUES (%s, %s, %s, %s, %s);
                    """
                    cur.execute(insert_query, (
                        current_user_db_type, current_user_id,
                        partner_db_type, partner_id,
                        new_msg_content.strip()
                    ))
                    conn.commit()
                    st.rerun()
                except Exception as e: st.error(f"Error sending: {e}")
            elif send_button and not new_msg_content.strip(): st.warning("Message cannot be empty.")
    else:
        st.info("Select a user, investor, or service provider from the sidebar to start chatting, or choose a recent chat.")

    cur.close(); conn.close()


# Main Application
def main():
    """Main function to run the Streamlit application."""
    
    st.set_page_config(page_title="GrowBis", page_icon="🚀", layout="wide", initial_sidebar_state="expanded")
    init_db()
    ai_models = load_ai_models()
    st.markdown("""<style>.main {background-color: #f5f5f5;} .sidebar .sidebar-content {background-color: #2c3e50; color: white;} 
                   h1 {color: #2c3e50;} .stButton>button {background-color: #3498db; color: white;} 
                   .stDownloadButton>button {background-color: #2ecc71; color: white;}</style>""", unsafe_allow_html=True)

    # Handle registration page display
    if st.session_state.get('show_registration_form'):
        reg_type = st.session_state.show_registration_form
        # Keep show_registration_form in state until back button is clicked
        
        if reg_type == 'investor': investor_registration_page()
        elif reg_type == 'service_provider': service_provider_registration_page()
        
        if st.button("Back to Login Selection", key=f"back_to_login_from_{reg_type}"):
            if 'show_registration_form' in st.session_state: del st.session_state.show_registration_form
            # Keep user_type_for_login and registered_email to prefill login
            st.rerun()
        return
    
    # Handle post-registration redirect to login
    if st.session_state.get('show_login_main'): 
        if 'show_login_main' in st.session_state: del st.session_state.show_login_main
        # user_type_for_login and registered_email should be set to prefill login
        # Fall through to login_page()

    # Check authentication for logged-in users
    if not check_auth():
        login_page()
        return

    # --- User is Authenticated ---
    user_type = st.session_state.user_type
    user_name = st.session_state.user_name
    entity_id = st.session_state.logged_in_entity_id

    st.sidebar.title(f"Welcome, {user_name}!")
    if st.sidebar.button("Logout", key="main_logout_btn"):
        logout(); st.rerun()
    
    if user_type == 'business':
        conn = get_db_connection(); cur = conn.cursor()
        cur.execute("SELECT name, subscription_type, subscription_expiry FROM businesses WHERE id = %s", (entity_id,))
        business_info = cur.fetchone(); cur.close(); conn.close()
        if not business_info: st.error("Business info not found."); logout(); st.rerun(); return

        business_name_display, sub_type, sub_expiry = business_info[0], business_info[1], business_info[2]
        days_left = (sub_expiry - datetime.now().date()).days
        sub_status = "Expired" if days_left < 0 else (f"Expires in {days_left} days" if days_left < 30 else f"Active until {sub_expiry}")
        
        st.title(f"🚀 {business_name_display} - GrowBis Platform")
        st.markdown(f"**Subscription:** {sub_type.capitalize()} • {sub_status}")
        
        st.sidebar.markdown("---"); st.sidebar.subheader("Business Modules")
        modules = [
            "Dashboard", "Inventory & Billing", "HR Tools", "Project Manager", 
            "Document Generator", "Market Analysis Tool", "Market Doubt Assistant (AI Chatbot)",
            "Investor & Agent Dashboards", "Govt/Private Schemes & News Alerts", 
            "Opportunity Director", "Voice Navigation", "Pitching Helper", 
            "Strategy Generator", "Hiring Helper", "Tax & GST Filing", 
            "IPO & Cap Table Management", "Legal, CA & Insurance Marketplace",
            "Enterprise Intelligence Dashboards", "AI Market Forecasting", "Messaging"
        ]
        selected_module = st.sidebar.selectbox("Select Module", modules, key="business_module_select")

        # Business module routing
        if selected_module == "Dashboard":
            st.header("📊 Dashboard")
            st.write("### Business Overview (Quarterly)")
            today = datetime.now().date()
            cq_start, cq_end = get_quarter_dates(today)
            pq_start, pq_end = get_previous_quarter_dates(today)
            current_q_revenue = get_dashboard_financials(entity_id, cq_start, cq_end)
            prev_q_revenue = get_dashboard_financials(entity_id, pq_start, pq_end)
            total_monthly_salary = get_total_monthly_salary_expense(entity_id)
            current_q_expenses_est = total_monthly_salary * 3
            prev_q_expenses_est = total_monthly_salary * 3 
            current_q_profit_est = current_q_revenue - current_q_expenses_est
            prev_q_profit_est = prev_q_revenue - prev_q_expenses_est

            def format_currency(value): return f"${value:,.0f}"
            def calculate_delta_string(current, previous):
                if previous == 0: return "N/A" if current == 0 else "New"
                return f"{(current - previous) / abs(previous) * 100:.1f}%" if previous != 0 else "N/A"

            col1, col2, col3 = st.columns(3)
            with col1: st.metric("Revenue (Current Qtr)", format_currency(current_q_revenue), f"{calculate_delta_string(current_q_revenue, prev_q_revenue)} vs Prev. Qtr")
            with col2: 
                st.metric("Est. Expenses (Current Qtr)", format_currency(current_q_expenses_est), f"{calculate_delta_string(current_q_expenses_est, prev_q_expenses_est)} vs Prev. Qtr (Est.)")
                st.caption("Expenses estimated based on 3x current monthly salaries.")
            with col3: st.metric("Est. Profit (Current Qtr)", format_currency(current_q_profit_est), f"{calculate_delta_string(current_q_profit_est, prev_q_profit_est)} vs Prev. Qtr (Est.)")
            
            st.write("### Recent Activity")
            activities = get_recent_activities_for_dashboard(entity_id, limit=5)
            if activities:
                for activity in activities:
                    with st.expander(f"{activity['type']}: {activity['detail']}"): st.write(f"⏱️ {activity['time_string']}")
            else: st.info("No recent activity.")
        elif selected_module == "Inventory & Billing": inventory_module(entity_id, ai_models)
        elif selected_module == "HR Tools": hr_module(entity_id, ai_models)
        elif selected_module == "Project Manager": project_module(entity_id, ai_models)
        elif selected_module == "Document Generator": document_module(entity_id, ai_models)
        elif selected_module == "Market Analysis Tool": market_analysis_module(entity_id, ai_models)
        elif selected_module == "Market Doubt Assistant (AI Chatbot)": chatbot_module(entity_id, ai_models)
        elif selected_module == "Investor & Agent Dashboards": investor_dashboard(entity_id, ai_models)
        elif selected_module == "Govt/Private Schemes & News Alerts": schemes_module(entity_id, ai_models)
        elif selected_module == "Opportunity Director": opportunities_module(entity_id, ai_models)
        elif selected_module == "Voice Navigation": voice_navigation(entity_id, ai_models)
        elif selected_module == "Pitching Helper": pitching_helper(entity_id, ai_models)
        elif selected_module == "Strategy Generator": strategy_generator(entity_id, ai_models)
        elif selected_module == "Hiring Helper": hiring_helper(entity_id, ai_models)
        elif selected_module == "Tax & GST Filing": tax_module(entity_id, ai_models)
        elif selected_module == "IPO & Cap Table Management": ipo_module(entity_id, ai_models)
        elif selected_module == "Legal, CA & Insurance Marketplace": legal_marketplace(entity_id, ai_models)
        elif selected_module == "Enterprise Intelligence Dashboards": enterprise_intelligence(entity_id, ai_models)
        elif selected_module == "AI Market Forecasting": market_forecasting(entity_id, ai_models)
        elif selected_module == "Messaging": chat_module(entity_id, ai_models) 
        else: st.write(f"Module {selected_module} selected.")

    elif user_type == 'investor':
        st.sidebar.markdown("---"); st.sidebar.subheader("Investor Tools")
        investor_portal(entity_id, ai_models)
        
    elif user_type == 'service_provider':
        st.sidebar.markdown("---"); st.sidebar.subheader(f"{st.session_state.service_type.capitalize()} Tools")
        service_provider_portal(entity_id, ai_models)

    st.sidebar.markdown("---"); st.sidebar.markdown("### About GrowBis\n- **Version**: 2.1 (Multi-Role)")

if __name__ == "__main__":
    main()
